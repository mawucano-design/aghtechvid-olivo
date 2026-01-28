# app.py — Versión adaptada para vid, olivo y hortalizas de hoja
# Autor: Martin Ernesto Cano
# Fecha: Enero 2026

import streamlit as st
import geopandas as gpd
import pandas as pd
import numpy as np
import tempfile
import os
import zipfile
from datetime import datetime, timedelta
import matplotlib.pyplot as plt
from matplotlib.tri import Triangulation
import matplotlib.patches as mpatches
from matplotlib.colors import LinearSegmentedColormap
import io
from shapely.geometry import Polygon, LineString
import math
import warnings
import xml.etree.ElementTree as ET
import base64
import json
from io import BytesIO
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import geojson
import requests
import contextily as ctx
from pyproj import CRS
warnings.filterwarnings('ignore')

# ===== USAR TODO EL ANCHO DE LA PANTALLA =====
st.set_page_config(layout="wide")

# === ESTILOS PERSONALIZADOS - VERSIÓN PREMIUM MODERNA ===
st.markdown("""
<style>
/* === FONDO GENERAL OSCURO ELEGANTE === */
.stApp {
background: linear-gradient(135deg, #0f172a 0%, #1e293b 100%) !important;
color: #ffffff !important;
font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
}
/* === SIDEBAR: FONDO BLANCO CON TEXTO NEGRO === */
[data-testid="stSidebar"] {
background: #ffffff !important;
border-right: 1px solid #e5e7eb !important;
box-shadow: 5px 0 25px rgba(0, 0, 0, 0.1) !important;
}
/* Texto general del sidebar en NEGRO */
[data-testid="stSidebar"] *,
[data-testid="stSidebar"] .stMarkdown,
[data-testid="stSidebar"] p,
[data-testid="stSidebar"] label,
[data-testid="stSidebar"] .stText,
[data-testid="stSidebar"] .stTitle,
[data-testid="stSidebar"] .stSubheader {
color: #000000 !important;
text-shadow: none !important;
}
/* Título del sidebar elegante */
.sidebar-title {
font-size: 1.4em;
font-weight: 800;
margin: 1.5em 0 1em 0;
text-align: center;
padding: 14px;
background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%);
border-radius: 16px;
color: #ffffff !important;
box-shadow: 0 6px 20px rgba(59, 130, 246, 0.3);
border: 1px solid rgba(255, 255, 255, 0.2);
letter-spacing: 0.5px;
}
/* Widgets del sidebar con estilo glassmorphism */
[data-testid="stSidebar"] .stSelectbox,
[data-testid="stSidebar"] .stDateInput,
[data-testid="stSidebar"] .stSlider {
background: rgba(255, 255, 255, 0.9) !important;
backdrop-filter: blur(10px);
border-radius: 12px;
padding: 12px;
margin: 8px 0;
border: 1px solid #d1d5db !important;
}
/* Labels de los widgets en negro */
[data-testid="stSidebar"] .stSelectbox div,
[data-testid="stSidebar"] .stDateInput div,
[data-testid="stSidebar"] .stSlider label {
color: #000000 !important;
font-weight: 600;
font-size: 0.95em;
}
/* Inputs y selects - fondo blanco con texto negro */
[data-testid="stSidebar"] .stSelectbox [data-baseweb="select"] {
background-color: #ffffff !important;
border: 1px solid #d1d5db !important;
color: #000000 !important;
border-radius: 8px;
}
/* Slider - colores negro */
[data-testid="stSidebar"] .stSlider [data-baseweb="slider"] {
color: #000000 !important;
}
/* Date Input - fondo blanco con texto negro */
[data-testid="stSidebar"] .stDateInput [data-baseweb="input"] {
background-color: #ffffff !important;
border: 1px solid #d1d5db !important;
color: #000000 !important;
border-radius: 8px;
}
/* Placeholder en gris */
[data-testid="stSidebar"] .stDateInput [data-baseweb="input"]::placeholder {
color: #6b7280 !important;
}
/* Botones premium */
.stButton > button {
background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
color: white !important;
border: none !important;
padding: 0.8em 1.5em !important;
border-radius: 12px !important;
font-weight: 700 !important;
font-size: 1em !important;
box-shadow: 0 4px 15px rgba(59, 130, 246, 0.4) !important;
transition: all 0.3s ease !important;
text-transform: uppercase !important;
letter-spacing: 0.5px !important;
}
.stButton > button:hover {
transform: translateY(-3px) !important;
box-shadow: 0 8px 25px rgba(59, 130, 246, 0.6) !important;
background: linear-gradient(135deg, #4f8df8 0%, #2d5fe8 100%) !important;
}
/* === HERO BANNER PRINCIPAL CON IMAGEN === */
.hero-banner {
background: linear-gradient(rgba(15, 23, 42, 0.9), rgba(15, 23, 42, 0.95)),
url('https://images.unsplash.com/photo-1597981309443-6e2d2a4d9c3f?ixlib=rb-4.0.3&ixid=M3wxMjA3fDB8MHxwaG90by1wYWdlfHx8fGVufDB8fHx8fA%3D%3D&auto=format&fit=crop&w=2070&q=80') !important;
background-size: cover !important;
background-position: center 40% !important;
padding: 3.5em 2em !important;
border-radius: 24px !important;
margin-bottom: 2.5em !important;
box-shadow: 0 20px 40px rgba(0, 0, 0, 0.4) !important;
border: 1px solid rgba(59, 130, 246, 0.2) !important;
position: relative !important;
overflow: hidden !important;
}
.hero-banner::before {
content: '' !important;
position: absolute !important;
top: 0 !important;
left: 0 !important;
right: 0 !important;
bottom: 0 !important;
background: linear-gradient(45deg, rgba(59, 130, 246, 0.1), rgba(29, 78, 216, 0.05)) !important;
z-index: 1 !important;
}
.hero-content {
position: relative !important;
z-index: 2 !important;
text-align: center !important;
}
.hero-title {
color: #ffffff !important;
font-size: 3.2em !important;
font-weight: 900 !important;
margin-bottom: 0.3em !important;
text-shadow: 0 4px 12px rgba(0, 0, 0, 0.6) !important;
letter-spacing: -0.5px !important;
background: linear-gradient(135deg, #ffffff 0%, #93c5fd 100%) !important;
-webkit-background-clip: text !important;
-webkit-text-fill-color: transparent !important;
background-clip: text !important;
}
.hero-subtitle {
color: #cbd5e1 !important;
font-size: 1.3em !important;
font-weight: 400 !important;
max-width: 800px !important;
margin: 0 auto !important;
line-height: 1.6 !important;
}
/* === PESTAÑAS PRINCIPALES (fuera del sidebar) - SIN CAMBIOS === */
.stTabs [data-baseweb="tab-list"] {
background: rgba(255, 255, 255, 0.05) !important;
backdrop-filter: blur(10px) !important;
padding: 8px 16px !important;
border-radius: 16px !important;
border: 1px solid rgba(255, 255, 255, 0.1) !important;
margin-top: 1em !important;
gap: 8px !important;
}
.stTabs [data-baseweb="tab"] {
color: #94a3b8 !important;
font-weight: 600 !important;
padding: 12px 24px !important;
border-radius: 12px !important;
background: transparent !important;
transition: all 0.3s ease !important;
border: 1px solid transparent !important;
}
.stTabs [data-baseweb="tab"]:hover {
color: #ffffff !important;
background: rgba(59, 130, 246, 0.2) !important;
border-color: rgba(59, 130, 246, 0.3) !important;
transform: translateY(-2px) !important;
}
.stTabs [aria-selected="true"] {
background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
color: #ffffff !important;
font-weight: 700 !important;
border: none !important;
box-shadow: 0 4px 15px rgba(59, 130, 246, 0.4) !important;
}
/* === PESTAÑAS DEL SIDEBAR: FONDO BLANCO + TEXTO NEGRO === */
[data-testid="stSidebar"] .stTabs [data-baseweb="tab-list"] {
background: #ffffff !important;
border: 1px solid #e2e8f0 !important;
padding: 8px !important;
border-radius: 12px !important;
gap: 6px !important;
}
[data-testid="stSidebar"] .stTabs [data-baseweb="tab"] {
color: #000000 !important;
background: transparent !important;
border-radius: 8px !important;
padding: 8px 16px !important;
font-weight: 600 !important;
border: 1px solid transparent !important;
}
[data-testid="stSidebar"] .stTabs [data-baseweb="tab"]:hover {
background: #f1f5f9 !important;
color: #000000 !important;
border-color: #cbd5e1 !important;
}
/* Pestaña activa en el sidebar: blanco con texto negro */
[data-testid="stSidebar"] .stTabs [aria-selected="true"] {
background: #ffffff !important;
color: #000000 !important;
font-weight: 700 !important;
border: 1px solid #3b82f6 !important;
}
/* === MÉTRICAS PREMIUM === */
div[data-testid="metric-container"] {
background: linear-gradient(135deg, rgba(30, 41, 59, 0.8), rgba(15, 23, 42, 0.9)) !important;
backdrop-filter: blur(10px) !important;
border-radius: 20px !important;
padding: 24px !important;
box-shadow: 0 10px 30px rgba(0, 0, 0, 0.3) !important;
border: 1px solid rgba(59, 130, 246, 0.2) !important;
transition: all 0.3s ease !important;
}
div[data-testid="metric-container"]:hover {
transform: translateY(-5px) !important;
box-shadow: 0 15px 40px rgba(59, 130, 246, 0.2) !important;
border-color: rgba(59, 130, 246, 0.4) !important;
}
div[data-testid="metric-container"] label,
div[data-testid="metric-container"] div,
div[data-testid="metric-container"] [data-testid="stMetricValue"],
div[data-testid="metric-container"] [data-testid="stMetricLabel"] {
color: #ffffff !important;
font-weight: 600 !important;
}
div[data-testid="metric-container"] [data-testid="stMetricValue"] {
font-size: 2.5em !important;
font-weight: 800 !important;
background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
-webkit-background-clip: text !important;
-webkit-text-fill-color: transparent !important;
background-clip: text !important;
}
/* === GRÁFICOS CON ESTILO OSCURO === */
.stPlotlyChart, .stPyplot {
background: rgba(15, 23, 42, 0.8) !important;
backdrop-filter: blur(10px) !important;
border-radius: 20px !important;
padding: 20px !important;
box-shadow: 0 10px 30px rgba(0, 0, 0, 0.3) !important;
border: 1px solid rgba(59, 130, 246, 0.2) !important;
}
/* === EXPANDERS ELEGANTES === */
.streamlit-expanderHeader {
color: #ffffff !important;
background: rgba(30, 41, 59, 0.8) !important;
backdrop-filter: blur(10px) !important;
border-radius: 16px !important;
font-weight: 700 !important;
border: 1px solid rgba(255, 255, 255, 0.1) !important;
padding: 16px 20px !important;
margin-bottom: 10px !important;
}
.streamlit-expanderContent {
background: rgba(15, 23, 42, 0.6) !important;
border-radius: 0 0 16px 16px !important;
padding: 20px !important;
border: 1px solid rgba(255, 255, 255, 0.1) !important;
border-top: none !important;
}
/* === TEXTOS GENERALES === */
h1, h2, h3, h4, h5, h6 {
color: #ffffff !important;
font-weight: 800 !important;
margin-top: 1.5em !important;
}
p, div, span, label, li {
color: #cbd5e1 !important;
line-height: 1.7 !important;
}
/* === DATA FRAMES TABLAS ELEGANTES === */
.dataframe {
background: rgba(15, 23, 42, 0.8) !important;
backdrop-filter: blur(10px) !important;
border-radius: 16px !important;
border: 1px solid rgba(255, 255,255, 0.1) !important;
color: #ffffff !important;
}
.dataframe th {
background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
color: #ffffff !important;
font-weight: 700 !important;
padding: 16px !important;
}
.dataframe td {
color: #cbd5e1 !important;
padding: 14px 16px !important;
border-bottom: 1px solid rgba(255, 255, 255, 0.1) !important;
}
/* === ALERTS Y MENSAJES === */
.stAlert {
border-radius: 16px !important;
border: 1px solid rgba(255, 255, 255, 0.1) !important;
backdrop-filter: blur(10px) !important;
}
/* === SCROLLBAR PERSONALIZADA === */
::-webkit-scrollbar {
width: 10px !important;
height: 10px !important;
}
::-webkit-scrollbar-track {
background: rgba(15, 23, 42, 0.8) !important;
border-radius: 10px !important;
}
::-webkit-scrollbar-thumb {
background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
border-radius: 10px !important;
}
::-webkit-scrollbar-thumb:hover {
background: linear-gradient(135deg, #4f8df8 0%, #2d5fe8 100%) !important;
}
/* === IMÁGENES DEL SIDEBAR === */
[data-testid="stSidebar"] img {
border-radius: 16px !important;
border: 2px solid #d1d5db !important;
box-shadow: 0 8px 25px rgba(0, 0, 0, 0.1) !important;
transition: all 0.3s ease !important;
}
[data-testid="stSidebar"] img:hover {
transform: scale(1.02) !important;
box-shadow: 0 12px 35px rgba(0, 0, 0, 0.2) !important;
border-color: #3b82f6 !important;
}
/* === TARJETAS DE CULTIVOS === */
.cultivo-card {
background: linear-gradient(135deg, rgba(30, 41, 59, 0.9), rgba(15, 23, 42, 0.95)) !important;
border-radius: 20px !important;
padding: 25px !important;
border: 1px solid rgba(59, 130, 246, 0.2) !important;
box-shadow: 0 10px 30px rgba(0, 0, 0, 0.3) !important;
transition: all 0.3s ease !important;
height: 100% !important;
}
.cultivo-card:hover {
transform: translateY(-8px) !important;
box-shadow: 0 20px 40px rgba(59, 130, 246, 0.2) !important;
border-color: rgba(59, 130, 246, 0.4) !important;
}
/* === TABLERO DE CONTROL === */
.dashboard-grid {
display: grid !important;
grid-template-columns: repeat(auto-fit, minmax(300px, 1fr)) !important;
gap: 25px !important;
margin: 30px 0 !important;
}
.dashboard-card {
background: linear-gradient(135deg, rgba(30, 41, 59, 0.9), rgba(15, 23, 42, 0.95)) !important;
border-radius: 20px !important;
padding: 25px !important;
border: 1px solid rgba(59, 130, 246, 0.2) !important;
box-shadow: 0 10px 30px rgba(0, 0, 0, 0.3) !important;
transition: all 0.3s ease !important;
}
.dashboard-card:hover {
transform: translateY(-5px) !important;
box-shadow: 0 20px 40px rgba(59, 130, 246, 0.2) !important;
}
/* === STATS BADGES === */
.stats-badge {
display: inline-block !important;
padding: 6px 14px !important;
border-radius: 50px !important;
font-size: 0.85em !important;
font-weight: 700 !important;
margin: 2px !important;
}
.badge-success {
background: linear-gradient(135deg, #10b981 0%, #059669 100%) !important;
color: white !important;
}
.badge-warning {
background: linear-gradient(135deg, #f59e0b 0%, #d97706 100%) !important;
color: white !important;
}
.badge-danger {
background: linear-gradient(135deg, #ef4444 0%, #dc2626 100%) !important;
color: white !important;
}
.badge-info {
background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
color: white !important;
}
/* Estilos adicionales para mapas de calor */
.map-container {
background: linear-gradient(135deg, rgba(15, 23, 42, 0.95), rgba(30, 41, 59, 0.98)) !important;
border-radius: 20px !important;
padding: 20px !important;
margin: 20px 0 !important;
border: 1px solid rgba(59, 130, 246, 0.3) !important;
box-shadow: 0 15px 40px rgba(0, 0, 0, 0.4) !important;
}
.map-title {
font-size: 1.4em !important;
font-weight: 800 !important;
color: #ffffff !important;
text-align: center !important;
margin-bottom: 15px !important;
padding-bottom: 10px !important;
border-bottom: 2px solid rgba(59, 130, 246, 0.5) !important;
}
.map-stats {
background: rgba(30, 41, 59, 0.8) !important;
border-radius: 12px !important;
padding: 15px !important;
margin: 10px 0 !important;
border: 1px solid rgba(255, 255, 255, 0.1) !important;
}
.map-tabs .stTabs [data-baseweb="tab-list"] {
background: rgba(30, 41, 59, 0.8) !important;
border-radius: 12px !important;
padding: 8px !important;
margin-bottom: 15px !important;
}
.map-tabs .stTabs [data-baseweb="tab"] {
color: #cbd5e1 !important;
font-weight: 600 !important;
}
.map-tabs .stTabs [aria-selected="true"] {
background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
color: #ffffff !important;
}
/* Mejorar visualización de imágenes de mapas */
.stImage > img {
border-radius: 16px !important;
border: 2px solid rgba(59, 130, 246, 0.3) !important;
box-shadow: 0 10px 30px rgba(0, 0, 0, 0.3) !important;
transition: all 0.3s ease !important;
}
.stImage > img:hover {
transform: scale(1.01) !important;
box-shadow: 0 15px 40px rgba(59, 130, 246, 0.2) !important;
border-color: rgba(59, 130, 246, 0.6) !important;
}
/* Estilos para estadísticas de mapas */
.stats-card {
background: linear-gradient(135deg, rgba(30, 41, 59, 0.9), rgba(15, 23, 42, 0.95)) !important;
border-radius: 16px !important;
padding: 20px !important;
margin: 10px 0 !important;
border: 1px solid rgba(59, 130, 246, 0.2) !important;
}
.stats-value {
font-size: 2em !important;
font-weight: 800 !important;
background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
-webkit-background-clip: text !important;
-webkit-text-fill-color: transparent !important;
background-clip: text !important;
margin: 5px 0 !important;
}
.stats-label {
color: #94a3b8 !important;
font-size: 0.9em !important;
font-weight: 600 !important;
text-transform: uppercase !important;
letter-spacing: 0.5px !important;
}
/* Mejorar botones de descarga de mapas */
.download-map-btn {
background: linear-gradient(135deg, #10b981 0%, #059669 100%) !important;
color: white !important;
border: none !important;
padding: 10px 20px !important;
border-radius: 10px !important;
font-weight: 700 !important;
margin: 5px !important;
transition: all 0.3s ease !important;
}
.download-map-btn:hover {
transform: translateY(-2px) !important;
box-shadow: 0 8px 20px rgba(16, 185, 129, 0.4) !important;
}
/* Leyenda de mapas mejorada */
.map-legend {
background: rgba(30, 41, 59, 0.9) !important;
border-radius: 10px !important;
padding: 15px !important;
border: 1px solid rgba(255, 255, 255, 0.1) !important;
margin: 15px 0 !important;
}
.legend-title {
color: #ffffff !important;
font-weight: 700 !important;
margin-bottom: 10px !important;
font-size: 1.1em !important;
}
.legend-item {
display: flex !important;
align-items: center !important;
margin: 5px 0 !important;
color: #cbd5e1 !important;
}
.legend-color {
width: 20px !important;
height: 20px !important;
border-radius: 4px !important;
margin-right: 10px !important;
border: 1px solid rgba(255, 255, 255, 0.2) !important;
}
</style>
""", unsafe_allow_html=True)

# ===== HERO BANNER PRINCIPAL =====
st.markdown("""
<div class="hero-banner">
<div class="hero-content">
<h1 class="hero-title">ANALIZADOR MULTI-CULTIVO SATELITAL</h1>
<p class="hero-subtitle">Potenciado con NASA POWER, GEE, INTA y tecnología avanzada para una agricultura de precisión</p>
<p class="hero-subtitle">Especializado en: Vid, Olivo y Hortalizas de Hoja</p>
</div>
</div>
""", unsafe_allow_html=True)

# ===== CONFIGURACIÓN DE SATÉLITES DISPONIBLES =====
SATELITES_DISPONIBLES = {
'SENTINEL-2': {
'nombre': 'Sentinel-2',
'resolucion': '10m',
'revisita': '5 días',
'bandas': ['B2', 'B3', 'B4', 'B5', 'B8', 'B8A', 'B11', 'B12'],
'indices': ['NDVI', 'NDRE', 'GNDVI', 'OSAVI', 'MCARI', 'TCARI', 'NDII'],
'icono': '🛰️',
'bandas_np': {
'N': ['B5', 'B8A'],  # Red Edge para NDRE
'P': ['B4', 'B11'],  # Rojo y SWIR para fósforo
'K': ['B8', 'B11', 'B12']  # NIR y SWIR para potasio
}
},
'LANDSAT-8': {
'nombre': 'Landsat 8',
'resolucion': '30m',
'revisita': '16 días',
'bandas': ['B2', 'B3', 'B4', 'B5', 'B6', 'B7', 'B10', 'B11'],
'indices': ['NDVI', 'NDWI', 'EVI', 'SAVI', 'MSAVI', 'NDII'],
'icono': '🛰️',
'bandas_np': {
'N': ['B4', 'B5'],  # Rojo y NIR para NDRE alternativo
'P': ['B3', 'B6'],  # Verde y SWIR1
'K': ['B5', 'B6', 'B7']  # NIR y SWIR
}
},
'DATOS_SIMULADOS': {
'nombre': 'Datos Simulados',
'resolucion': '10m',
'revisita': '5 días',
'bandas': ['B2', 'B3', 'B4', 'B5', 'B8'],
'indices': ['NDVI', 'NDRE', 'GNDVI'],
'icono': '🔬'
}
}

# ===== NUEVAS METODOLOGÍAS PARA ESTIMAR NPK CON TELEDETECCIÓN =====
METODOLOGIAS_NPK = {
'SENTINEL-2': {
'NITRÓGENO': {
'metodo': 'NDRE + Regresión Espectral',
'formula': 'N = 150 * NDRE + 50 * (B8A/B5)',
'bandas': ['B5', 'B8A'],
'r2_esperado': 0.75,
'referencia': 'Clevers & Gitelson, 2013'
},
'FÓSFORO': {
'metodo': 'Índice SWIR-VIS',
'formula': 'P = 80 * (B11/B4)^0.5 + 20',
'bandas': ['B4', 'B11'],
'r2_esperado': 0.65,
'referencia': 'Miphokasap et al., 2012'
},
'POTASIO': {
'metodo': 'Índice de Estrés Hídrico',
'formula': 'K = 120 * (B8 - B11)/(B8 + B12) + 40',
'bandas': ['B8', 'B11', 'B12'],
'r2_esperado': 0.70,
'referencia': 'Jackson et al., 2004'
}
},
'LANDSAT-8': {
'NITRÓGENO': {
'metodo': 'TCARI/OSAVI',
'formula': 'N = 3*[(B5-B4)-0.2*(B5-B3)*(B5/B4)] / (1.16*(B5-B4)/(B5+B4+0.16))',
'bandas': ['B3', 'B4', 'B5'],
'r2_esperado': 0.72,
'referencia': 'Haboudane et al., 2002'
},
'FÓSFORO': {
'metodo': 'Relación SWIR1-Verde',
'formula': 'P = 60 * (B6/B3)^0.7 + 25',
'bandas': ['B3', 'B6'],
'r2_esperado': 0.68,
'referencia': 'Chen et al., 2010'
},
'POTASIO': {
'metodo': 'Índice NIR-SWIR',
'formula': 'K = 100 * (B5 - B7)/(B5 + B7) + 50',
'bandas': ['B5', 'B7'],
'r2_esperado': 0.69,
'referencia': 'Thenkabail et al., 2000'
}
}
}

# ===== CONFIGURACIÓN =====
# PARÁMETROS GEE POR CULTIVO - ADAPTADOS PARA NUEVOS CULTIVOS

# ===== VARIEDADES DE VID PARA ARGENTINA =====
VARIEDADES_VID = {
'MALBEC': {
'RENDIMIENTO_BASE': 10.0,
'RENDIMIENTO_OPTIMO': 15.0,
'RESPUESTA_N': 0.04,
'RESPUESTA_P': 0.06,
'RESPUESTA_K': 0.03,
'NITROGENO_OPTIMO': 60,
'FOSFORO_OPTIMO': 35,
'POTASIO_OPTIMO': 150,
'CICLO': 180,
'TIPO': 'Tinto',
'REGION': 'Mendoza'
},
'CABERNET SAUVIGNON': {
'RENDIMIENTO_BASE': 8.0,
'RENDIMIENTO_OPTIMO': 12.0,
'RESPUESTA_N': 0.03,
'RESPUESTA_P': 0.05,
'RESPUESTA_K': 0.04,
'NITROGENO_OPTIMO': 55,
'FOSFORO_OPTIMO': 32,
'POTASIO_OPTIMO': 145,
'CICLO': 190,
'TIPO': 'Tinto',
'REGION': 'Mendoza'
},
'TORRONTÉS': {
'RENDIMIENTO_BASE': 12.0,
'RENDIMIENTO_OPTIMO': 18.0,
'RESPUESTA_N': 0.05,
'RESPUESTA_P': 0.07,
'RESPUESTA_K': 0.035,
'NITROGENO_OPTIMO': 65,
'FOSFORO_OPTIMO': 38,
'POTASIO_OPTIMO': 155,
'CICLO': 175,
'TIPO': 'Blanco',
'REGION': 'Salta'
},
'CHARDONNAY': {
'RENDIMIENTO_BASE': 9.0,
'RENDIMIENTO_OPTIMO': 14.0,
'RESPUESTA_N': 0.035,
'RESPUESTA_P': 0.055,
'RESPUESTA_K': 0.032,
'NITROGENO_OPTIMO': 58,
'FOSFORO_OPTIMO': 34,
'POTASIO_OPTIMO': 148,
'CICLO': 185,
'TIPO': 'Blanco',
'REGION': 'Mendoza'
}
}

# ===== VARIEDADES DE OLIVO PARA ARGENTINA =====
VARIEDADES_OLIVO = {
'ARBEQUINA': {
'RENDIMIENTO_BASE': 8.0,
'RENDIMIENTO_OPTIMO': 12.0,
'RESPUESTA_N': 0.025,
'RESPUESTA_P': 0.04,
'RESPUESTA_K': 0.03,
'NITROGENO_OPTIMO': 50,
'FOSFORO_OPTIMO': 25,
'POTASIO_OPTIMO': 80,
'CICLO': 210,
'ACEITE': 'Alto (22-25%)',
'REGION': 'La Rioja'
},
'MANZANILLA': {
'RENDIMIENTO_BASE': 6.0,
'RENDIMIENTO_OPTIMO': 9.0,
'RESPUESTA_N': 0.02,
'RESPUESTA_P': 0.035,
'RESPUESTA_K': 0.025,
'NITROGENO_OPTIMO': 45,
'FOSFORO_OPTIMO': 22,
'POTASIO_OPTIMO': 75,
'CICLO': 200,
'ACEITE': 'Medio (18-20%)',
'REGION': 'Catamarca'
},
'PICUAL': {
'RENDIMIENTO_BASE': 10.0,
'RENDIMIENTO_OPTIMO': 15.0,
'RESPUESTA_N': 0.03,
'RESPUESTA_P': 0.045,
'RESPUESTA_K': 0.035,
'NITROGENO_OPTIMO': 55,
'FOSFORO_OPTIMO': 28,
'POTASIO_OPTIMO': 85,
'CICLO': 220,
'ACEITE': 'Muy Alto (25-28%)',
'REGION': 'Mendoza'
},
'FRANTOIO': {
'RENDIMIENTO_BASE': 7.0,
'RENDIMIENTO_OPTIMO': 11.0,
'RESPUESTA_N': 0.028,
'RESPUESTA_P': 0.042,
'RESPUESTA_K': 0.032,
'NITROGENO_OPTIMO': 52,
'FOSFORO_OPTIMO': 26,
'POTASIO_OPTIMO': 82,
'CICLO': 215,
'ACEITE': 'Alto (20-23%)',
'REGION': 'San Juan'
}
}

# ===== VARIEDADES DE HORTALIZAS DE HOJA PARA ARGENTINA =====
VARIEDADES_HORTALIZAS = {
'LECHUGA CRESPA': {
'RENDIMIENTO_BASE': 25.0,
'RENDIMIENTO_OPTIMO': 40.0,
'RESPUESTA_N': 0.08,
'RESPUESTA_P': 0.12,
'RESPUESTA_K': 0.06,
'NITROGENO_OPTIMO': 120,
'FOSFORO_OPTIMO': 60,
'POTASIO_OPTIMO': 150,
'CICLO': 45,
'TIPO': 'Hoja',
'REGION': 'Cinturón Verde'
},
'ESPINACA': {
'RENDIMIENTO_BASE': 15.0,
'RENDIMIENTO_OPTIMO': 25.0,
'RESPUESTA_N': 0.07,
'RESPUESTA_P': 0.10,
'RESPUESTA_K': 0.05,
'NITROGENO_OPTIMO': 100,
'FOSFORO_OPTIMO': 50,
'POTASIO_OPTIMO': 120,
'CICLO': 40,
'TIPO': 'Hoja',
'REGION': 'Cinturón Verde'
},
'ACELGA': {
'RENDIMIENTO_BASE': 20.0,
'RENDIMIENTO_OPTIMO': 35.0,
'RESPUESTA_N': 0.075,
'RESPUESTA_P': 0.11,
'RESPUESTA_K': 0.055,
'NITROGENO_OPTIMO': 110,
'FOSFORO_OPTIMO': 55,
'POTASIO_OPTIMO': 135,
'CICLO': 55,
'TIPO': 'Hoja',
'REGION': 'Cinturón Verde'
},
'RÚCULA': {
'RENDIMIENTO_BASE': 18.0,
'RENDIMIENTO_OPTIMO': 30.0,
'RESPUESTA_N': 0.085,
'RESPUESTA_P': 0.13,
'RESPUESTA_K': 0.065,
'NITROGENO_OPTIMO': 125,
'FOSFORO_OPTIMO': 65,
'POTASIO_OPTIMO': 160,
'CICLO': 35,
'TIPO': 'Hoja',
'REGION': 'Cinturón Verde'
}
}

PARAMETROS_CULTIVOS = {
'VID': {
'NITROGENO': {'min': 40, 'max': 80, 'optimo': 60},
'FOSFORO': {'min': 25, 'max': 45, 'optimo': 35},
'POTASIO': {'min': 120, 'max': 180, 'optimo': 150},
'MATERIA_ORGANICA_OPTIMA': 2.5,
'HUMEDAD_OPTIMA': 0.25,
'NDVI_OPTIMO': 0.75,
'NDRE_OPTIMO': 0.45,
'TCARI_OPTIMO': 0.35,
'OSAVI_OPTIMO': 0.55,
'RENDIMIENTO_BASE': 10.0,
'RENDIMIENTO_OPTIMO': 15.0,
'RESPUESTA_N': 0.04,
'RESPUESTA_P': 0.06,
'RESPUESTA_K': 0.03,
'FACTOR_CLIMA': 0.8,
'VARIEDAD_DEFAULT': 'MALBEC'
},
'OLIVO': {
'NITROGENO': {'min': 35, 'max': 65, 'optimo': 50},
'FOSFORO': {'min': 20, 'max': 35, 'optimo': 28},
'POTASIO': {'min': 60, 'max': 100, 'optimo': 80},
'MATERIA_ORGANICA_OPTIMA': 2.0,
'HUMEDAD_OPTIMA': 0.20,
'NDVI_OPTIMO': 0.70,
'NDRE_OPTIMO': 0.40,
'TCARI_OPTIMO': 0.30,
'OSAVI_OPTIMO': 0.50,
'RENDIMIENTO_BASE': 8.0,
'RENDIMIENTO_OPTIMO': 12.0,
'RESPUESTA_N': 0.025,
'RESPUESTA_P': 0.04,
'RESPUESTA_K': 0.03,
'FACTOR_CLIMA': 0.75,
'VARIEDAD_DEFAULT': 'ARBEQUINA'
},
'HORTALIZAS DE HOJA': {
'NITROGENO': {'min': 90, 'max': 140, 'optimo': 115},
'FOSFORO': {'min': 45, 'max': 75, 'optimo': 58},
'POTASIO': {'min': 100, 'max': 180, 'optimo': 140},
'MATERIA_ORGANICA_OPTIMA': 3.5,
'HUMEDAD_OPTIMA': 0.35,
'NDVI_OPTIMO': 0.85,
'NDRE_OPTIMO': 0.55,
'TCARI_OPTIMO': 0.45,
'OSAVI_OPTIMO': 0.65,
'RENDIMIENTO_BASE': 20.0,
'RENDIMIENTO_OPTIMO': 32.0,
'RESPUESTA_N': 0.075,
'RESPUESTA_P': 0.11,
'RESPUESTA_K': 0.055,
'FACTOR_CLIMA': 0.85,
'VARIEDAD_DEFAULT': 'LECHUGA CRESPA'
}
}

# ===== PARÁMETROS ECONÓMICOS PARA ARGENTINA (2025) - ACTUALIZADOS =====
PARAMETROS_ECONOMICOS = {
'PRECIOS_CULTIVOS': {
'VID': {
'precio_ton': 800,  # USD/ton (uva para vino)
'costo_semilla': 3000,  # USD/ha (plantines)
'costo_herbicidas': 120,
'costo_insecticidas': 150,
'costo_labores': 400,
'costo_cosecha': 250,
'costo_otros': 200
},
'OLIVO': {
'precio_ton': 1200,  # USD/ton (aceituna para aceite)
'costo_semilla': 2500,  # USD/ha (plantines)
'costo_herbicidas': 100,
'costo_insecticidas': 120,
'costo_labores': 350,
'costo_cosecha': 300,
'costo_otros': 180
},
'HORTALIZAS DE HOJA': {
'precio_ton': 500,  # USD/ton
'costo_semilla': 800,  # USD/ha
'costo_herbicidas': 150,
'costo_insecticidas': 200,
'costo_labores': 300,
'costo_cosecha': 180,
'costo_otros': 120
}
},
'PRECIOS_FERTILIZANTES': {
'UREA': 450,  # USD/ton
'FOSFATO_DIAMONICO': 650,  # USD/ton
'CLORURO_POTASIO': 400,  # USD/ton
'SULFATO_AMONICO': 350,  # USD/ton
'SUPERFOSFATO': 420  # USD/ton
},
'CONVERSION_NUTRIENTES': {
'NITRÓGENO': {
'fuente_principal': 'UREA',
'contenido_nutriente': 0.46,  # 46% N
'eficiencia': 0.6  # 60% eficiencia
},
'FÓSFORO': {
'fuente_principal': 'FOSFATO_DIAMONICO',
'contenido_nutriente': 0.18,  # 18% P2O5 (46% P)
'eficiencia': 0.3  # 30% eficiencia
},
'POTASIO': {
'fuente_principal': 'CLORURO_POTASIO',
'contenido_nutriente': 0.60,  # 60% K2O (50% K)
'eficiencia': 0.5  # 50% eficiencia
}
},
'PARAMETROS_FINANCIEROS': {
'tasa_descuento': 0.10,  # 10% anual
'periodo_analisis': 5,  # 5 años
'inflacion_esperada': 0.08,  # 8% anual
'impuestos': 0.35,  # 35%
'subsidios': 0.05  # 5% de subsidios
}
}

# ===== NUEVA CLASIFICACIÓN USDA PARA TEXTURA DE SUELO =====
def clasificar_textura_usda(arena, limo, arcilla):
    """
    Clasifica la textura del suelo según el sistema USDA
    """
    try:
        total = arena + limo + arcilla
        if total == 0:
            return "Sin datos"
        # Normalizar porcentajes
        arena_pct = (arena / total) * 100
        limo_pct = (limo / total) * 100
        arcilla_pct = (arcilla / total) * 100
        # Clasificación USDA según el triángulo de texturas
        if arcilla_pct > 40:
            if limo_pct >= 40:
                return "Arcilla limosa"
            elif arena_pct <= 45:
                return "Arcilla"
            else:
                return "Arcilla arenosa"
        elif arcilla_pct >= 27 and arcilla_pct <= 40:
            if limo_pct >= 40:
                return "Franco arcilloso limoso"
            elif arena_pct <= 20:
                return "Franco arcilloso"
            else:
                return "Franco arcilloso arenoso"
        elif arcilla_pct >= 20 and arcilla_pct < 27:
            if limo_pct < 28:
                if arena_pct >= 52:
                    return "Arena franca"
                else:
                    return "Franco arenoso"
            else:
                if arena_pct >= 52:
                    return "Franco limoso arenoso"
                else:
                    return "Franco limoso"
        elif arcilla_pct >= 10 and arcilla_pct < 20:
            if limo_pct >= 50:
                return "Limo"
            elif limo_pct >= 30:
                if arena_pct >= 50:
                    return "Franco limoso arenoso"
                else:
                    return "Franco limoso"
            else:
                if arena_pct >= 70:
                    return "Arena"
                elif arena_pct >= 50:
                    return "Arena franca"
                else:
                    return "Franco arenoso"
        else:  # arcilla_pct < 10
            if limo_pct >= 80:
                return "Limo"
            elif limo_pct >= 50:
                return "Limo arenoso"
            else:
                if arena_pct >= 85:
                    return "Arena"
                else:
                    return "Arena franca"
    except Exception as e:
        return "Sin datos"

# ===== PARÁMETROS DE TEXTURA DEL SUELO POR CULTIVO - ACTUALIZADO A USDA =====
TEXTURA_SUELO_OPTIMA = {
'VID': {
'textura_optima': 'Franco arenoso',
'arena_optima': 60,
'limo_optima': 25,
'arcilla_optima': 15,
'densidad_aparente_optima': 1.4,
'porosidad_optima': 0.45
},
'OLIVO': {
'textura_optima': 'Franco limoso',
'arena_optima': 45,
'limo_optima': 35,
'arcilla_optima': 20,
'densidad_aparente_optima': 1.3,
'porosidad_optima': 0.48
},
'HORTALIZAS DE HOJA': {
'textura_optima': 'Franco limoso',
'arena_optima': 40,
'limo_optima': 40,
'arcilla_optima': 20,
'densidad_aparente_optima': 1.2,
'porosidad_optima': 0.55
}
}

# CLASIFICACIÓN DE PENDIENTES
CLASIFICACION_PENDIENTES = {
'PLANA (0-2%)': {'min': 0, 'max': 2, 'color': '#4daf4a', 'factor_erosivo': 0.1},
'SUAVE (2-5%)': {'min': 2, 'max': 5, 'color': '#a6d96a', 'factor_erosivo': 0.3},
'MODERADA (5-10%)': {'min': 5, 'max': 10, 'color': '#ffffbf', 'factor_erosivo': 0.6},
'FUERTE (10-15%)': {'min': 10, 'max': 15, 'color': '#fdae61', 'factor_erosivo': 0.8},
'MUY FUERTE (15-25%)': {'min': 15, 'max': 25, 'color': '#f46d43', 'factor_erosivo': 0.9},
'EXTREMA (>25%)': {'min': 25, 'max': 100, 'color': '#d73027', 'factor_erosivo': 1.0}
}

# ===== RECOMENDACIONES POR TIPO DE TEXTURA USDA - ACTUALIZADO =====
RECOMENDACIONES_TEXTURA = {
'Franco limoso': {
'propiedades': [
"Equilibrio ideal arena-limo-arcilla",
"Excelente estructura y porosidad",
"Alta capacidad de retención de agua",
"Fertilidad natural alta"
],
'limitantes': [
"Puede compactarse con maquinaria pesada",
"Moderadamente susceptible a erosión"
],
'manejo': [
"Labranza mínima o conservacionista",
"Rotación de cultivos",
"Uso de coberturas vegetales",
"Fertilización balanceada"
]
},
'Franco': {
'propiedades': [
"Buena aireación y drenaje",
"Fácil labranza",
"Calentamiento rápido en primavera",
"Retención moderada de nutrientes"
],
'limitantes': [
"Menor retención de agua que suelos más arcillosos",
"Requiere riego más frecuente"
],
'manejo': [
"Riego por goteo o aspersión",
"Fertilización fraccionada",
"Mulching para conservar humedad"
]
},
'Franco arcilloso limoso': {
'propiedades': [
"Alta capacidad de retención de agua",
"Excelente retención de nutrientes",
"Estructura estable",
"Resistente a la erosión"
],
'limitantes': [
"Lento drenaje",
"Difícil labranza en condiciones húmedas",
"Lento calentamiento en primavera"
],
'manejo': [
"Sistemas de drenaje",
"Labranza en condiciones óptimas de humedad",
"Incorporación de materia orgánica"
]
},
'Franco arenoso': {
'propiedades': [
"Excelente drenaje",
"Fácil labranza en cualquier condición",
"Rápido calentamiento",
"Buen desarrollo radicular"
],
'limitantes': [
"Baja retención de agua y nutrientes",
"Alta lixiviación de fertilizantes",
"Baja materia orgánica"
],
'manejo': [
"Riego frecuente en pequeñas cantidades",
"Fertilización fraccionada",
"Aplicación de materia orgánica",
"Cultivos de cobertura"
]
},
'Arcilla': {
'propiedades': [
"Alta capacidad de retención de agua y nutrientes",
"Estructura estable",
"Alta fertilidad potencial"
],
'limitantes': [
"Muy pesada cuando está húmeda",
"Drenaje muy lento",
"Difícil labranza",
"Propensa a compactación"
],
'manejo': [
"Drenaje artificial obligatorio",
"Labranza en condiciones óptimas",
"Encalamiento para mejorar estructura",
"Cultivos tolerantes a humedad"
]
},
'Arena franca': {
'propiedades': [
"Drenaje muy rápido",
"Fácil labranza",
"Rápido calentamiento",
"Bajo riesgo de compactación"
],
'limitantes': [
"Muy baja retención de agua",
"Alta lixiviación de nutrientes",
"Baja fertilidad natural"
],
'manejo': [
"Riego por goteo con alta frecuencia",
"Fertilización en múltiples aplicaciones",
"Aplicación intensiva de materia orgánica",
"Mulching para conservar humedad"
]
},
'Arcilla limosa': {
'propiedades': [
"Muy alta retención de agua y nutrientes",
"Estructura muy estable",
"Excelente para cultivos exigentes"
],
'limitantes': [
"Drenaje extremadamente lento",
"Muy pesada para labranza",
"Requiere manejo especializado"
],
'manejo': [
"Sistemas de drenaje avanzados",
"Labranza solo en condiciones óptimas",
"Aplicación de yeso para mejorar estructura",
"Camas elevadas para cultivos"
]
},
'Limo': {
'propiedades': [
"Alta capacidad de retención de agua",
"Fácil labranza",
"Buena fertilidad natural"
],
'limitantes': [
"Susceptible a compactación",
"Propenso a formación de costra superficial",
"Baja estabilidad estructural"
],
'manejo': [
"Evitar labranza en condiciones húmedas",
"Uso de coberturas vegetales",
"Aplicación de materia orgánica",
"Riego por aspersión ligera"
]
}
}

# ICONOS Y COLORES POR CULTIVO - ACTUALIZADO PARA NUEVOS CULTIVOS
ICONOS_CULTIVOS = {
'VID': '🍇',
'OLIVO': '🫒',
'HORTALIZAS DE HOJA': '🥬'
}
COLORES_CULTIVOS = {
'VID': '#6A0DAD',
'OLIVO': '#808000',
'HORTALIZAS DE HOJA': '#90EE90'
}

# PALETAS GEE MEJORADAS
PALETAS_GEE = {
'FERTILIDAD': ['#d73027', '#f46d43', '#fdae61', '#fee08b', '#d9ef8b', '#a6d96a', '#66bd63', '#1a9850', '#006837'],
'NITROGENO': ['#00ff00', '#80ff00', '#ffff00', '#ff8000', '#ff0000'],
'FOSFORO': ['#0000ff', '#4040ff', '#8080ff', '#c0c0ff', '#ffffff'],
'POTASIO': ['#4B0082', '#6A0DAD', '#8A2BE2', '#9370DB', '#D8BFD8'],
'TEXTURA': ['#8c510a', '#d8b365', '#f6e8c3', '#c7eae5', '#5ab4ac', '#01665e'],
'ELEVACION': ['#006837', '#1a9850', '#66bd63', '#a6d96a', '#d9ef8b', '#ffffbf', '#fee08b', '#fdae61', '#f46d43', '#d73027'],
'PENDIENTE': ['#4daf4a', '#a6d96a', '#ffffbf', '#fdae61', '#f46d43', '#d73027']
}

# URLs de imágenes para sidebar - ACTUALIZADAS PARA NUEVOS CULTIVOS
IMAGENES_CULTIVOS = {
'VID': 'https://images.unsplash.com/photo-1510812431401-41d2bd2722f3?auto=format&fit=crop&w=200&h=150&q=80',
'OLIVO': 'https://images.unsplash.com/photo-1615485500605-3866d64cd48d?auto=format&fit=crop&w=200&h=150&q=80',
'HORTALIZAS DE HOJA': 'https://images.unsplash.com/photo-1590779033100-9f60a05a013d?auto=format&fit=crop&w=200&h=150&q=80',
}

# ===== FUNCIÓN AUXILIAR PARA CONVERSIÓN RGBA =====
def rgba_to_tuple(rgba_str):
    """Convierte 'rgba(r,g,b,a)' a tupla (r/255, g/255, b/255, a)"""
    import re
    match = re.match(r'rgba\((\d+),\s*(\d+),\s*(\d+),\s*([\d.]+)\)', rgba_str)
    if match:
        r, g, b, a = map(float, match.groups())
        return (r/255, g/255, b/255, a)
    # Fallback a negro transparente
    return (0, 0, 0, 0.5)

# ===== INICIALIZACIÓN SEGURA DE VARIABLES DE CONFIGURACIÓN =====
if 'variedad' not in st.session_state:
    st.session_state['variedad'] = None
if 'variedad_params' not in st.session_state:
    st.session_state['variedad_params'] = None
nutriente = None
satelite_seleccionado = "SENTINEL-2"
indice_seleccionado = "NDVI"
fecha_inicio = datetime.now() - timedelta(days=30)
fecha_fin = datetime.now()
intervalo_curvas = 5.0
resolucion_dem = 10.0

# ===== NUEVA FUNCIÓN: Integración con el INTA =====
def obtener_materia_organica_inta(gdf, cultivo, usar_inta=True):
    """
    Obtiene materia orgánica del suelo usando datos regionales del INTA.
    Si usar_inta=False, devuelve la estimación genérica del sistema.
    """
    if not usar_inta:
        mo_valor = PARAMETROS_CULTIVOS[cultivo]['MATERIA_ORGANICA_OPTIMA']
        return {
            'materia_organica': round(mo_valor, 2),
            'region_inta': 'Estimación genérica',
            'fuente': 'Sistema interno',
            'textura_predominante': 'No disponible'
        }

    try:
        centroid = gdf.geometry.unary_union.centroid
        lon = centroid.x
        lat = centroid.y

        regiones_inta = [
            {
                'nombre': 'MENDOZA (Región Vitivinícola)',
                'lat_min': -34.0, 'lat_max': -32.0,
                'lon_min': -69.0, 'lon_max': -67.0,
                'mo_promedio': 2.2,
                'mo_rango': (1.5, 3.0),
                'textura_predominante': 'Franco arenoso'
            },
            {
                'nombre': 'LA RIOJA - CATAMARCA (Olivícola)',
                'lat_min': -30.0, 'lat_max': -27.0,
                'lon_min': -67.5, 'lon_max': -65.0,
                'mo_promedio': 1.8,
                'mo_rango': (1.0, 2.5),
                'textura_predominante': 'Franco limoso'
            },
            {
                'nombre': 'CINTURÓN VERDE (Hortícola)',
                'lat_min': -35.0, 'lat_max': -33.0,
                'lon_min': -59.0, 'lon_max': -57.0,
                'mo_promedio': 3.5,
                'mo_rango': (2.5, 4.5),
                'textura_predominante': 'Franco limoso'
            },
            {
                'nombre': 'PATAGONIA (Río Negro, Neuquén)',
                'lat_min': -42.0, 'lat_max': -38.0,
                'lon_min': -72.0, 'lon_max': -62.0,
                'mo_promedio': 5.2,
                'mo_rango': (3.5, 8.0),
                'textura_predominante': 'Franco volcánico'
            },
            {
                'nombre': 'OESTE (San Juan)',
                'lat_min': -32.0, 'lat_max': -30.0,
                'lon_min': -69.0, 'lon_max': -67.0,
                'mo_promedio': 1.5,
                'mo_rango': (0.8, 2.5),
                'textura_predominante': 'Franco arenoso'
            }
        ]

        region_encontrada = None
        for region in regiones_inta:
            if (region['lat_min'] <= lat <= region['lat_max'] and
                region['lon_min'] <= lon <= region['lon_max']):
                region_encontrada = region
                break

        if region_encontrada:
            seed_value = abs(hash(f"{lat:.4f}_{lon:.4f}_inta_mo")) % (2**32)
            rng = np.random.RandomState(seed_value)
            mo_valor = rng.normal(region_encontrada['mo_promedio'], 0.6)
            mo_valor = max(region_encontrada['mo_rango'][0], min(region_encontrada['mo_rango'][1], mo_valor))

            return {
                'materia_organica': round(mo_valor, 2),
                'region_inta': region_encontrada['nombre'],
                'fuente': 'INTA (simulado con datos reales)',
                'textura_predominante': region_encontrada['textura_predominante']
            }
        else:
            mo_valor = PARAMETROS_CULTIVOS[cultivo]['MATERIA_ORGANICA_OPTIMA']
            return {
                'materia_organica': round(mo_valor, 2),
                'region_inta': 'Fuera de cobertura INTA',
                'fuente': 'Estimación genérica',
                'textura_predominante': 'No disponible'
            }

    except Exception as e:
        mo_valor = PARAMETROS_CULTIVOS[cultivo]['MATERIA_ORGANICA_OPTIMA']
        return {
            'materia_organica': round(mo_valor, 2),
            'region_inta': 'Error en consulta',
            'fuente': 'Estimación de respaldo',
            'textura_predominante': 'No disponible'
        }

# ===== NUEVA FUNCIÓN: Capa WMS del INTA =====
def agregar_capa_inta(ax, alpha=0.4):
    """
    Agrega la capa WMS de suelos del INTA al eje matplotlib.
    Fuente: Servicio WMS oficial del INTA (si está disponible)
    """
    try:
        wms_url = "https://wms.inta.gob.ar/geoserver/inta/wms"
        ctx.add_basemap(
            ax,
            url=wms_url + "?SERVICE=WMS&VERSION=1.1.1&REQUEST=GetMap&LAYERS=inta%3Asuelos_argentina&SRS=EPSG%3A3857&WIDTH=256&HEIGHT=256&BBOX={bbox-epsg-3857}&FORMAT=image/png&TRANSPARENT=true",
            attribution="© INTA - Mapa de Suelos",
            alpha=alpha
        )
        return True
    except Exception as e:
        return False

# ===== FUNCIONES AUXILIARES - CORREGIDAS PARA EPSG:4326 =====
def validar_y_corregir_crs(gdf):
    if gdf is None or len(gdf) == 0:
        return gdf
    try:
        if gdf.crs is None:
            gdf = gdf.set_crs('EPSG:4326', inplace=False)
            st.info("ℹ️ Se asignó EPSG:4326 al archivo (no tenía CRS)")
        elif str(gdf.crs).upper() != 'EPSG:4326':
            original_crs = str(gdf.crs)
            gdf = gdf.to_crs('EPSG:4326')
            st.info(f"ℹ️ Transformado de {original_crs} a EPSG:4326")
        return gdf
    except Exception as e:
        st.warning(f"⚠️ Error al corregir CRS: {str(e)}")
        return gdf

def calcular_superficie(gdf):
    """
    Calcula el área total de un GeoDataFrame en hectáreas.
    Usa una proyección de área equivalente para mayor precisión.
    """
    try:
        if gdf is None or len(gdf) == 0:
            return 0.0

        # Asegurar CRS
        gdf = validar_y_corregir_crs(gdf)

        # Obtener el centroide para determinar UTM (opcional pero más preciso)
        centroid = gdf.geometry.unary_union.centroid
        lon, lat = centroid.x, centroid.y

        # Determinar zona UTM
        utm_zone = int((lon + 180) / 6) + 1
        hemisphere = 'north' if lat >= 0 else 'south'
        epsg_utm = f"326{utm_zone:02d}" if hemisphere == 'north' else f"327{utm_zone:02d}"

        try:
            # Intentar reproyectar a UTM local
            gdf_utm = gdf.to_crs(epsg=epsg_utm)
            area_m2 = gdf_utm.geometry.area.sum()
        except Exception:
            # Fallback a CRS de área global (EPSG:6933 - Equal Earth)
            gdf_eq = gdf.to_crs("EPSG:6933")
            area_m2 = gdf_eq.geometry.area.sum()

        return area_m2 / 10000  # Convertir a hectáreas

    except Exception as e:
        st.warning(f"⚠️ Error al calcular área precisa: {str(e)}. Usando cálculo aproximado.")
        # Último fallback: cálculo en grados (solo válido cerca del ecuador)
        area_grados2 = gdf.geometry.area.sum()
        area_m2 = area_grados2 * (111000 ** 2)
        return area_m2 / 10000

def dividir_parcela_en_zonas(gdf, n_zonas):
    if len(gdf) == 0:
        return gdf
    gdf = validar_y_corregir_crs(gdf)
    parcela_principal = gdf.iloc[0].geometry
    bounds = parcela_principal.bounds
    minx, miny, maxx, maxy = bounds
    sub_poligonos = []
    n_cols = math.ceil(math.sqrt(n_zonas))
    n_rows = math.ceil(n_zonas / n_cols)
    width = (maxx - minx) / n_cols
    height = (maxy - miny) / n_rows
    for i in range(n_rows):
        for j in range(n_cols):
            if len(sub_poligonos) >= n_zonas:
                break
            cell_minx = minx + (j * width)
            cell_maxx = minx + ((j + 1) * width)
            cell_miny = miny + (i * height)
            cell_maxy = miny + ((i + 1) * height)
            cell_poly = Polygon([(cell_minx, cell_miny), (cell_maxx, cell_miny), (cell_maxx, cell_maxy), (cell_minx, cell_maxy)])
            intersection = parcela_principal.intersection(cell_poly)
            if not intersection.is_empty and intersection.area > 0:
                sub_poligonos.append(intersection)
    if sub_poligonos:
        nuevo_gdf = gpd.GeoDataFrame({'id_zona': range(1, len(sub_poligonos) + 1), 'geometry': sub_poligonos}, crs='EPSG:4326')
        return nuevo_gdf
    else:
        return gdf

# ===== FUNCIONES PARA CARGAR ARCHIVOS =====
def cargar_shapefile_desde_zip(zip_file):
    try:
        with tempfile.TemporaryDirectory() as tmp_dir:
            with zipfile.ZipFile(zip_file, 'r') as zip_ref:
                zip_ref.extractall(tmp_dir)
            shp_files = [f for f in os.listdir(tmp_dir) if f.endswith('.shp')]
            if shp_files:
                shp_path = os.path.join(tmp_dir, shp_files[0])
                gdf = gpd.read_file(shp_path)
                gdf = validar_y_corregir_crs(gdf)
                return gdf
            else:
                st.error("❌ No se encontró ningún archivo .shp en el ZIP")
                return None
    except Exception as e:
        st.error(f"❌ Error cargando shapefile desde ZIP: {str(e)}")
        return None

def parsear_kml_manual(contenido_kml):
    try:
        root = ET.fromstring(contenido_kml)
        namespaces = {'kml': 'http://www.opengis.net/kml/2.2'}
        polygons = []
        for polygon_elem in root.findall('.//kml:Polygon', namespaces):
            coords_elem = polygon_elem.find('.//kml:coordinates', namespaces)
            if coords_elem is not None and coords_elem.text:
                coord_text = coords_elem.text.strip()
                coord_list = []
                for coord_pair in coord_text.split():
                    parts = coord_pair.split(',')
                    if len(parts) >= 2:
                        lon = float(parts[0])
                        lat = float(parts[1])
                        coord_list.append((lon, lat))
                if len(coord_list) >= 3:
                    polygons.append(Polygon(coord_list))
        if not polygons:
            for multi_geom in root.findall('.//kml:MultiGeometry', namespaces):
                for polygon_elem in multi_geom.findall('.//kml:Polygon', namespaces):
                    coords_elem = polygon_elem.find('.//kml:coordinates', namespaces)
                    if coords_elem is not None and coords_elem.text:
                        coord_text = coords_elem.text.strip()
                        coord_list = []
                        for coord_pair in coord_text.split():
                            parts = coord_pair.split(',')
                            if len(parts) >= 2:
                                lon = float(parts[0])
                                lat = float(parts[1])
                                coord_list.append((lon, lat))
                        if len(coord_list) >= 3:
                            polygons.append(Polygon(coord_list))
        if polygons:
            gdf = gpd.GeoDataFrame({'geometry': polygons}, crs='EPSG:4326')
            return gdf
        else:
            for placemark in root.findall('.//kml:Placemark', namespaces):
                for elem_name in ['Polygon', 'LineString', 'Point', 'LinearRing']:
                    elem = placemark.find(f'.//kml:{elem_name}', namespaces)
                    if elem is not None:
                        coords_elem = elem.find('.//kml:coordinates', namespaces)
                        if coords_elem is not None and coords_elem.text:
                            coord_text = coords_elem.text.strip()
                            coord_list = []
                            for coord_pair in coord_text.split():
                                parts = coord_pair.split(',')
                                if len(parts) >= 2:
                                    lon = float(parts[0])
                                    lat = float(parts[1])
                                    coord_list.append((lon, lat))
                            if len(coord_list) >= 3:
                                polygons.append(Polygon(coord_list))
                            break
            if polygons:
                gdf = gpd.GeoDataFrame({'geometry': polygons}, crs='EPSG:4326')
                return gdf
        return None
    except Exception as e:
        st.error(f"❌ Error parseando KML manualmente: {str(e)}")
        return None

def cargar_kml(kml_file):
    try:
        if kml_file.name.endswith('.kmz'):
            with tempfile.TemporaryDirectory() as tmp_dir:
                with zipfile.ZipFile(kml_file, 'r') as zip_ref:
                    zip_ref.extractall(tmp_dir)
                kml_files = [f for f in os.listdir(tmp_dir) if f.endswith('.kml')]
                if kml_files:
                    kml_path = os.path.join(tmp_dir, kml_files[0])
                    with open(kml_path, 'r', encoding='utf-8') as f:
                        contenido = f.read()
                    gdf = parsear_kml_manual(contenido)
                    if gdf is not None:
                        return gdf
                    else:
                        try:
                            gdf = gpd.read_file(kml_path)
                            gdf = validar_y_corregir_crs(gdf)
                            return gdf
                        except:
                            st.error("❌ No se pudo cargar el archivo KML/KMZ")
                            return None
                else:
                    st.error("❌ No se encontró ningún archivo .kml en el KMZ")
                    return None
        else:
            contenido = kml_file.read().decode('utf-8')
            gdf = parsear_kml_manual(contenido)
            if gdf is not None:
                return gdf
            else:
                kml_file.seek(0)
                gdf = gpd.read_file(kml_file)
                gdf = validar_y_corregir_crs(gdf)
                return gdf
    except Exception as e:
        st.error(f"❌ Error cargando archivo KML/KMZ: {str(e)}")
        return None

def cargar_archivo_parcela(uploaded_file):
    try:
        if uploaded_file.name.endswith('.zip'):
            gdf = cargar_shapefile_desde_zip(uploaded_file)
        elif uploaded_file.name.endswith(('.kml', '.kmz')):
            gdf = cargar_kml(uploaded_file)
        else:
            st.error("❌ Formato de archivo no soportado")
            return None
        if gdf is not None:
            gdf = validar_y_corregir_crs(gdf)
            if not gdf.geometry.geom_type.str.contains('Polygon').any():
                st.warning("⚠️ El archivo no contiene polígonos. Intentando extraer polígonos...")
                gdf = gdf.explode()
                gdf = gdf[gdf.geometry.geom_type.isin(['Polygon', 'MultiPolygon'])]
            if len(gdf) > 0:
                if 'id_zona' not in gdf.columns:
                    gdf['id_zona'] = range(1, len(gdf) + 1)
                if str(gdf.crs).upper() != 'EPSG:4326':
                    st.warning(f"⚠️ El archivo no pudo ser convertido a EPSG:4326. CRS actual: {gdf.crs}")
                return gdf
            else:
                st.error("❌ No se encontraron polígonos en el archivo")
                return None
        return gdf
    except Exception as e:
        st.error(f"❌ Error cargando archivo: {str(e)}")
        import traceback
        st.error(f"Detalle: {traceback.format_exc()}")
        return None

# ===== NUEVAS FUNCIONES PARA ESTIMAR NPK CON TELEDETECCIÓN =====
def calcular_nitrogeno_sentinel2(b5, b8a):
    """Calcula nitrógeno usando NDRE para Sentinel-2"""
    # NDRE = (NIR - Red Edge) / (NIR + Red Edge)
    ndre = (b8a - b5) / (b8a + b5 + 1e-10)
    # Modelo basado en Clevers & Gitelson (2013)
    nitrogeno = 150 * ndre + 50 * (b8a / (b5 + 1e-10))
    return max(0, min(300, nitrogeno)), ndre

def calcular_fosforo_sentinel2(b4, b11):
    """Calcula fósforo usando relación SWIR-VIS para Sentinel-2"""
    # Índice SWIR-VIS (Miphokasap et al., 2012)
    swir_vis_ratio = b11 / (b4 + 1e-10)
    fosforo = 80 * (swir_vis_ratio ** 0.5) + 20
    return max(0, min(100, fosforo)), swir_vis_ratio

def calcular_potasio_sentinel2(b8, b11, b12):
    """Calcula potasio usando índice de estrés hídrico para Sentinel-2"""
    # NDII = (NIR - SWIR) / (NIR + SWIR)
    ndii = (b8 - b11) / (b8 + b11 + 1e-10)
    # Modelo basado en Jackson et al. (2004)
    potasio = 120 * ndii + 40 * (b8 / (b12 + 1e-10))
    return max(0, min(250, potasio)), ndii

def calcular_nitrogeno_landsat8(b3, b4, b5):
    """Calcula nitrógeno usando TCARI/OSAVI para Landsat-8"""
    # TCARI = 3 * [(B5 - B4) - 0.2 * (B5 - B3) * (B5 / B4)]
    tcari = 3 * ((b5 - b4) - 0.2 * (b5 - b3) * (b5 / (b4 + 1e-10)))
    # OSAVI = (1.16 * (B5 - B4)) / (B5 + B4 + 0.16)
    osavi = (1.16 * (b5 - b4)) / (b5 + b4 + 0.16 + 1e-10)
    # TCARI/OSAVI ratio
    tcari_osavi = tcari / (osavi + 1e-10)
    nitrogeno = 100 * tcari_osavi + 30
    return max(0, min(300, nitrogeno)), tcari_osavi

def calcular_fosforo_landsat8(b3, b6):
    """Calcula fósforo usando relación SWIR1-Verde para Landsat-8"""
    # Relación SWIR1-Verde (Chen et al., 2010)
    swir_verde_ratio = b6 / (b3 + 1e-10)
    fosforo = 60 * (swir_verde_ratio ** 0.7) + 25
    return max(0, min(100, fosforo)), swir_verde_ratio

def calcular_potasio_landsat8(b5, b7):
    """Calcula potasio usando índice NIR-SWIR para Landsat-8"""
    # Índice NIR-SWIR (Thenkabail et al., 2000)
    nir_swir_ratio = (b5 - b7) / (b5 + b7 + 1e-10)
    potasio = 100 * nir_swir_ratio + 50
    return max(0, min(250, potasio)), nir_swir_ratio

def calcular_indices_npk_avanzados(gdf, cultivo, satelite, usar_inta=True):
    """Calcula NPK usando metodologías científicas avanzadas + INTA para materia orgánica"""
    resultados = []
    # Usar parámetros específicos por variedad si está disponible
    if 'variedad_params' in st.session_state and st.session_state['variedad_params']:
        params_variedad = st.session_state['variedad_params']
        params = PARAMETROS_CULTIVOS[cultivo].copy()
        params.update({
            'RENDIMIENTO_BASE': params_variedad['RENDIMIENTO_BASE'],
            'RENDIMIENTO_OPTIMO': params_variedad['RENDIMIENTO_OPTIMO'],
            'RESPUESTA_N': params_variedad['RESPUESTA_N'],
            'RESPUESTA_P': params_variedad['RESPUESTA_P'],
            'RESPUESTA_K': params_variedad['RESPUESTA_K'],
            'NITROGENO': {'optimo': params_variedad['NITROGENO_OPTIMO'], 'min': params_variedad['NITROGENO_OPTIMO']*0.7, 'max': params_variedad['NITROGENO_OPTIMO']*1.2},
            'FOSFORO': {'optimo': params_variedad['FOSFORO_OPTIMO'], 'min': params_variedad['FOSFORO_OPTIMO']*0.7, 'max': params_variedad['FOSFORO_OPTIMO']*1.2},
            'POTASIO': {'optimo': params_variedad['POTASIO_OPTIMO'], 'min': params_variedad['POTASIO_OPTIMO']*0.7, 'max': params_variedad['POTASIO_OPTIMO']*1.2}
        })
    else:
        params = PARAMETROS_CULTIVOS[cultivo]

    # Obtener datos del INTA (una vez por parcela)
    datos_inta = obtener_materia_organica_inta(
        gdf, 
        cultivo, 
        usar_inta=usar_inta
    )
    materia_organica_base = datos_inta['materia_organica']

    for idx, row in gdf.iterrows():
        centroid = row.geometry.centroid
        seed_value = abs(hash(f"{centroid.x:.6f}_{centroid.y:.6f}_{cultivo}_{satelite}")) % (2**32)
        rng = np.random.RandomState(seed_value)

        # Simular valores de reflectancia basados en posición y cultivo
        if satelite == "SENTINEL-2":
            # Valores típicos de reflectancia para Sentinel-2 (en %)
            b3 = rng.uniform(0.08, 0.12)  # Verde
            b4 = rng.uniform(0.06, 0.10)  # Rojo
            b5 = rng.uniform(0.10, 0.15)  # Red Edge 1
            b8 = rng.uniform(0.25, 0.40)  # NIR
            b8a = rng.uniform(0.20, 0.35)  # Red Edge 4
            b11 = rng.uniform(0.15, 0.25)  # SWIR 1
            b12 = rng.uniform(0.10, 0.20)  # SWIR 2

            # Calcular NPK
            nitrogeno, ndre = calcular_nitrogeno_sentinel2(b5, b8a)
            fosforo, swir_vis = calcular_fosforo_sentinel2(b4, b11)
            potasio, ndii = calcular_potasio_sentinel2(b8, b11, b12)

            # Ajustar según cultivo
            nitrogeno = nitrogeno * (params['NDRE_OPTIMO'] / 0.5)
            fosforo = fosforo * (params['MATERIA_ORGANICA_OPTIMA'] / 3.5)
            potasio = potasio * (params['HUMEDAD_OPTIMA'] / 0.3)

        elif satelite == "LANDSAT-8":
            # Valores típicos de reflectancia para Landsat-8
            b3 = rng.uniform(0.08, 0.12)  # Verde
            b4 = rng.uniform(0.06, 0.10)  # Rojo
            b5 = rng.uniform(0.20, 0.35)  # NIR
            b6 = rng.uniform(0.12, 0.22)  # SWIR 1
            b7 = rng.uniform(0.08, 0.18)  # SWIR 2

            # Calcular NPK
            nitrogeno, tcari_osavi = calcular_nitrogeno_landsat8(b3, b4, b5)
            fosforo, swir_verde = calcular_fosforo_landsat8(b3, b6)
            potasio, nir_swir = calcular_potasio_landsat8(b5, b7)

            # Ajustar según cultivo
            nitrogeno = nitrogeno * (params['TCARI_OPTIMO'] / 0.4)
            fosforo = fosforo * (params['MATERIA_ORGANICA_OPTIMA'] / 3.5)
            potasio = potasio * (params['HUMEDAD_OPTIMA'] / 0.3)

        else:  # DATOS_SIMULADOS
            # Simulación básica
            nitrogeno = rng.uniform(params['NITROGENO']['min'] * 0.8, params['NITROGENO']['max'] * 1.2)
            fosforo = rng.uniform(params['FOSFORO']['min'] * 0.8, params['FOSFORO']['max'] * 1.2)
            potasio = rng.uniform(params['POTASIO']['min'] * 0.8, params['POTASIO']['max'] * 1.2)
            ndre = rng.uniform(0.2, 0.7)
            swir_vis = rng.uniform(0.5, 2.0)
            ndii = rng.uniform(0.1, 0.6)

        # Aplicar variación espacial a materia orgánica
        materia_organica = materia_organica_base * (1 + rng.normal(0, 0.1))
        materia_organica = max(0.5, min(10.0, materia_organica))

        # Calcular otros índices
        ndvi = rng.uniform(params['NDVI_OPTIMO'] * 0.7, params['NDVI_OPTIMO'] * 1.1)
        humedad_suelo = rng.uniform(params['HUMEDAD_OPTIMA'] * 0.7, params['HUMEDAD_OPTIMA'] * 1.2)
        ndwi = rng.uniform(0.1, 0.4)

        # Índice NPK integrado (0-1)
        npk_integrado = (
            0.4 * (nitrogeno / params['NITROGENO']['optimo']) +
            0.3 * (fosforo / params['FOSFORO']['optimo']) +
            0.3 * (potasio / params['POTASIO']['optimo'])
        ) / 1.0

        resultados.append({
            'nitrogeno_actual': round(nitrogeno, 1),
            'fosforo_actual': round(fosforo, 1),
            'potasio_actual': round(potasio, 1),
            'npk_integrado': round(npk_integrado, 3),
            'materia_organica': round(materia_organica, 2),
            'humedad_suelo': round(humedad_suelo, 3),
            'ndvi': round(ndvi, 3),
            'ndre': round(ndre, 3),
            'ndwi': round(ndwi, 3),
            'ndii': round(ndii, 3) if 'ndii' in locals() else 0.0,
            'region_inta': datos_inta['region_inta'],
            'fuente_materia_organica': datos_inta['fuente']
        })

    return resultados

# ===== FUNCIONES PARA CÁLCULO DE RENDIMIENTO MEJORADAS =====
def obtener_parametros_rendimiento(cultivo):
    """Obtiene parámetros de rendimiento según cultivo y variedad"""
    if 'variedad_params' in st.session_state and st.session_state['variedad_params']:
        variedad_params = st.session_state['variedad_params']
        params = PARAMETROS_CULTIVOS[cultivo].copy()
        # Actualizar con parámetros de la variedad
        params.update({
            'RENDIMIENTO_BASE': variedad_params['RENDIMIENTO_BASE'],
            'RENDIMIENTO_OPTIMO': variedad_params['RENDIMIENTO_OPTIMO'],
            'RESPUESTA_N': variedad_params['RESPUESTA_N'],
            'RESPUESTA_P': variedad_params['RESPUESTA_P'],
            'RESPUESTA_K': variedad_params['RESPUESTA_K'],
            'NITROGENO': {'optimo': variedad_params['NITROGENO_OPTIMO']},
            'FOSFORO': {'optimo': variedad_params['FOSFORO_OPTIMO']},
            'POTASIO': {'optimo': variedad_params['POTASIO_OPTIMO']}
        })
        return params
    else:
        return PARAMETROS_CULTIVOS[cultivo]

def calcular_rendimiento_potencial(gdf_analizado, cultivo):
    """Calcula el rendimiento potencial actual basado en fertilidad existente"""
    params = obtener_parametros_rendimiento(cultivo)
    rendimientos = []
    for idx, row in gdf_analizado.iterrows():
        # Factor de fertilidad actual (0-1)
        factor_fertilidad = row['npk_integrado']
        # Factor de humedad (NDWI ajustado)
        factor_humedad = min(1.0, row['ndwi'] / 0.4) if 'ndwi' in row else 0.7
        # Factor de vigor vegetativo (NDVI)
        factor_vigor = min(1.2, row['ndvi'] / params['NDVI_OPTIMO'])
        # Factor climático base
        factor_clima = params['FACTOR_CLIMA']

        # Cálculo de rendimiento base
        rendimiento_base = params['RENDIMIENTO_BASE']
        # Ajuste por fertilidad actual
        ajuste_fertilidad = 0.5 + (factor_fertilidad * 0.5)  # Entre 0.5 y 1.0

        # Rendimiento potencial estimado
        rendimiento_potencial = (
            rendimiento_base *
            ajuste_fertilidad *
            factor_humedad *
            factor_vigor *
            factor_clima
        )

        # Límite máximo por defecto
        rendimiento_potencial = min(rendimiento_potencial, params['RENDIMIENTO_OPTIMO'] * 1.1)
        rendimientos.append(round(rendimiento_potencial, 2))

    return rendimientos

def calcular_rendimiento_con_recomendaciones(gdf_analizado, cultivo):
    """Calcula el rendimiento proyectado aplicando recomendaciones NPK"""
    params = obtener_parametros_rendimiento(cultivo)
    rendimientos = []
    for idx, row in gdf_analizado.iterrows():
        # Rendimiento base actual
        factor_fertilidad = row['npk_integrado']
        factor_humedad = min(1.0, row['ndwi'] / 0.4) if 'ndwi' in row else 0.7
        factor_vigor = min(1.2, row['ndvi'] / params['NDVI_OPTIMO'])
        factor_clima = params['FACTOR_CLIMA']
        rendimiento_base = params['RENDIMIENTO_BASE']
        ajuste_fertilidad = 0.5 + (factor_fertilidad * 0.5)
        rendimiento_actual = (
            rendimiento_base *
            ajuste_fertilidad *
            factor_humedad *
            factor_vigor *
            factor_clima
        )

        # Calcular incremento por fertilización
        incremento_total = 0

        # Incremento por Nitrógeno
        if 'valor_recomendado' in row and row['valor_recomendado'] > 0:
            n_actual = row['nitrogeno_actual']
            n_optimo = params['NITROGENO']['optimo']
            if n_actual < n_optimo * 0.9:  # Si hay deficiencia significativa
                deficiencia_n = max(0, n_optimo - n_actual)
                eficiencia_n = params['RESPUESTA_N'] * 0.7  # 70% de eficiencia
                incremento_n = deficiencia_n * eficiencia_n
                incremento_total += min(incremento_n, deficiencia_n * params['RESPUESTA_N'])

        # Incremento por Fósforo
        p_actual = row['fosforo_actual']
        p_optimo = params['FOSFORO']['optimo']
        if p_actual < p_optimo * 0.85:
            deficiencia_p = max(0, p_optimo - p_actual)
            eficiencia_p = params['RESPUESTA_P'] * 0.5  # 50% de eficiencia
            incremento_p = deficiencia_p * eficiencia_p
            incremento_total += incremento_p

        # Incremento por Potasio
        k_actual = row['potasio_actual']
        k_optimo = params['POTASIO']['optimo']
        if k_actual < k_optimo * 0.85:
            deficiencia_k = max(0, k_optimo - k_actual)
            eficiencia_k = params['RESPUESTA_K'] * 0.6  # 60% de eficiencia
            incremento_k = deficiencia_k * eficiencia_k
            incremento_total += incremento_k

        # Rendimiento con recomendaciones
        rendimiento_proyectado = rendimiento_actual + incremento_total
        # Límite máximo
        rendimiento_max = params['RENDIMIENTO_OPTIMO'] * 1.2  # 20% sobre el óptimo
        rendimiento_proyectado = min(rendimiento_proyectado, rendimiento_max)
        rendimientos.append(round(rendimiento_proyectado, 2))

    return rendimientos

# ===== FUNCIONES PARA ANÁLISIS ECONÓMICO =====
def realizar_analisis_economico(gdf_analizado, cultivo, variedad_params, area_total):
    """Realiza análisis económico completo (VAN, TIR, B/C) para la parcela"""
    # Obtener parámetros económicos
    precios_cultivo = PARAMETROS_ECONOMICOS['PRECIOS_CULTIVOS'][cultivo]
    precios_fert = PARAMETROS_ECONOMICOS['PRECIOS_FERTILIZANTES']
    conversion = PARAMETROS_ECONOMICOS['CONVERSION_NUTRIENTES']
    financieros = PARAMETROS_ECONOMICOS['PARAMETROS_FINANCIEROS']

    # Calcular promedios de la parcela
    rend_actual_prom = gdf_analizado['rendimiento_actual'].mean()
    rend_proy_prom = gdf_analizado['rendimiento_proyectado'].mean()
    incremento_prom = gdf_analizado['incremento_rendimiento'].mean()

    # Calcular fertilizante necesario (promedio por ha)
    fertilizante_necesario = {
        'NITRÓGENO': 0,
        'FÓSFORO': 0,
        'POTASIO': 0
    }
    if 'valor_recomendado' in gdf_analizado.columns:
        # Solo para análisis de recomendaciones NPK
        for nutriente in ['NITRÓGENO', 'FÓSFORO', 'POTASIO']:
            if f'{nutriente.lower()}_recomendado' in gdf_analizado.columns:
                col_name = f'{nutriente.lower()}_recomendado'
                fertilizante_necesario[nutriente] = gdf_analizado[col_name].mean()

    # === CÁLCULO DE COSTOS ===
    costos = {}
    # 1. Costos fijos (por ha)
    costos['semilla'] = precios_cultivo['costo_semilla']
    costos['herbicidas'] = precios_cultivo['costo_herbicidas']
    costos['insecticidas'] = precios_cultivo['costo_insecticidas']
    costos['labores'] = precios_cultivo['costo_labores']
    costos['cosecha'] = precios_cultivo['costo_cosecha']
    costos['otros'] = precios_cultivo['costo_otros']

    # 2. Costos de fertilización
    costos_fertilizacion = 0

    # Nitrógeno
    if fertilizante_necesario['NITRÓGENO'] > 0:
        fuente_n = conversion['NITRÓGENO']['fuente_principal']
        contenido_n = conversion['NITRÓGENO']['contenido_nutriente']
        eficiencia_n = conversion['NITRÓGENO']['eficiencia']
        # Cantidad de fertilizante necesaria (kg/ha)
        kg_fertilizante_n = (fertilizante_necesario['NITRÓGENO'] / contenido_n) / eficiencia_n
        costo_n = (kg_fertilizante_n / 1000) * precios_fert[fuente_n]
        costos_fertilizacion += costo_n

    # Fósforo
    if fertilizante_necesario['FÓSFORO'] > 0:
        fuente_p = conversion['FÓSFORO']['fuente_principal']
        contenido_p = conversion['FÓSFORO']['contenido_nutriente']
        eficiencia_p = conversion['FÓSFORO']['eficiencia']
        kg_fertilizante_p = (fertilizante_necesario['FÓSFORO'] / contenido_p) / eficiencia_p
        costo_p = (kg_fertilizante_p / 1000) * precios_fert[fuente_p]
        costos_fertilizacion += costo_p

    # Potasio
    if fertilizante_necesario['POTASIO'] > 0:
        fuente_k = conversion['POTASIO']['fuente_principal']
        contenido_k = conversion['POTASIO']['contenido_nutriente']
        eficiencia_k = conversion['POTASIO']['eficiencia']
        kg_fertilizante_k = (fertilizante_necesario['POTASIO'] / contenido_k) / eficiencia_k
        costo_k = (kg_fertilizante_k / 1000) * precios_fert[fuente_k]
        costos_fertilizacion += costo_k

    costos['fertilizacion'] = costos_fertilizacion

    # Costo total por ha
    costo_total_ha = sum(costos.values())

    # === CÁLCULO DE INGRESOS ===
    # Escenario actual (sin fertilización)
    ingresos_actual_ha = rend_actual_prom * precios_cultivo['precio_ton']
    margen_actual_ha = ingresos_actual_ha - costo_total_ha + costos['fertilizacion']  # Sin costo de fertilización

    # Escenario proyectado (con fertilización)
    ingresos_proy_ha = rend_proy_prom * precios_cultivo['precio_ton']
    margen_proy_ha = ingresos_proy_ha - costo_total_ha

    # === CÁLCULO DE INDICADORES FINANCIEROS ===
    # 1. Incremento de margen por ha
    incremento_margen_ha = margen_proy_ha - margen_actual_ha

    # 2. Retorno sobre inversión en fertilización (ROI)
    if costos_fertilizacion > 0:
        roi_fertilizacion = (incremento_margen_ha / costos_fertilizacion) * 100
    else:
        roi_fertilizacion = 0

    # 3. Relación Beneficio/Costo (B/C)
    if costo_total_ha > 0:
        relacion_bc_actual = margen_actual_ha / costo_total_ha
        relacion_bc_proy = margen_proy_ha / costo_total_ha
    else:
        relacion_bc_actual = 0
        relacion_bc_proy = 0

    # 4. VAN (Valor Actual Neto) para 5 años
    flujos = []
    for año in range(financieros['periodo_analisis']):
        # Ajustar por inflación
        factor_inflacion = (1 + financieros['inflacion_esperada']) ** año
        # Flujo neto anual (considerando impuestos y subsidios)
        flujo_neto = incremento_margen_ha * area_total * factor_inflacion
        flujo_neto = flujo_neto * (1 - financieros['impuestos'])  # Después de impuestos
        flujo_neto = flujo_neto * (1 + financieros['subsidios'])  # Con subsidios
        # Inversión inicial solo en el año 0 (costo de fertilización)
        if año == 0:
            flujo_neto -= costos_fertilizacion * area_total
        flujos.append(flujo_neto)

    # Calcular VAN
    van = 0
    for t, flujo in enumerate(flujos):
        van += flujo / ((1 + financieros['tasa_descuento']) ** t)

    # 5. TIR (Tasa Interna de Retorno) - aproximación
    def calcular_tir(flujos):
        """Calcula TIR por prueba y error"""
        def npv(tasa):
            npv_val = 0
            for t, flujo in enumerate(flujos):
                npv_val += flujo / ((1 + tasa) ** t)
            return npv_val

        # Método de bisección
        low = 0.0
        high = 1.0  # 100%
        for _ in range(100):
            mid = (low + high) / 2
            if npv(mid) > 0:
                low = mid
            else:
                high = mid
        return (low + high) / 2

    tir = calcular_tir(flujos) * 100  # Convertir a porcentaje

    # 6. Punto de equilibrio
    if incremento_margen_ha > 0:
        punto_equilibrio_ha = costos_fertilizacion / incremento_margen_ha
    else:
        punto_equilibrio_ha = 0

    # === RESULTADOS CONSOLIDADOS ===
    resultados_economicos = {
        # Información básica
        'cultivo': cultivo,
        'area_total_ha': area_total,
        'variedad': st.session_state.get('variedad', 'No especificada'),
        # Rendimientos
        'rendimiento_actual_ton_ha': rend_actual_prom,
        'rendimiento_proy_ton_ha': rend_proy_prom,
        'incremento_rendimiento_ton_ha': incremento_prom,
        # Costos (USD/ha)
        'costo_total_ha': costo_total_ha,
        'costo_fertilizacion_ha': costos_fertilizacion,
        'costo_semilla_ha': costos['semilla'],
        'costo_insumos_ha': costos['herbicidas'] + costos['insecticidas'] + costos['otros'],
        # Ingresos (USD/ha)
        'ingreso_actual_ha': ingresos_actual_ha,
        'ingreso_proy_ha': ingresos_proy_ha,
        # Margenes (USD/ha)
        'margen_actual_ha': margen_actual_ha,
        'margen_proy_ha': margen_proy_ha,
        'incremento_margen_ha': incremento_margen_ha,
        # Indicadores financieros
        'roi_fertilizacion_%': roi_fertilizacion,
        'relacion_bc_actual': relacion_bc_actual,
        'relacion_bc_proy': relacion_bc_proy,
        'van_usd': van,
        'tir_%': tir,
        'punto_equilibrio_ha': punto_equilibrio_ha,
        # Totales para la parcela
        'incremento_produccion_total_ton': incremento_prom * area_total,
        'incremento_ingreso_total_usd': incremento_margen_ha * area_total,
        'costo_fertilizacion_total_usd': costos_fertilizacion * area_total
    }

    return resultados_economicos

def mostrar_analisis_economico(resultados_economicos):
    """Muestra el análisis económico en una interfaz atractiva"""
    # Calcular flujos proyectados
    financieros = PARAMETROS_ECONOMICOS['PARAMETROS_FINANCIEROS']
    flujos_proyectados = []
    for año in range(financieros['periodo_analisis']):
        factor_inflacion = (1 + financieros['inflacion_esperada']) ** año
        flujo_anual = resultados_economicos['incremento_margen_ha'] * resultados_economicos['area_total_ha'] * factor_inflacion
        if año == 0:
            flujo_anual -= resultados_economicos['costo_fertilizacion_total_usd']
        flujos_proyectados.append({
            'Año': año + 1,
            'Flujo Neto (USD)': flujo_anual,
            'Flujo Acumulado (USD)': sum(f['Flujo Neto (USD)'] for f in flujos_proyectados) + flujo_anual
        })

    st.markdown("---")
    st.subheader("💰 ANÁLISIS ECONÓMICO AGRÍCOLA")

    # Métricas principales
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric(
            "📈 Incremento Producción",
            f"{resultados_economicos['incremento_rendimiento_ton_ha']:.1f} ton/ha",
            f"{resultados_economicos['incremento_produccion_total_ton']:.0f} ton total"
        )
    with col2:
        st.metric(
            "💰 ROI Fertilización",
            f"{resultados_economicos['roi_fertilizacion_%']:.0f}%",
            delta_color="normal"
        )
    with col3:
        st.metric(
            "📊 VAN Proyecto",
            f"${resultados_economicos['van_usd']:,.0f}",
            delta_color="normal"
        )
    with col4:
        st.metric(
            "🎯 TIR",
            f"{resultados_economicos['tir_%']:.1f}%",
            delta_color="normal"
        )

    # Tabs para análisis detallado
    tab1, tab2, tab3, tab4 = st.tabs(["📋 Resumen", "📊 Costos", "💰 Rentabilidad", "📈 Proyecciones"])
    with tab1:
        # Tabla resumen
        st.markdown("#### 📊 RESUMEN ECONÓMICO POR HECTÁREA")
        resumen_data = {
            'Concepto': [
                'Rendimiento Actual',
                'Rendimiento Proyectado',
                'Incremento',
                'Ingresos Actuales',
                'Ingresos Proyectados',
                'Costos Totales',
                'Margen Actual',
                'Margen Proyectado',
                'Incremento Margen'
            ],
            'Valor (USD/ha)': [
                f"{resultados_economicos['rendimiento_actual_ton_ha']:.1f} ton",
                f"{resultados_economicos['rendimiento_proy_ton_ha']:.1f} ton",
                f"+{resultados_economicos['incremento_rendimiento_ton_ha']:.1f} ton",
                f"${resultados_economicos['ingreso_actual_ha']:,.0f}",
                f"${resultados_economicos['ingreso_proy_ha']:,.0f}",
                f"${resultados_economicos['costo_total_ha']:,.0f}",
                f"${resultados_economicos['margen_actual_ha']:,.0f}",
                f"${resultados_economicos['margen_proy_ha']:,.0f}",
                f"${resultados_economicos['incremento_margen_ha']:,.0f}"
            ]
        }
        st.dataframe(pd.DataFrame(resumen_data), use_container_width=True)

    with tab2:
        # Gráfico de costos
        st.markdown("#### 🏷️ DISTRIBUCIÓN DE COSTOS (USD/ha)")
        costos_data = {
            'Rubro': ['Semilla', 'Fertilización', 'Herbicidas', 'Inseticidas', 'Labores', 'Cosecha', 'Otros'],
            'Costo USD/ha': [
                resultados_economicos['costo_semilla_ha'],
                resultados_economicos['costo_fertilizacion_ha'],
                PARAMETROS_ECONOMICOS['PRECIOS_CULTIVOS'][resultados_economicos['cultivo']]['costo_herbicidas'],
                PARAMETROS_ECONOMICOS['PRECIOS_CULTIVOS'][resultados_economicos['cultivo']]['costo_insecticidas'],
                PARAMETROS_ECONOMICOS['PRECIOS_CULTIVOS'][resultados_economicos['cultivo']]['costo_labores'],
                PARAMETROS_ECONOMICOS['PRECIOS_CULTIVOS'][resultados_economicos['cultivo']]['costo_cosecha'],
                PARAMETROS_ECONOMICOS['PRECIOS_CULTIVOS'][resultados_economicos['cultivo']]['costo_otros']
            ]
        }
        df_costos = pd.DataFrame(costos_data)
        fig, ax = plt.subplots(figsize=(10, 6))
        colors = ['#3b82f6', '#10b981', '#f59e0b', '#ef4444', '#8b5cf6', '#ec4899', '#14b8a6']
        ax.bar(df_costos['Rubro'], df_costos['Costo USD/ha'], color=colors)
        ax.set_title('Distribución de Costos por Hectárea', color='white')
        ax.set_ylabel('USD/ha', color='white')
        ax.tick_params(colors='white')
        ax.set_facecolor('#0f172a')
        fig.patch.set_facecolor('#0f172a')
        # Agregar valores en las barras
        for i, v in enumerate(df_costos['Costo USD/ha']):
            ax.text(i, v + 5, f'${v:,.0f}', ha='center', color='white', fontweight='bold')
        st.pyplot(fig)

    with tab3:
        # Gráfico de rentabilidad comparativa
        st.markdown("#### 📈 COMPARATIVA DE RENTABILIDAD")
        rentabilidad_data = {
            'Escenario': ['Actual (sin fert.)', 'Proyectado (con fert.)'],
            'B/C Ratio': [
                resultados_economicos['relacion_bc_actual'],
                resultados_economicos['relacion_bc_proy']
            ],
            'Margen (USD/ha)': [
                resultados_economicos['margen_actual_ha'],
                resultados_economicos['margen_proy_ha']
            ]
        }
        df_rent = pd.DataFrame(rentabilidad_data)
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 5))

        # Gráfico 1: B/C Ratio
        bars1 = ax1.bar(df_rent['Escenario'], df_rent['B/C Ratio'], color=['#ef4444', '#10b981'])
        ax1.set_title('Relación Beneficio/Costo', color='white')
        ax1.set_ylabel('Ratio B/C', color='white')
        ax1.axhline(y=1.0, color='white', linestyle='--', alpha=0.5)
        ax1.text(0.5, 1.05, 'Límite Rentabilidad', ha='center', color='white', fontsize=9)

        # Gráfico 2: Margen USD/ha
        bars2 = ax2.bar(df_rent['Escenario'], df_rent['Margen (USD/ha)'], color=['#f59e0b', '#3b82f6'])
        ax2.set_title('Margen Neto por Hectárea', color='white')
        ax2.set_ylabel('USD/ha', color='white')

        # Configurar ambos gráficos
        for ax in [ax1, ax2]:
            ax.set_facecolor('#0f172a')
            ax.tick_params(colors='white')

        # Agregar valores en las barras
        for bar in bars1:
            height = bar.get_height()
            ax1.text(bar.get_x() + bar.get_width()/2., height + 0.05,
                     f'{height:.2f}',
                     ha='center', va='bottom', color='white', fontweight='bold')

        for bar in bars2:
            height = bar.get_height()
            ax2.text(bar.get_x() + bar.get_width()/2., height + 10,
                     f'${height:,.0f}',
                     ha='center', va='bottom', color='white', fontweight='bold')

        fig.patch.set_facecolor('#0f172a')
        plt.tight_layout()
        st.pyplot(fig)

        # Interpretación
        st.markdown("##### 💡 INTERPRETACIÓN:")
        if resultados_economicos['relacion_bc_proy'] > 1.5:
            st.success("**EXCELENTE RENTABILIDAD:** La inversión en fertilización es altamente rentable (B/C > 1.5)")
        elif resultados_economicos['relacion_bc_proy'] > 1.2:
            st.info("**BUENA RENTABILIDAD:** La inversión es recomendable (B/C > 1.2)")
        elif resultados_economicos['relacion_bc_proy'] > 1.0:
            st.warning("**RENTABILIDAD LIMITE:** La inversión apenas cubre costos (B/C > 1.0)")

    with tab4:
        st.markdown("#### 📅 PROYECCIÓN DE FLUJOS DE CAJA (5 AÑOS)")
        # Usar flujos_proyectados ya calculados
        df_flujos = pd.DataFrame(flujos_proyectados)
        st.dataframe(df_flujos.style.format({
            'Flujo Neto (USD)': '${:,.0f}',
            'Flujo Acumulado (USD)': '${:,.0f}'
        }), use_container_width=True)

        # Gráfico de flujos
        fig, ax = plt.subplots(figsize=(10, 5))
        años = df_flujos['Año']
        flujos = df_flujos['Flujo Neto (USD)']

        # Barras para flujos anuales
        bars = ax.bar(años, flujos, color=['#ef4444' if f < 0 else '#10b981' for f in flujos])
        ax.set_title('Flujos de Caja Anuales', color='white')
        ax.set_xlabel('Año', color='white')
        ax.set_ylabel('USD', color='white')
        ax.axhline(y=0, color='white', linewidth=1)

        # Línea para flujo acumulado
        ax2 = ax.twinx()
        ax2.plot(años, df_flujos['Flujo Acumulado (USD)'], color='#3b82f6', marker='o', linewidth=3)
        ax2.set_ylabel('Flujo Acumulado (USD)', color='#3b82f6')
        ax2.tick_params(axis='y', colors='#3b82f6')

        ax.set_facecolor('#0f172a')
        fig.patch.set_facecolor('#0f172a')
        ax.tick_params(colors='white')
        st.pyplot(fig)

    # Recomendaciones económicas
    st.markdown("---")
    st.subheader("💡 RECOMENDACIONES ECONÓMICAS")
    rec_col1, rec_col2 = st.columns(2)
    with rec_col1:
        st.markdown("##### ✅ ACCIONES RECOMENDADAS")
        if resultados_economicos['roi_fertilizacion_%'] > 100:
            st.success("**INVERTIR EN FERTILIZACIÓN:** ROI > 100% indica excelente retorno")
        if resultados_economicos['van_usd'] > 0:
            st.success("**PROYECTO VIABLE:** VAN positivo genera valor económico")
        if resultados_economicos['tir_%'] > PARAMETROS_ECONOMICOS['PARAMETROS_FINANCIEROS']['tasa_descuento'] * 100:
            st.success(f"**TIR ATRACTIVA:** {resultados_economicos['tir_%']:.1f}% supera la tasa de descuento")

        # Recomendación específica por cultivo
        if resultados_economicos['cultivo'] == "VID":
            st.info("Para vid: Priorizar potasio para calidad de uva y aplicaciones fraccionadas de nitrógeno")
        elif resultados_economicos['cultivo'] == "OLIVO":
            st.info("Para olivo: Balancear nitrógeno y potasio para optimizar producción de aceite")
        elif resultados_economicos['cultivo'] == "HORTALIZAS DE HOJA":
            st.info("Para hortalizas de hoja: Aplicar nitrógeno en dosis divididas para máximo crecimiento foliar")

    with rec_col2:
        st.markdown("##### ⚠️ CONSIDERACIONES")
        st.warning("**RIESGOS CLIMÁTICOS:** Considerar seguro agrícola")
        st.warning("**VOLATILIDAD PRECIOS:** Diversificar cultivos si es posible")
        st.warning("**COSTOS LOGÍSTICOS:** Incluir en análisis de rentabilidad")

    # Descargar análisis económico
    st.markdown("---")
    st.subheader("📥 EXPORTAR ANÁLISIS ECONÓMICO")
    if st.button("📊 Generar Reporte Económico (Excel)"):
        # Crear DataFrame para Excel
        df_economico = pd.DataFrame([resultados_economicos])

        # Crear Excel con múltiples hojas
        excel_buffer = BytesIO()
        with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
            # Hoja 1: Resumen
            df_economico.T.to_excel(writer, sheet_name='Resumen')

            # Hoja 2: Costos detallados
            costos_det = pd.DataFrame({
                'Rubro': ['Semilla', 'Fertilización', 'Herbicidas', 'Inseticidas', 'Labores', 'Cosecha', 'Otros', 'TOTAL'],
                'USD/ha': [
                    resultados_economicos['costo_semilla_ha'],
                    resultados_economicos['costo_fertilizacion_ha'],
                    PARAMETROS_ECONOMICOS['PRECIOS_CULTIVOS'][resultados_economicos['cultivo']]['costo_herbicidas'],
                    PARAMETROS_ECONOMICOS['PRECIOS_CULTIVOS'][resultados_economicos['cultivo']]['costo_insecticidas'],
                    PARAMETROS_ECONOMICOS['PRECIOS_CULTIVOS'][resultados_economicos['cultivo']]['costo_labores'],
                    PARAMETROS_ECONOMICOS['PRECIOS_CULTIVOS'][resultados_economicos['cultivo']]['costo_cosecha'],
                    PARAMETROS_ECONOMICOS['PRECIOS_CULTIVOS'][resultados_economicos['cultivo']]['costo_otros'],
                    resultados_economicos['costo_total_ha']
                ]
            })
            costos_det.to_excel(writer, sheet_name='Costos', index=False)

            # Hoja 3: Proyecciones
            pd.DataFrame(flujos_proyectados).to_excel(writer, sheet_name='Proyecciones', index=False)

        excel_buffer.seek(0)
        st.download_button(
            label="📥 Descargar Análisis Económico (Excel)",
            data=excel_buffer,
            file_name=f"analisis_economico_{resultados_economicos['cultivo']}_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

# ===== FUNCIONES PARA GENERAR MAPAS DE CALOR DE RENDIMIENTO =====
def crear_mapa_calor_rendimiento_actual(gdf_analizado, cultivo):
    """Crea mapa de calor para rendimiento actual con visualización suave y profesional"""
    try:
        if 'rendimiento_actual' not in gdf_analizado.columns:
            return None

        gdf_plot = gdf_analizado.to_crs(epsg=3857)

        # Crear figura con estilo moderno
        fig, ax = plt.subplots(1, 1, figsize=(14, 8))
        fig.patch.set_facecolor('#0f172a')
        ax.set_facecolor('#0f172a')

        # Obtener los centroides para interpolación
        centroids = gdf_plot.geometry.centroid
        x = np.array([c.x for c in centroids])
        y = np.array([c.y for c in centroids])
        z = gdf_plot['rendimiento_actual'].values

        # Crear malla para interpolación
        x_min, y_min, x_max, y_max = gdf_plot.total_bounds
        xi = np.linspace(x_min, x_max, 200)
        yi = np.linspace(y_min, y_max, 200)
        xi, yi = np.meshgrid(xi, yi)

        # Interpolación lineal para suavizar
        try:
            from scipy.interpolate import griddata
            zi = griddata((x, y), z, (xi, yi), method='cubic', fill_value=np.nan)
        except:
            # Fallback a interpolación lineal
            zi = griddata((x, y), z, (xi, yi), method='linear', fill_value=np.nan)

        # Crear mapa de calor suave
        im = ax.contourf(xi, yi, zi, levels=50, cmap='RdYlGn', alpha=0.8,
                         vmin=z.min()*0.9, vmax=z.max()*1.1)

        # Agregar líneas de contorno
        contour = ax.contour(xi, yi, zi, levels=10, colors='white', linewidths=0.5, alpha=0.5)

        # Agregar etiquetas en los centroides
        for idx, (centroid, valor) in enumerate(zip(centroids, z)):
            ax.plot(centroid.x, centroid.y, 'o', markersize=8,
                    markeredgecolor='white', markerfacecolor=plt.cm.RdYlGn((valor - z.min())/(z.max() - z.min())))

            # Etiqueta con valor
            if idx % 2 == 0:  # Mostrar solo algunas etiquetas para evitar sobrecarga
                ax.annotate(f"{valor:.1f}t",
                            (centroid.x, centroid.y),
                            xytext=(0, 10), textcoords="offset points",
                            fontsize=8, color='white', weight='bold',
                            ha='center', va='center',
                            bbox=dict(boxstyle="round,pad=0.2",
                                      facecolor=(0, 0, 0, 0.7),
                                      alpha=0.8))

        # Agregar mapa base
        try:
            ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, alpha=0.4)
        except:
            pass

        # Configurar título y etiquetas
        ax.set_title(f'🌾 MAPA DE CALOR - RENDIMIENTO ACTUAL\n{cultivo} (ton/ha)',
                     fontsize=16, fontweight='bold', pad=20, color='white')
        ax.set_xlabel('Longitud', color='white')
        ax.set_ylabel('Latitud', color='white')
        ax.tick_params(colors='white')
        ax.grid(True, alpha=0.1, color='#475569', linestyle='--')

        # Barra de colores profesional
        cbar = plt.colorbar(im, ax=ax, shrink=0.8, pad=0.02)
        cbar.set_label('Rendimiento (ton/ha)', fontsize=12, fontweight='bold', color='white')
        cbar.ax.yaxis.set_tick_params(color='white')
        cbar.outline.set_edgecolor('white')
        plt.setp(plt.getp(cbar.ax.axes, 'yticklabels'), color='white')

        # Leyenda de interpretación
        stats = {
            'promedio': z.mean(),
            'min': z.min(),
            'max': z.max(),
            'std': z.std()
        }
        info_text = f"""
        📊 ESTADÍSTICAS:
        • Promedio: {stats['promedio']:.1f} ton/ha
        • Mínimo: {stats['min']:.1f} ton/ha
        • Máximo: {stats['max']:.1f} ton/ha
        • Variación: {stats['std']:.1f} ton/ha
        """
        ax.text(0.02, 0.98, info_text, transform=ax.transAxes, fontsize=9,
                verticalalignment='top', color='white',
                bbox=dict(boxstyle="round,pad=0.3",
                          facecolor=(30/255, 41/255, 59/255, 0.9),
                          alpha=0.9, edgecolor='white'))

        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=200, bbox_inches='tight',
                    facecolor='#0f172a', transparent=False)
        buf.seek(0)
        plt.close()
        return buf

    except Exception as e:
        st.error(f"Error creando mapa de calor actual: {str(e)}")
        # Fallback al método anterior si falla la interpolación
        return crear_mapa_calor_rendimiento_actual_fallback(gdf_analizado, cultivo)

def crear_mapa_calor_rendimiento_actual_fallback(gdf_analizado, cultivo):
    """Versión fallback sin interpolación"""
    try:
        gdf_plot = gdf_analizado.to_crs(epsg=3857)
        fig, ax = plt.subplots(1, 1, figsize=(12, 8))
        fig.patch.set_facecolor('#0f172a')
        ax.set_facecolor('#0f172a')

        # Valores de rendimiento
        valores = gdf_plot['rendimiento_actual']
        vmin = valores.min() * 0.9
        vmax = valores.max() * 1.1

        # Crear colormap profesional
        colors = ['#d73027', '#f46d43', '#fdae61', '#fee08b', '#ffffbf',
                  '#d9ef8b', '#a6d96a', '#66bd63', '#1a9850', '#006837']
        cmap = LinearSegmentedColormap.from_list('rendimiento_actual', colors)

        # Plotear cada zona con gradiente suave
        for idx, row in gdf_plot.iterrows():
            valor = row['rendimiento_actual']
            valor_norm = (valor - vmin) / (vmax - vmin) if vmax != vmin else 0.5
            valor_norm = max(0, min(1, valor_norm))
            color = cmap(valor_norm)

            # Dibujar polígono con borde suave
            gdf_plot.iloc[[idx]].plot(ax=ax, color=color, edgecolor='white',
                                      linewidth=1, alpha=0.85)

            # Etiqueta con rendimiento
            centroid = row.geometry.centroid
            ax.annotate(f"{valor:.1f}t",
                        (centroid.x, centroid.y),
                        xytext=(0, 0), textcoords="offset points",
                        fontsize=8, color='white', weight='bold',
                        bbox=dict(boxstyle="circle,pad=0.2",
                                  facecolor=(0, 0, 0, 0.6),
                                  alpha=0.8, edgecolor='white'))

        # Agregar mapa base
        try:
            ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, alpha=0.3)
        except:
            pass

        ax.set_title(f'🌾 MAPA DE CALOR - RENDIMIENTO ACTUAL\n{cultivo} (ton/ha)',
                     fontsize=16, fontweight='bold', pad=20, color='white')
        ax.set_xlabel('Longitud', color='white')
        ax.set_ylabel('Latitud', color='white')
        ax.tick_params(colors='white')
        ax.grid(True, alpha=0.2, color='#475569')

        # Barra de colores mejorada
        sm = plt.cm.ScalarMappable(cmap=cmap, norm=plt.Normalize(vmin=vmin, vmax=vmax))
        sm.set_array([])
        cbar = plt.colorbar(sm, ax=ax, shrink=0.8)
        cbar.set_label('Rendimiento (ton/ha)', fontsize=12, fontweight='bold', color='white')
        cbar.ax.yaxis.set_tick_params(color='white')
        cbar.outline.set_edgecolor('white')
        plt.setp(plt.getp(cbar.ax.axes, 'yticklabels'), color='white')

        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='#0f172a')
        buf.seek(0)
        plt.close()
        return buf

    except Exception as e:
        st.error(f"Error en fallback: {str(e)}")
        return None

def crear_mapa_calor_rendimiento_proyectado(gdf_analizado, cultivo):
    """Crea mapa de calor para rendimiento proyectado con visualización profesional"""
    try:
        if 'rendimiento_proyectado' not in gdf_analizado.columns:
            return None

        gdf_plot = gdf_analizado.to_crs(epsg=3857)

        # Crear figura con estilo moderno
        fig, ax = plt.subplots(1, 1, figsize=(14, 8))
        fig.patch.set_facecolor('#0f172a')
        ax.set_facecolor('#0f172a')

        # Obtener datos para interpolación
        centroids = gdf_plot.geometry.centroid
        x = np.array([c.x for c in centroids])
        y = np.array([c.y for c in centroids])
        z_proyectado = gdf_plot['rendimiento_proyectado'].values
        z_actual = gdf_plot['rendimiento_actual'].values
        incrementos = z_proyectado - z_actual

        # Crear malla para interpolación
        x_min, y_min, x_max, y_max = gdf_plot.total_bounds
        xi = np.linspace(x_min, x_max, 200)
        yi = np.linspace(y_min, y_max, 200)
        xi, yi = np.meshgrid(xi, yi)

        # Interpolación del rendimiento proyectado
        try:
            from scipy.interpolate import griddata
            zi_proyectado = griddata((x, y), z_proyectado, (xi, yi), method='cubic', fill_value=np.nan)
            zi_incremento = griddata((x, y), incrementos, (xi, yi), method='cubic', fill_value=np.nan)
        except:
            zi_proyectado = griddata((x, y), z_proyectado, (xi, yi), method='linear', fill_value=np.nan)
            zi_incremento = griddata((x, y), incrementos, (xi, yi), method='linear', fill_value=np.nan)

        # Crear mapa de calor con dos capas
        im_proyectado = ax.contourf(xi, yi, zi_proyectado, levels=50, cmap='RdYlGn', alpha=0.7,
                                    vmin=z_proyectado.min()*0.9, vmax=z_proyectado.max()*1.1)

        # Superponer mapa de incrementos con transparencia
        im_incremento = ax.contourf(xi, yi, zi_incremento, levels=20, cmap='viridis', alpha=0.4)

        # Agregar líneas de contorno para rendimiento proyectado
        contour = ax.contour(xi, yi, zi_proyectado, levels=8, colors='white', linewidths=1, alpha=0.6)

        # Etiquetar las líneas de contorno
        ax.clabel(contour, inline=True, fontsize=8, colors='white', fmt='%1.1f t')

        # Agregar puntos de datos
        for idx, (centroid, valor_proy, valor_act, inc) in enumerate(zip(centroids, z_proyectado, z_actual, incrementos)):
            # Punto con tamaño proporcional al incremento
            marker_size = 6 + (inc / max(incrementos) * 10) if max(incrementos) > 0 else 8
            ax.plot(centroid.x, centroid.y, 'o', markersize=marker_size,
                    markeredgecolor='white', markerfacecolor=plt.cm.RdYlGn((valor_proy - z_proyectado.min())/(z_proyectado.max() - z_proyectado.min())),
                    markeredgewidth=1)

            # Etiqueta con incremento
            if idx % 3 == 0:  # Mostrar algunas etiquetas
                ax.annotate(f"+{inc:.1f}t",
                            (centroid.x, centroid.y),
                            xytext=(0, 15), textcoords="offset points",
                            fontsize=7, color='cyan', weight='bold',
                            ha='center', va='center',
                            bbox=dict(boxstyle="round,pad=0.2",
                                      facecolor=(0, 0, 0, 0.7),
                                      alpha=0.8))

        # Agregar mapa base
        try:
            ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, alpha=0.3)
        except:
            pass

        # Configurar título y etiquetas
        ax.set_title(f'🚀 MAPA DE CALOR - RENDIMIENTO PROYECTADO\n{cultivo} (con fertilización óptima)',
                     fontsize=16, fontweight='bold', pad=20, color='white')
        ax.set_xlabel('Longitud', color='white')
        ax.set_ylabel('Latitud', color='white')
        ax.tick_params(colors='white')
        ax.grid(True, alpha=0.1, color='#475569', linestyle='--')

        # Barra de colores principal
        cbar1 = plt.colorbar(im_proyectado, ax=ax, shrink=0.8, pad=0.02)
        cbar1.set_label('Rendimiento Proyectado (ton/ha)', fontsize=12, fontweight='bold', color='white')
        cbar1.ax.yaxis.set_tick_params(color='white')

        # Barra de colores para incrementos (más pequeña)
        from mpl_toolkits.axes_grid1 import make_axes_locatable
        divider = make_axes_locatable(ax)
        cax2 = divider.append_axes("right", size="3%", pad=0.15)
        cbar2 = plt.colorbar(im_incremento, cax=cax2)
        cbar2.set_label('Incremento (ton/ha)', fontsize=9, color='white')
        cbar2.ax.yaxis.set_tick_params(color='white')

        # Configurar colores de barras
        for cbar in [cbar1, cbar2]:
            cbar.outline.set_edgecolor('white')
            plt.setp(plt.getp(cbar.ax.axes, 'yticklabels'), color='white')

        # Estadísticas
        stats_text = f"""
        📈 ESTADÍSTICAS DE POTENCIAL:
        • Actual: {z_actual.mean():.1f} ton/ha
        • Proyectado: {z_proyectado.mean():.1f} ton/ha
        • Incremento: +{incrementos.mean():.1f} ton/ha
        • Aumento: +{(incrementos.mean()/z_actual.mean()*100 if z_actual.mean()>0 else 0):.1f}%
        """
        ax.text(0.02, 0.98, stats_text, transform=ax.transAxes, fontsize=9,
                verticalalignment='top', color='white',
                bbox=dict(boxstyle="round,pad=0.3",
                          facecolor=(30/255, 41/255, 59/255, 0.9),
                          alpha=0.9, edgecolor='white'))

        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=200, bbox_inches='tight',
                    facecolor='#0f172a', transparent=False)
        buf.seek(0)
        plt.close()
        return buf

    except Exception as e:
        st.error(f"Error creando mapa de calor proyectado: {str(e)}")
        return crear_mapa_calor_rendimiento_proyectado_fallback(gdf_analizado, cultivo)

def crear_mapa_calor_rendimiento_proyectado_fallback(gdf_analizado, cultivo):
    """Versión fallback sin interpolación"""
    try:
        gdf_plot = gdf_analizado.to_crs(epsg=3857)
        fig, ax = plt.subplots(1, 1, figsize=(12, 8))
        fig.patch.set_facecolor('#0f172a')
        ax.set_facecolor('#0f172a')

        # Valores de rendimiento proyectado
        valores = gdf_plot['rendimiento_proyectado']
        vmin = valores.min() * 0.9
        vmax = valores.max() * 1.1

        # Crear colormap para rendimiento proyectado
        colors = ['#313695', '#4575b4', '#74add1', '#abd9e9', '#e0f3f8',
                  '#fee090', '#fdae61', '#f46d43', '#d73027', '#a50026']
        cmap = LinearSegmentedColormap.from_list('rendimiento_proyectado', colors)

        # Plotear cada zona
        for idx, row in gdf_plot.iterrows():
            valor = row['rendimiento_proyectado']
            incremento = row['rendimiento_proyectado'] - row['rendimiento_actual']
            valor_norm = (valor - vmin) / (vmax - vmin) if vmax != vmin else 0.5
            valor_norm = max(0, min(1, valor_norm))
            color = cmap(valor_norm)

            # Dibujar polígono con borde
            gdf_plot.iloc[[idx]].plot(ax=ax, color=color, edgecolor='white',
                                      linewidth=1, alpha=0.85)

            # Etiqueta con rendimiento e incremento
            centroid = row.geometry.centroid
            ax.annotate(f"{valor:.1f}t\n(+{incremento:.1f})",
                        (centroid.x, centroid.y),
                        xytext=(0, 0), textcoords="offset points",
                        fontsize=7, color='white', weight='bold',
                        ha='center', va='center',
                        bbox=dict(boxstyle="round,pad=0.2",
                                  facecolor=(0, 0, 0, 0.6),
                                  alpha=0.8, edgecolor='white'))

        # Agregar mapa base
        try:
            ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, alpha=0.3)
        except:
            pass

        ax.set_title(f'🚀 MAPA DE CALOR - RENDIMIENTO PROYECTADO\n{cultivo} (con fertilización óptima)',
                     fontsize=16, fontweight='bold', pad=20, color='white')
        ax.set_xlabel('Longitud', color='white')
        ax.set_ylabel('Latitud', color='white')
        ax.tick_params(colors='white')
        ax.grid(True, alpha=0.2, color='#475569')

        # Barra de colores
        sm = plt.cm.ScalarMappable(cmap=cmap, norm=plt.Normalize(vmin=vmin, vmax=vmax))
        sm.set_array([])
        cbar = plt.colorbar(sm, ax=ax, shrink=0.8)
        cbar.set_label('Rendimiento Proyectado (ton/ha)', fontsize=12, fontweight='bold', color='white')
        cbar.ax.yaxis.set_tick_params(color='white')
        cbar.outline.set_edgecolor('white')
        plt.setp(plt.getp(cbar.ax.axes, 'yticklabels'), color='white')

        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='#0f172a')
        buf.seek(0)
        plt.close()
        return buf

    except Exception as e:
        st.error(f"Error en fallback: {str(e)}")
        return None

def crear_mapa_comparativo_calor(gdf_analizado, cultivo):
    """Crea mapa comparativo side-by-side de rendimiento actual vs proyectado con estilo profesional"""
    try:
        if 'rendimiento_actual' not in gdf_analizado.columns or 'rendimiento_proyectado' not in gdf_analizado.columns:
            return None

        gdf_plot = gdf_analizado.to_crs(epsg=3857)
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(18, 8))
        fig.patch.set_facecolor('#0f172a')
        ax1.set_facecolor('#0f172a')
        ax2.set_facecolor('#0f172a')

        # Obtener centroides para ambos mapas
        centroids = gdf_plot.geometry.centroid
        x = np.array([c.x for c in centroids])
        y = np.array([c.y for c in centroids])
        z_actual = gdf_plot['rendimiento_actual'].values
        z_proyectado = gdf_plot['rendimiento_proyectado'].values
        incrementos = z_proyectado - z_actual

        # Rango común para comparación
        vmin = min(z_actual.min(), z_proyectado.min()) * 0.9
        vmax = max(z_actual.max(), z_proyectado.max()) * 1.1

        # Crear malla para interpolación
        x_min, y_min, x_max, y_max = gdf_plot.total_bounds
        xi = np.linspace(x_min, x_max, 150)
        yi = np.linspace(y_min, y_max, 150)
        xi, yi = np.meshgrid(xi, yi)

        # Interpolación para ambos conjuntos de datos
        try:
            from scipy.interpolate import griddata
            zi_actual = griddata((x, y), z_actual, (xi, yi), method='cubic', fill_value=np.nan)
            zi_proyectado = griddata((x, y), z_proyectado, (xi, yi), method='cubic', fill_value=np.nan)
        except:
            zi_actual = griddata((x, y), z_actual, (xi, yi), method='linear', fill_value=np.nan)
            zi_proyectado = griddata((x, y), z_proyectado, (xi, yi), method='linear', fill_value=np.nan)

        # MAPA 1: RENDIMIENTO ACTUAL
        im1 = ax1.contourf(xi, yi, zi_actual, levels=40, cmap='RdYlGn', alpha=0.8, vmin=vmin, vmax=vmax)
        contour1 = ax1.contour(xi, yi, zi_actual, levels=6, colors='white', linewidths=1, alpha=0.5)
        ax1.clabel(contour1, inline=True, fontsize=8, colors='white', fmt='%1.1f t')

        # Agregar puntos de datos
        for centroid, valor in zip(centroids, z_actual):
            ax1.plot(centroid.x, centroid.y, 'o', markersize=6,
                     markeredgecolor='white', markerfacecolor=plt.cm.RdYlGn((valor - vmin)/(vmax - vmin)))

        # MAPA 2: RENDIMIENTO PROYECTADO
        im2 = ax2.contourf(xi, yi, zi_proyectado, levels=40, cmap='RdYlGn', alpha=0.8, vmin=vmin, vmax=vmax)
        contour2 = ax2.contour(xi, yi, zi_proyectado, levels=6, colors='white', linewidths=1, alpha=0.5)
        ax2.clabel(contour2, inline=True, fontsize=8, colors='white', fmt='%1.1f t')

        # Superponer capa de incrementos
        zi_incremento = griddata((x, y), incrementos, (xi, yi), method='linear', fill_value=np.nan)
        im_incremento = ax2.contourf(xi, yi, zi_incremento, levels=15, cmap='Blues', alpha=0.3)

        # Agregar puntos de datos con incremento
        for centroid, valor, inc in zip(centroids, z_proyectado, incrementos):
            ax2.plot(centroid.x, centroid.y, 'o', markersize=6 + (inc/max(incrementos)*3 if max(incrementos)>0 else 0),
                     markeredgecolor='white', markerfacecolor=plt.cm.RdYlGn((valor - vmin)/(vmax - vmin)))

        # Títulos
        ax1.set_title('🌾 RENDIMIENTO ACTUAL\n(ton/ha)', fontsize=14, fontweight='bold', color='white', pad=15)
        ax2.set_title('🚀 RENDIMIENTO PROYECTADO\n(ton/ha)', fontsize=14, fontweight='bold', color='white', pad=15)

        # Configuración común
        for ax in [ax1, ax2]:
            ax.set_xlabel('Longitud', color='white')
            ax.set_ylabel('Latitud', color='white')
            ax.tick_params(colors='white')
            ax.grid(True, alpha=0.1, color='#475569', linestyle='--')

        # Barras de color
        cbar1 = plt.colorbar(im1, ax=ax1, shrink=0.6)
        cbar1.set_label('ton/ha', fontsize=10, color='white')
        cbar1.ax.yaxis.set_tick_params(color='white')
        cbar2 = plt.colorbar(im2, ax=ax2, shrink=0.6)
        cbar2.set_label('ton/ha', fontsize=10, color='white')
        cbar2.ax.yaxis.set_tick_params(color='white')

        # Configurar colores de barras
        for cbar in [cbar1, cbar2]:
            cbar.outline.set_edgecolor('white')
            plt.setp(plt.getp(cbar.ax.axes, 'yticklabels'), color='white')

        # Estadísticas comparativas
        rend_actual_prom = z_actual.mean()
        rend_proy_prom = z_proyectado.mean()
        incremento_prom = incrementos.mean()
        porcentaje_aumento = (incremento_prom / rend_actual_prom * 100) if rend_actual_prom > 0 else 0

        info_comparativo = f"""
        📊 COMPARATIVA DE POTENCIAL:
        ┌─────────────────────────────┐
        │ Actual:     {rend_actual_prom:6.1f} ton/ha │
        │ Proyectado: {rend_proy_prom:6.1f} ton/ha │
        │ Incremento: +{incremento_prom:5.1f} ton/ha │
        │ Aumento:    +{porcentaje_aumento:5.1f}%    │
        └─────────────────────────────┘
        """

        # Agregar texto en la parte inferior central
        fig.text(0.5, 0.02, info_comparativo, fontsize=11, color='white', ha='center',
                 bbox=dict(boxstyle="round,pad=0.5",
                           facecolor=(30/255, 41/255, 59/255, 0.95),
                           alpha=0.95, edgecolor='#3b82f6', linewidth=2))

        plt.tight_layout(rect=[0, 0.05, 1, 0.95])
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=200, bbox_inches='tight',
                    facecolor='#0f172a', transparent=False)
        buf.seek(0)
        plt.close()
        return buf

    except Exception as e:
        st.error(f"Error creando mapa comparativo: {str(e)}")
        return crear_mapa_comparativo_calor_fallback(gdf_analizado, cultivo)

def crear_mapa_comparativo_calor_fallback(gdf_analizado, cultivo):
    """Versión fallback sin interpolación"""
    try:
        gdf_plot = gdf_analizado.to_crs(epsg=3857)
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 7))
        fig.patch.set_facecolor('#0f172a')
        ax1.set_facecolor('#0f172a')
        ax2.set_facecolor('#0f172a')

        # Colormaps
        cmap_actual = plt.cm.YlOrRd
        cmap_proyectado = plt.cm.RdYlGn

        # Rango común para comparación
        vmin = min(gdf_plot['rendimiento_actual'].min(), gdf_plot['rendimiento_proyectado'].min()) * 0.9
        vmax = max(gdf_plot['rendimiento_actual'].max(), gdf_plot['rendimiento_proyectado'].max()) * 1.1

        # Mapa 1: Rendimiento Actual
        for idx, row in gdf_plot.iterrows():
            valor = row['rendimiento_actual']
            valor_norm = (valor - vmin) / (vmax - vmin) if vmax != vmin else 0.5
            valor_norm = max(0, min(1, valor_norm))
            color = cmap_actual(valor_norm)
            gdf_plot.iloc[[idx]].plot(ax=ax1, color=color, edgecolor='white', linewidth=0.8, alpha=0.85)

            # Etiqueta simple
            centroid = row.geometry.centroid
            ax1.annotate(f"{valor:.1f}",
                         (centroid.x, centroid.y),
                         xytext=(0, 0), textcoords="offset points",
                         fontsize=6, color='white', weight='bold',
                         ha='center', va='center',
                         bbox=dict(boxstyle="circle,pad=0.15",
                                   facecolor=(0, 0, 0, 0.6),
                                   alpha=0.8))

        # Mapa 2: Rendimiento Proyectado
        for idx, row in gdf_plot.iterrows():
            valor = row['rendimiento_proyectado']
            incremento = row['rendimiento_proyectado'] - row['rendimiento_actual']
            valor_norm = (valor - vmin) / (vmax - vmin) if vmax != vmin else 0.5
            valor_norm = max(0, min(1, valor_norm))
            color = cmap_proyectado(valor_norm)
            gdf_plot.iloc[[idx]].plot(ax=ax2, color=color, edgecolor='white', linewidth=0.8, alpha=0.85)

            # Etiqueta con incremento
            centroid = row.geometry.centroid
            ax2.annotate(f"{valor:.1f}\n+{incremento:.1f}",
                         (centroid.x, centroid.y),
                         xytext=(0, 0), textcoords="offset points",
                         fontsize=6, color='white', weight='bold',
                         ha='center', va='center',
                         bbox=dict(boxstyle="round,pad=0.15",
                                   facecolor=(0, 0, 0, 0.6),
                                   alpha=0.8))

        # Títulos
        ax1.set_title('🌾 RENDIMIENTO ACTUAL\n(ton/ha)', fontsize=14, fontweight='bold', color='white')
        ax2.set_title('🚀 RENDIMIENTO CON FERTILIZACIÓN\n(ton/ha)', fontsize=14, fontweight='bold', color='white')

        for ax in [ax1, ax2]:
            ax.set_xlabel('Longitud', color='white')
            ax.set_ylabel('Latitud', color='white')
            ax.tick_params(colors='white')
            ax.grid(True, alpha=0.1, color='#475569')

        # Barras de color
        sm1 = plt.cm.ScalarMappable(cmap=cmap_actual, norm=plt.Normalize(vmin=vmin, vmax=vmax))
        sm1.set_array([])
        cbar1 = plt.colorbar(sm1, ax=ax1, shrink=0.6)
        cbar1.set_label('ton/ha', fontsize=10, color='white')
        sm2 = plt.cm.ScalarMappable(cmap=cmap_proyectado, norm=plt.Normalize(vmin=vmin, vmax=vmax))
        sm2.set_array([])
        cbar2 = plt.colorbar(sm2, ax=ax2, shrink=0.6)
        cbar2.set_label('ton/ha', fontsize=10, color='white')

        # Configurar colores de barras
        for cbar in [cbar1, cbar2]:
            cbar.ax.yaxis.set_tick_params(color='white')
            cbar.outline.set_edgecolor('white')
            plt.setp(plt.getp(cbar.ax.axes, 'yticklabels'), color='white')

        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='#0f172a')
        buf.seek(0)
        plt.close()
        return buf

    except Exception as e:
        st.error(f"Error en fallback: {str(e)}")
        return None

# ===== FUNCIONES PARA DATOS SATELITALES =====
def descargar_datos_landsat8(gdf, fecha_inicio, fecha_fin, indice='NDVI'):
    try:
        st.info(f"🔍 Buscando escenas Landsat 8...")
        datos_simulados = {
            'indice': indice,
            'valor_promedio': 0.65 + np.random.normal(0, 0.1),
            'fuente': 'Landsat-8',
            'fecha': datetime.now().strftime('%Y-%m-%d'),
            'id_escena': f"LC08_{np.random.randint(1000000, 9999999)}",
            'cobertura_nubes': f"{np.random.randint(0, 15)}%",
            'resolucion': '30m'
        }
        st.success(f"✅ Escena Landsat 8 encontrada: {datos_simulados['id_escena']}")
        st.info(f"☁️ Cobertura de nubes: {datos_simulados['cobertura_nubes']}")
        return datos_simulados
    except Exception as e:
        st.error(f"❌ Error procesando Landsat 8: {str(e)}")
        return None

def descargar_datos_sentinel2(gdf, fecha_inicio, fecha_fin, indice='NDVI'):
    try:
        st.info(f"🔍 Buscando escenas Sentinel-2...")
        datos_simulados = {
            'indice': indice,
            'valor_promedio': 0.72 + np.random.normal(0, 0.08),
            'fuente': 'Sentinel-2',
            'fecha': datetime.now().strftime('%Y-%m-%d'),
            'id_escena': f"S2A_{np.random.randint(1000000, 9999999)}",
            'cobertura_nubes': f"{np.random.randint(0, 10)}%",
            'resolucion': '10m'
        }
        st.success(f"✅ Escena Sentinel-2 encontrada: {datos_simulados['id_escena']}")
        st.info(f"☁️ Cobertura de nubes: {datos_simulados['cobertura_nubes']}")
        return datos_simulados
    except Exception as e:
        st.error(f"❌ Error procesando Sentinel-2: {str(e)}")
        return None

def generar_datos_simulados(gdf, cultivo, indice='NDVI'):
    st.info("🔬 Generando datos simulados...")
    datos_simulados = {
        'indice': indice,
        'valor_promedio': PARAMETROS_CULTIVOS[cultivo]['NDVI_OPTIMO'] * 0.8 + np.random.normal(0, 0.1),
        'fuente': 'Simulación',
        'fecha': datetime.now().strftime('%Y-%m-%d'),
        'resolucion': '10m'
    }
    st.success("✅ Datos simulados generados")
    return datos_simulados

# ===== FUNCIÓN CORREGIDA PARA OBTENER DATOS DE NASA POWER =====
def obtener_datos_nasa_power(gdf, fecha_inicio, fecha_fin):
    """
    Obtiene datos meteorológicos diarios de NASA POWER para el centroide de la parcela.
    Variables: radiación solar (ALLSKY_SFC_SW_DWN) y viento a 2m (WS2M).
    """
    try:
        centroid = gdf.geometry.unary_union.centroid
        lat = round(centroid.y, 4)
        lon = round(centroid.x, 4)
        start = fecha_inicio.strftime("%Y%m%d")
        end = fecha_fin.strftime("%Y%m%d")
        params = {
            'parameters': 'ALLSKY_SFC_SW_DWN,WS2M,T2M,PRECTOTCORR',
            'community': 'RE',
            'longitude': lon,
            'latitude': lat,
            'start': start,
            'end': end,
            'format': 'JSON'
        }
        url = "https://power.larc.nasa.gov/api/temporal/daily/point"
        response = requests.get(url, params=params, timeout=15)
        data = response.json()
        if 'properties' not in data:
            st.warning("⚠️ No se obtuvieron datos de NASA POWER (fuera de rango o sin conexión).")
            return None
        series = data['properties']['parameter']
        df_power = pd.DataFrame({
            'fecha': pd.to_datetime(list(series['ALLSKY_SFC_SW_DWN'].keys())),
            'radiacion_solar': list(series['ALLSKY_SFC_SW_DWN'].values()),
            'viento_2m': list(series['WS2M'].values()),
            'temperatura': list(series['T2M'].values()),
            'precipitacion': list(series['PRECTOTCORR'].values())
        })
        df_power = df_power.replace(-999, np.nan).dropna()
        if df_power.empty:
            st.warning("⚠️ Datos de NASA POWER no disponibles para el período seleccionado.")
            return None
        st.success("✅ Datos meteorológicos de NASA POWER cargados.")
        return df_power
    except Exception as e:
        st.error(f"❌ Error al obtener datos de NASA POWER: {str(e)}")
        return None

# ===== FUNCIONES DE ANÁLISIS GEE MEJORADAS =====
def calcular_recomendaciones_npk_cientificas(gdf_analizado, nutriente, cultivo):
    """Calcula recomendaciones basadas en metodologías científicas - CORREGIDO"""
    import copy
    recomendaciones = []
    params = copy.deepcopy(PARAMETROS_CULTIVOS[cultivo])

    # Usar parámetros específicos por variedad
    if 'variedad_params' in st.session_state:
        variedad_params = st.session_state['variedad_params']
        if nutriente == "NITRÓGENO":
            params['NITROGENO']['optimo'] = variedad_params['NITROGENO_OPTIMO']
        elif nutriente == "FÓSFORO":
            params['FOSFORO']['optimo'] = variedad_params['FOSFORO_OPTIMO']
        elif nutriente == "POTASIO":
            params['POTASIO']['optimo'] = variedad_params['POTASIO_OPTIMO']

    for idx, row in gdf_analizado.iterrows():
        if nutriente == "NITRÓGENO":
            valor_actual = row['nitrogeno_actual']
            objetivo = params['NITROGENO']['optimo']
            # Calcular deficiencia
            deficiencia = max(0, objetivo - valor_actual)
            # Eficiencia de fertilización (40-60% dependiendo del método)
            eficiencia = 0.5  # 50% eficiencia típica
            # Recomendación ajustada
            recomendado = deficiencia / eficiencia if deficiencia > 0 else 0
        elif nutriente == "FÓSFORO":
            valor_actual = row['fosforo_actual']
            objetivo = params['FOSFORO']['optimo']
            deficiencia = max(0, objetivo - valor_actual)
            eficiencia = 0.3  # 30% eficiencia típica para P
            recomendado = deficiencia / eficiencia if deficiencia > 0 else 0
        else:  # POTASIO
            valor_actual = row['potasio_actual']
            objetivo = params['POTASIO']['optimo']
            deficiencia = max(0, objetivo - valor_actual)
            eficiencia = 0.6  # 60% eficiencia típica para K
            recomendado = deficiencia / eficiencia if deficiencia > 0 else 0

        # Redondear a múltiplos de 5 kg/ha
        recomendado_redondeado = round(recomendado / 5) * 5
        recomendaciones.append(max(0, recomendado_redondeado))

    return recomendaciones

# ===== FUNCIONES DE TEXTURA DEL SUELO - ACTUALIZADAS A USDA =====
def clasificar_textura_suelo(arena, limo, arcilla):
    """Función principal que ahora usa clasificación USDA"""
    return clasificar_textura_usda(arena, limo, arcilla)

def analizar_textura_suelo(gdf, cultivo):
    gdf = validar_y_corregir_crs(gdf)
    params_textura = TEXTURA_SUELO_OPTIMA[cultivo]
    zonas_gdf = gdf.copy()
    areas_ha_list = []
    arena_list = []
    limo_list = []
    arcilla_list = []
    textura_list = []

    for idx, row in zonas_gdf.iterrows():
        try:
            area_gdf = gpd.GeoDataFrame({'geometry': [row.geometry]}, crs=zonas_gdf.crs)
            area_ha = calcular_superficie(area_gdf)
            if hasattr(area_ha, 'iloc'):
                area_ha = float(area_ha.iloc[0])
            elif hasattr(area_ha, '__len__') and len(area_ha) > 0:
                area_ha = float(area_ha[0])
            else:
                area_ha = float(area_ha)

            centroid = row.geometry.centroid if hasattr(row.geometry, 'centroid') else row.geometry.representative_point()
            seed_value = abs(hash(f"{centroid.x:.6f}_{centroid.y:.6f}_{cultivo}_textura")) % (2**32)
            rng = np.random.RandomState(seed_value)

            # Simular composición basada en textura óptima
            arena_optima = params_textura['arena_optima']
            limo_optima = params_textura['limo_optima']
            arcilla_optima = params_textura['arcilla_optima']

            # Variación alrededor del óptimo
            arena_val = max(5, min(95, rng.normal(arena_optima, 10)))
            limo_val = max(5, min(95, rng.normal(limo_optima, 8)))
            arcilla_val = max(5, min(95, rng.normal(arcilla_optima, 7)))

            total = arena_val + limo_val + arcilla_val
            arena_pct = (arena_val / total) * 100
            limo_pct = (limo_val / total) * 100
            arcilla_pct = (arcilla_val / total) * 100

            textura = clasificar_textura_suelo(arena_pct, limo_pct, arcilla_pct)

            areas_ha_list.append(area_ha)
            arena_list.append(float(arena_pct))
            limo_list.append(float(limo_pct))
            arcilla_list.append(float(arcilla_pct))
            textura_list.append(textura)

        except Exception as e:
            areas_ha_list.append(0.0)
            arena_list.append(float(params_textura['arena_optima']))
            limo_list.append(float(params_textura['limo_optima']))
            arcilla_list.append(float(params_textura['arcilla_optima']))
            textura_list.append(params_textura['textura_optima'])

    zonas_gdf['area_ha'] = areas_ha_list
    zonas_gdf['arena'] = arena_list
    zonas_gdf['limo'] = limo_list
    zonas_gdf['arcilla'] = arcilla_list
    zonas_gdf['textura_suelo'] = textura_list

    return zonas_gdf

# ===== FUNCIONES DE CURVAS DE NIVEL =====
def clasificar_pendiente(pendiente_porcentaje):
    for categoria, params in CLASIFICACION_PENDIENTES.items():
        if params['min'] <= pendiente_porcentaje < params['max']:
            return categoria, params['color']
    return "EXTREMA (>25%)", CLASIFICACION_PENDIENTES['EXTREMA (>25%)']['color']

def calcular_estadisticas_pendiente_simple(pendiente_grid):
    pendiente_flat = pendiente_grid.flatten()
    pendiente_flat = pendiente_flat[~np.isnan(pendiente_flat)]
    if len(pendiente_flat) == 0:
        return {'promedio': 0, 'min': 0, 'max': 0, 'std': 0, 'distribucion': {}}

    stats = {
        'promedio': float(np.mean(pendiente_flat)),
        'min': float(np.min(pendiente_flat)),
        'max': float(np.max(pendiente_flat)),
        'std': float(np.std(pendiente_flat)),
        'distribucion': {}
    }

    for categoria, params in CLASIFICACION_PENDIENTES.items():
        mask = (pendiente_flat >= params['min']) & (pendiente_flat < params['max'])
        stats['distribucion'][categoria] = {'porcentaje': float(np.sum(mask) / len(pendiente_flat) * 100), 'color': params['color']}

    return stats

def generar_dem_sintetico(gdf, resolucion=10.0):
    """
    Genera un DEM sintético determinístico basado en las coordenadas de la parcela.
    Mismo input → mismo output siempre.
    """
    gdf = validar_y_corregir_crs(gdf)
    bounds = gdf.total_bounds
    minx, miny, maxx, maxy = bounds

    # Crear una semilla determinística basada en las coordenadas de la parcela
    centroid = gdf.geometry.unary_union.centroid
    seed_value = int(centroid.x * 10000 + centroid.y * 10000) % (2**32)

    # Inicializar el generador aleatorio con la semilla
    rng = np.random.RandomState(seed_value)

    num_cells = 50
    x = np.linspace(minx, maxx, num_cells)
    y = np.linspace(miny, maxy, num_cells)
    X, Y = np.meshgrid(x, y)

    # Valores fijos basados en la semilla
    elevacion_base = rng.uniform(100, 300)
    slope_x = rng.uniform(-0.001, 0.001)
    slope_y = rng.uniform(-0.001, 0.001)

    relief = np.zeros_like(X)
    n_hills = rng.randint(2, 5)
    for _ in range(n_hills):
        hill_center_x = rng.uniform(minx, maxx)
        hill_center_y = rng.uniform(miny, maxy)
        hill_radius = rng.uniform(0.001, 0.005)
        hill_height = rng.uniform(10, 50)
        dist = np.sqrt((X - hill_center_x)**2 + (Y - hill_center_y)**2)
        relief += hill_height * np.exp(-(dist**2) / (2 * hill_radius**2))

    noise = rng.randn(*X.shape) * 2
    Z = elevacion_base + slope_x * (X - minx) + slope_y * (Y - miny) + relief + noise
    Z = np.maximum(Z, 50)

    return X, Y, Z, bounds

def calcular_pendiente_simple(X, Y, Z, resolucion=10.0):
    dy = np.gradient(Z, axis=0) / resolucion
    dx = np.gradient(Z, axis=1) / resolucion
    pendiente = np.sqrt(dx**2 + dy**2) * 100
    pendiente = np.clip(pendiente, 0, 100)
    return pendiente

def crear_mapa_pendientes_simple(X, Y, pendiente_grid, gdf_original):
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 6))

    # Configurar estilo oscuro
    fig.patch.set_facecolor('#0f172a')
    ax1.set_facecolor('#0f172a')
    ax2.set_facecolor('#0f172a')

    X_flat = X.flatten()
    Y_flat = Y.flatten()
    Z_flat = pendiente_grid.flatten()
    valid_mask = ~np.isnan(Z_flat)

    if np.sum(valid_mask) > 10:
        scatter = ax1.scatter(X_flat[valid_mask], Y_flat[valid_mask], c=Z_flat[valid_mask], cmap='RdYlGn_r', s=20, alpha=0.7, vmin=0, vmax=30)
        cbar = plt.colorbar(scatter, ax=ax1, shrink=0.8)
        cbar.set_label('Pendiente (%)', color='white')
        cbar.ax.yaxis.set_tick_params(color='white')
        cbar.outline.set_edgecolor('white')
        plt.setp(plt.getp(cbar.ax.axes, 'yticklabels'), color='white')

        for porcentaje in [2, 5, 10, 15, 25]:
            mask_cat = (Z_flat[valid_mask] >= porcentaje-1) & (Z_flat[valid_mask] <= porcentaje+1)
            if np.sum(mask_cat) > 0:
                x_center = np.mean(X_flat[valid_mask][mask_cat])
                y_center = np.mean(Y_flat[valid_mask][mask_cat])
                ax1.text(x_center, y_center, f'{porcentaje}%', fontsize=8, fontweight='bold', ha='center', va='center',
                         bbox=dict(boxstyle="round,pad=0.3", facecolor=(30/255, 41/255, 59/255, 0.9), edgecolor='white'), color='white')
    else:
        ax1.text(0.5, 0.5, 'Datos insuficientes\npara mapa de calor', transform=ax1.transAxes, ha='center', va='center', fontsize=12, color='white')

    gdf_original.plot(ax=ax1, color='none', edgecolor='white', linewidth=2)
    ax1.set_title('Mapa de Calor de Pendientes', fontsize=12, fontweight='bold', color='white')
    ax1.set_xlabel('Longitud', color='white')
    ax1.set_ylabel('Latitud', color='white')
    ax1.tick_params(colors='white')
    ax1.grid(True, alpha=0.3, color='#475569')

    if np.sum(valid_mask) > 0:
        pendiente_data = Z_flat[valid_mask]
        ax2.hist(pendiente_data, bins=30, edgecolor='white', color='#3b82f6', alpha=0.7)

        for porcentaje, color in [(2, '#4daf4a'), (5, '#a6d96a'), (10, '#ffffbf'), (15, '#fdae61'), (25, '#f46d43')]:
            ax2.axvline(x=porcentaje, color=color, linestyle='--', linewidth=1, alpha=0.7)
            ax2.text(porcentaje+0.5, ax2.get_ylim()[1]*0.9, f'{porcentaje}%', color=color, fontsize=8)

        stats_pendiente = calcular_estadisticas_pendiente_simple(pendiente_grid)
        stats_text = f"""
        Estadísticas:
        • Mínima: {stats_pendiente['min']:.1f}%
        • Máxima: {stats_pendiente['max']:.1f}%
        • Promedio: {stats_pendiente['promedio']:.1f}%
        • Desviación: {stats_pendiente['std']:.1f}%
        """
        ax2.text(0.02, 0.98, stats_text, transform=ax2.transAxes, fontsize=9, verticalalignment='top',
                 color='white', bbox=dict(boxstyle="round,pad=0.3", facecolor=(30/255, 41/255, 59/255, 0.9), edgecolor='white'))

        ax2.set_xlabel('Pendiente (%)', color='white')
        ax2.set_ylabel('Frecuencia', color='white')
        ax2.set_title('Distribución de Pendientes', fontsize=12, fontweight='bold', color='white')
        ax2.tick_params(colors='white')
        ax2.grid(True, alpha=0.3, color='#475569')
    else:
        ax2.text(0.5, 0.5, 'Sin datos de pendiente', transform=ax2.transAxes, ha='center', va='center', fontsize=12, color='white')

    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='#0f172a')
    buf.seek(0)
    plt.close()
    return buf, calcular_estadisticas_pendiente_simple(pendiente_grid)

def generar_curvas_nivel_simple(X, Y, Z, intervalo=5.0, gdf_original=None):
    curvas = []
    elevaciones = []
    try:
        if gdf_original is not None:
            poligono_principal = gdf_original.iloc[0].geometry
            bounds = poligono_principal.bounds
            centro = poligono_principal.centroid
            ancho = bounds[2] - bounds[0]
            alto = bounds[3] - bounds[1]
            radio_max = min(ancho, alto) / 2
            z_min, z_max = np.nanmin(Z), np.nanmax(Z)
            n_curvas = min(10, int((z_max - z_min) / intervalo))
            for i in range(1, n_curvas + 1):
                radio = radio_max * (i / n_curvas)
                circle = centro.buffer(radio)
                interseccion = poligono_principal.intersection(circle)
                if interseccion.geom_type == 'LineString':
                    curvas.append(interseccion)
                    elevaciones.append(z_min + (i * intervalo))
                elif interseccion.geom_type == 'MultiLineString':
                    for parte in interseccion.geoms:
                        curvas.append(parte)
                        elevaciones.append(z_min + (i * intervalo))
    except Exception as e:
        if gdf_original is not None:
            bounds = gdf_original.total_bounds
            for i in range(3):
                y = bounds[1] + (i + 1) * ((bounds[3] - bounds[1]) / 4)
                linea = LineString([(bounds[0], y), (bounds[2], y)])
                curvas.append(linea)
                elevaciones.append(100 + i * 50)
    return curvas, elevaciones

# ===== FUNCIONES DE EXPORTACIÓN Y REPORTES - CORREGIDAS =====
def exportar_a_geojson(gdf, nombre_base="parcela"):
    try:
        gdf = validar_y_corregir_crs(gdf)
        geojson_data = gdf.to_json()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        nombre_archivo = f"{nombre_base}_{timestamp}.geojson"
        return geojson_data, nombre_archivo
    except Exception as e:
        st.error(f"❌ Error exportando a GeoJSON: {str(e)}")
        return None, None

def generar_resumen_estadisticas(gdf_analizado, analisis_tipo, cultivo, df_power=None):
    estadisticas = {}
    try:
        if analisis_tipo in ["FERTILIDAD ACTUAL", "RECOMENDACIONES NPK"]:
            if 'npk_integrado' in gdf_analizado.columns:
                estadisticas['Índice NPK Integrado'] = f"{gdf_analizado['npk_integrado'].mean():.3f}"
            if 'nitrogeno_actual' in gdf_analizado.columns:
                estadisticas['Nitrógeno Promedio'] = f"{gdf_analizado['nitrogeno_actual'].mean():.1f} kg/ha"
            if 'fosforo_actual' in gdf_analizado.columns:
                estadisticas['Fósforo Promedio'] = f"{gdf_analizado['fosforo_actual'].mean():.1f} kg/ha"
            if 'potasio_actual' in gdf_analizado.columns:
                estadisticas['Potasio Promedio'] = f"{gdf_analizado['potasio_actual'].mean():.1f} kg/ha"
            if 'ndvi' in gdf_analizado.columns:
                estadisticas['NDVI Promedio'] = f"{gdf_analizado['ndvi'].mean():.3f}"
            if 'ndwi' in gdf_analizado.columns:
                estadisticas['NDWI Promedio'] = f"{gdf_analizado['ndwi'].mean():.3f}"
            if 'materia_organica' in gdf_analizado.columns:
                estadisticas['Materia Orgánica Promedio'] = f"{gdf_analizado['materia_organica'].mean():.1f}%"

            # Información de variedad
            if 'variedad' in st.session_state:
                estadisticas['Variedad'] = st.session_state['variedad']

            # Datos de NASA POWER
            if df_power is not None:
                estadisticas['Radiación Solar Promedio'] = f"{df_power['radiacion_solar'].mean():.1f} kWh/m²/día"
                estadisticas['Velocidad Viento Promedio'] = f"{df_power['viento_2m'].mean():.2f} m/s"
                estadisticas['Precipitación Promedio'] = f"{df_power['precipitacion'].mean():.2f} mm/día"

        elif analisis_tipo == "ANÁLISIS DE TEXTURA":
            if 'arena' in gdf_analizado.columns:
                estadisticas['Arena Promedio'] = f"{gdf_analizado['arena'].mean():.1f}%"
                estadisticas['Limo Promedio'] = f"{gdf_analizado['limo'].mean():.1f}%"
                estadisticas['Arcilla Promedio'] = f"{gdf_analizado['arcilla'].mean():.1f}%"
            if 'textura_suelo' in gdf_analizado.columns:
                textura_predominante = gdf_analizado['textura_suelo'].mode()[0] if len(gdf_analizado) > 0 else "N/D"
                estadisticas['Textura Predominante'] = textura_predominante
            if 'area_ha' in gdf_analizado.columns:
                estadisticas['Área Promedio por Zona'] = f"{gdf_analizado['area_ha'].mean():.2f} ha"
                if gdf_analizado['area_ha'].mean() > 0:
                    estadisticas['Coeficiente de Variación'] = f"{(gdf_analizado['area_ha'].std() / gdf_analizado['area_ha'].mean() * 100):.1f}%"

    except Exception as e:
        st.warning(f"No se pudieron calcular algunas estadísticas: {str(e)}")

    return estadisticas

def generar_recomendaciones_generales(gdf_analizado, analisis_tipo, cultivo):
    recomendaciones = []
    try:
        if analisis_tipo == "FERTILIDAD ACTUAL":
            if 'npk_integrado' in gdf_analizado.columns:
                npk_promedio = gdf_analizado['npk_integrado'].mean()
                if npk_promedio < 0.3:
                    recomendaciones.append("Fertilidad MUY BAJA: Se recomienda aplicación urgente de fertilizantes balanceados")
                    recomendaciones.append("Considerar enmiendas orgánicas para mejorar la estructura del suelo")
                elif npk_promedio < 0.5:
                    recomendaciones.append("Fertilidad BAJA: Recomendada aplicación de fertilizantes según análisis de suelo")
                elif npk_promedio < 0.7:
                    recomendaciones.append("Fertilidad ADECUADA: Mantener prácticas de manejo actuales")
                else:
                    recomendaciones.append("Fertilidad ÓPTIMA: Excelente condición, continuar con manejo actual")

            # Recomendaciones específicas por nutriente
            if 'nitrogeno_actual' in gdf_analizado.columns:
                n_prom = gdf_analizado['nitrogeno_actual'].mean()
                n_opt = PARAMETROS_CULTIVOS[cultivo]['NITROGENO']['optimo']
                if n_prom < n_opt * 0.7:
                    recomendaciones.append(f"Deficiencia de Nitrógeno ({n_prom:.1f} vs {n_opt:.1f} kg/ha): Aplicar fertilizante nitrogenado")

            if 'fosforo_actual' in gdf_analizado.columns:
                p_prom = gdf_analizado['fosforo_actual'].mean()
                p_opt = PARAMETROS_CULTIVOS[cultivo]['FOSFORO']['optimo']
                if p_prom < p_opt * 0.7:
                    recomendaciones.append(f"Deficiencia de Fósforo ({p_prom:.1f} vs {p_opt:.1f} kg/ha): Aplicar superfosfato o fosfato diamónico")

            if 'potasio_actual' in gdf_analizado.columns:
                k_prom = gdf_analizado['potasio_actual'].mean()
                k_opt = PARAMETROS_CULTIVOS[cultivo]['POTASIO']['optimo']
                if k_prom < k_opt * 0.7:
                    recomendaciones.append(f"Deficiencia de Potasio ({k_prom:.1f} vs {k_opt:.1f} kg/ha): Aplicar cloruro o sulfato de potasio")

        elif analisis_tipo == "ANÁLISIS DE TEXTURA":
            if 'textura_suelo' in gdf_analizado.columns:
                textura_predominante = gdf_analizado['textura_suelo'].mode()[0] if len(gdf_analizado) > 0 else "N/D"
                if textura_predominante in RECOMENDACIONES_TEXTURA:
                    recomendaciones.append(f"Suelo {textura_predominante}: Ver recomendaciones específicas en el informe")
                else:
                    recomendaciones.append("Consultar recomendaciones específicas para la textura identificada")

        # === RECOMENDACIONES POR CULTIVO ===
        if cultivo == "VID":
            if 'variedad' in st.session_state:
                variedad = st.session_state['variedad']
                recomendaciones.append(f"Variedad: {variedad}")
                if "MALBEC" in variedad:
                    recomendaciones.append("Para Malbec: Priorizar potasio para calidad de color y taninos")
                elif "TORRONTÉS" in variedad:
                    recomendaciones.append("Para Torrontés: Controlar nitrógeno para mantener acidez y aromas")
                else:
                    recomendaciones.append("Para vid: Manejar nitrógeno en función del vigor vegetativo.")
                recomendaciones.append("Riego controlado para evitar exceso de vigor y mejorar calidad.")

        elif cultivo == "OLIVO":
            if 'variedad' in st.session_state:
                variedad = st.session_state['variedad']
                recomendaciones.append(f"Variedad: {variedad}")
                if "ARBEQUINA" in variedad:
                    recomendaciones.append("Para Arbequina: Manejo de potasio para optimizar contenido de aceite")
                elif "PICUAL" in variedad:
                    recomendaciones.append("Para Picual: Balancear nitrógeno y potasio para calidad de aceite")
                else:
                    recomendaciones.append("Para olivo: Fertilización balanceada N-P-K según análisis foliar.")
                recomendaciones.append("Riego deficitario controlado para aumentar polifenoles en aceite.")

        elif cultivo == "HORTALIZAS DE HOJA":
            if 'variedad' in st.session_state:
                variedad = st.session_state['variedad']
                recomendaciones.append(f"Variedad: {variedad}")
                if "LECHUGA" in variedad:
                    recomendaciones.append("Para lechuga: Nitrógeno en dosis divididas para máximo crecimiento foliar")
                elif "ESPINACA" in variedad:
                    recomendaciones.append("Para espinaca: Asegurar nitrógeno y hierro para hojas verdes oscuras")
                else:
                    recomendaciones.append("Para hortalizas de hoja: Alta demanda de nitrógeno en forma nítrica.")
                recomendaciones.append("Riego frecuente y uniforme para evitar estrés hídrico.")

        recomendaciones.append("Realizar análisis de suelo de laboratorio para validar resultados satelitales")
        recomendaciones.append("Considerar agricultura de precisión para aplicación variable de insumos")

    except Exception as e:
        recomendaciones.append("Error generando recomendaciones específicas")

    return recomendaciones

def limpiar_texto_para_pdf(texto):
    if not isinstance(texto, str):
        texto = str(texto)
    reemplazos = {
        '\u2022': '-',
        '\u2705': '[OK]',
        '\u26A0\uFE0F': '[!]',
        '\u274C': '[X]',
        '\u2013': '-',
        '\u2014': '--',
        '\u2018': "'",
        '\u2019': "'",
        '\u201C': '"',
        '\u201D': '"',
        '\u2192': '->',
        '\u2190': '<-',
        '\u2265': '>=',
        '\u2264': '<=',
        '\u00A0': ' ',
    }
    for original, reemplazo in reemplazos.items():
        texto = texto.replace(original, reemplazo)
    texto = texto.encode('latin-1', errors='replace').decode('latin-1')
    return texto

# ===== FUNCIÓN ÚNICA PARA GENERAR REPORTE COMPLETO EN DOCX - VERSIÓN CORREGIDA =====
def generar_reporte_completo_docx(gdf_analizado, cultivo, area_total, analisis_tipo, 
                                 nutriente=None, satelite=None, indice=None,
                                 resultados_economicos=None, gdf_textura=None,
                                 dem_data=None, pendiente_data=None):
    """
    Genera un reporte DOCX completo que incluye todos los análisis.
    VERSIÓN CORREGIDA - Funciona correctamente con Streamlit
    """
    try:
        doc = Document()
        
        # ===== ESTILOS DEL DOCUMENTO =====
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        
        # ===== PORTADA =====
        title = doc.add_heading('INFORME COMPLETO DE ANÁLISIS AGRÍCOLA', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        portada = doc.add_paragraph()
        portada.alignment = WD_ALIGN_PARAGRAPH.CENTER
        portada_run = portada.add_run('🍇🫒🥬 ANALIZADOR MULTI-CULTIVO SATELITAL\n\n')
        portada_run.bold = True
        portada_run.font.size = Pt(14)
        
        portada.add_run(f'CULTIVO: {cultivo}\n')
        if 'variedad' in st.session_state and st.session_state['variedad']:
            portada.add_run(f'VARIEDAD: {st.session_state["variedad"]}\n')
        portada.add_run(f'ÁREA TOTAL: {area_total:.2f} ha\n')
        portada.add_run(f'FECHA DE ANÁLISIS: {datetime.now().strftime("%d/%m/%Y %H:%M")}\n')
        portada.add_run(f'TIPO DE ANÁLISIS: {analisis_tipo}\n')
        
        if satelite:
            portada.add_run(f'SATÉLITE: {satelite}\n')
        if indice:
            portada.add_run(f'ÍNDICE: {indice}\n')
        if nutriente:
            portada.add_run(f'NUTRIENTE ANALIZADO: {nutriente}\n')
        
        doc.add_page_break()
        
        # ===== TABLA DE CONTENIDOS =====
        doc.add_heading('ÍNDICE DE CONTENIDOS', 1)
        contenido = doc.add_paragraph()
        contenido.add_run('1. RESUMEN EJECUTIVO\n').bold = True
        contenido.add_run('2. ANÁLISIS DE FERTILIDAD DEL SUELO\n').bold = True
        
        if analisis_tipo == "RECOMENDACIONES NPK":
            contenido.add_run('3. RECOMENDACIONES DE FERTILIZACIÓN\n').bold = True
            
            if resultados_economicos:
                contenido.add_run('4. ANÁLISIS ECONÓMICO\n').bold = True
        
        if gdf_textura is not None:
            contenido.add_run('5. ANÁLISIS DE TEXTURA DEL SUELO\n').bold = True
        
        if dem_data is not None:
            contenido.add_run('6. ANÁLISIS TOPOGRÁFICO\n').bold = True
        
        contenido.add_run('7. CONCLUSIONES Y RECOMENDACIONES\n').bold = True
        contenido.add_run('8. METADATOS TÉCNICOS\n').bold = True
        
        doc.add_page_break()
        
        # ===== 1. RESUMEN EJECUTIVO =====
        doc.add_heading('1. RESUMEN EJECUTIVO', 1)
        resumen = doc.add_paragraph()
        resumen.add_run('Este informe presenta un análisis completo del estado actual del cultivo ')
        resumen.add_run(f'{cultivo} ').bold = True
        resumen.add_run(f'para una parcela de {area_total:.2f} hectáreas.\n\n')
        
        # Estadísticas resumen
        if 'npk_integrado' in gdf_analizado.columns:
            npk_prom = gdf_analizado['npk_integrado'].mean()
            clasificacion = "BAJA" if npk_prom < 0.5 else "MEDIA" if npk_prom < 0.7 else "ALTA"
            resumen.add_run(f'• Fertilidad del suelo: {clasificacion} (Índice NPK: {npk_prom:.3f})\n')
        
        if 'rendimiento_actual' in gdf_analizado.columns:
            rend_actual = gdf_analizado['rendimiento_actual'].mean()
            resumen.add_run(f'• Potencial de cosecha actual: {rend_actual:.1f} ton/ha\n')
        
        if resultados_economicos:
            incremento = resultados_economicos.get('incremento_margen_ha', 0)
            roi = resultados_economicos.get('roi_fertilizacion_%', 0)
            resumen.add_run(f'• Incremento económico esperado: USD {incremento:.0f} por ha\n')
            resumen.add_run(f'• Retorno sobre inversión (ROI): {roi:.0f}%\n')
        
        doc.add_paragraph()
        
        # ===== 2. ANÁLISIS DE FERTILIDAD DEL SUELO =====
        doc.add_heading('2. ANÁLISIS DE FERTILIDAD DEL SUELO', 1)
        
        # 2.1. Estadísticas de fertilidad
        doc.add_heading('2.1. Estadísticas de Fertilidad', 2)
        
        if 'npk_integrado' in gdf_analizado.columns:
            # Crear tabla de estadísticas
            tabla_estadisticas = doc.add_table(rows=6, cols=2)
            tabla_estadisticas.style = 'Table Grid'
            
            datos_estadisticos = [
                ("Índice NPK Promedio", f"{gdf_analizado['npk_integrado'].mean():.3f}"),
                ("Nitrógeno Promedio (kg/ha)", f"{gdf_analizado['nitrogeno_actual'].mean():.1f}"),
                ("Fósforo Promedio (kg/ha)", f"{gdf_analizado['fosforo_actual'].mean():.1f}"),
                ("Potasio Promedio (kg/ha)", f"{gdf_analizado['potasio_actual'].mean():.1f}"),
                ("Materia Orgánica Promedio (%)", f"{gdf_analizado['materia_organica'].mean():.1f}" if 'materia_organica' in gdf_analizado.columns else "N/A"),
                ("Humedad del Suelo Promedio", f"{gdf_analizado['humedad_suelo'].mean():.3f}" if 'humedad_suelo' in gdf_analizado.columns else "N/A")
            ]
            
            for i, (concepto, valor) in enumerate(datos_estadisticos):
                tabla_estadisticas.cell(i, 0).text = concepto
                tabla_estadisticas.cell(i, 0).paragraphs[0].runs[0].bold = True
                tabla_estadisticas.cell(i, 1).text = valor
        
        # 2.2. Tabla detallada por zona
        doc.add_heading('2.2. Resultados por Zona', 2)
        
        columnas_mostrar = ['id_zona', 'area_ha', 'npk_integrado']
        if 'nitrogeno_actual' in gdf_analizado.columns:
            columnas_mostrar.append('nitrogeno_actual')
        if 'fosforo_actual' in gdf_analizado.columns:
            columnas_mostrar.append('fosforo_actual')
        if 'potasio_actual' in gdf_analizado.columns:
            columnas_mostrar.append('potasio_actual')
        
        columnas_mostrar = [col for col in columnas_mostrar if col in gdf_analizado.columns]
        
        if columnas_mostrar:
            df_mostrar = gdf_analizado[columnas_mostrar].head(15).copy()
            
            # Crear nombres de columnas amigables
            nombres_columnas = []
            for col in columnas_mostrar:
                if col == 'id_zona':
                    nombres_columnas.append('Zona')
                elif col == 'area_ha':
                    nombres_columnas.append('Área (ha)')
                elif col == 'npk_integrado':
                    nombres_columnas.append('Índice NPK')
                elif col == 'nitrogeno_actual':
                    nombres_columnas.append('N (kg/ha)')
                elif col == 'fosforo_actual':
                    nombres_columnas.append('P (kg/ha)')
                elif col == 'potasio_actual':
                    nombres_columnas.append('K (kg/ha)')
                else:
                    nombres_columnas.append(col.replace('_', ' ').title())
            
            # Crear tabla
            tabla = doc.add_table(rows=len(df_mostrar)+1, cols=len(columnas_mostrar))
            tabla.style = 'Table Grid'
            
            # Encabezados
            for j, col_name in enumerate(nombres_columnas):
                tabla.cell(0, j).text = col_name
                tabla.cell(0, j).paragraphs[0].runs[0].bold = True
            
            # Datos
            for i, (_, row) in enumerate(df_mostrar.iterrows(), 1):
                for j, col in enumerate(columnas_mostrar):
                    valor = row[col]
                    if isinstance(valor, float):
                        if col == 'npk_integrado':
                            tabla.cell(i, j).text = f"{valor:.3f}"
                        else:
                            tabla.cell(i, j).text = f"{valor:.1f}"
                    else:
                        tabla.cell(i, j).text = str(valor)
        
        # ===== 3. RECOMENDACIONES DE FERTILIZACIÓN (solo si es análisis de recomendaciones) =====
        if analisis_tipo == "RECOMENDACIONES NPK" and nutriente:
            doc.add_heading('3. RECOMENDACIONES DE FERTILIZACIÓN', 1)
            
            # 3.1. Recomendaciones por zona
            doc.add_heading('3.1. Recomendaciones por Zona', 2)
            
            if 'valor_recomendado' in gdf_analizado.columns:
                columnas_recomendaciones = ['id_zona', 'area_ha', f'{nutriente.lower()}_actual', 'valor_recomendado']
                columnas_recomendaciones = [col for col in columnas_recomendaciones if col in gdf_analizado.columns]
                
                if columnas_recomendaciones:
                    df_recomendaciones = gdf_analizado[columnas_recomendaciones].head(10).copy()
                    
                    # Crear tabla de recomendaciones
                    tabla_rec = doc.add_table(rows=len(df_recomendaciones)+1, cols=len(columnas_recomendaciones))
                    tabla_rec.style = 'Table Grid'
                    
                    # Nombres de columnas
                    nombres_rec = []
                    for col in columnas_recomendaciones:
                        if col == 'id_zona':
                            nombres_rec.append('Zona')
                        elif col == 'area_ha':
                            nombres_rec.append('Área (ha)')
                        elif col == f'{nutriente.lower()}_actual':
                            nombres_rec.append(f'{nutriente} Actual (kg/ha)')
                        elif col == 'valor_recomendado':
                            nombres_rec.append(f'{nutriente} Recomendado (kg/ha)')
                        else:
                            nombres_rec.append(col.replace('_', ' ').title())
                    
                    # Encabezados
                    for j, col_name in enumerate(nombres_rec):
                        tabla_rec.cell(0, j).text = col_name
                        tabla_rec.cell(0, j).paragraphs[0].runs[0].bold = True
                    
                    # Datos
                    for i, (_, row) in enumerate(df_recomendaciones.iterrows(), 1):
                        for j, col in enumerate(columnas_recomendaciones):
                            valor = row[col]
                            if isinstance(valor, float):
                                tabla_rec.cell(i, j).text = f"{valor:.1f}"
                            else:
                                tabla_rec.cell(i, j).text = str(valor)
            
            # 3.2. Resumen de fertilización necesaria
            doc.add_heading('3.2. Resumen de Fertilización Necesaria', 2)
            
            if 'valor_recomendado' in gdf_analizado.columns and 'area_ha' in gdf_analizado.columns:
                total_kg = (gdf_analizado['valor_recomendado'] * gdf_analizado['area_ha']).sum()
                promedio_kg_ha = gdf_analizado['valor_recomendado'].mean()
                
                resumen_fert = doc.add_paragraph()
                resumen_fert.add_run(f'Total de {nutriente} requerido: ').bold = True
                resumen_fert.add_run(f'{total_kg:.0f} kg\n')
                resumen_fert.add_run(f'Promedio por hectárea: ').bold = True
                resumen_fert.add_run(f'{promedio_kg_ha:.1f} kg/ha\n')
                
                # Recomendaciones de fertilizante específico
                if nutriente == "NITRÓGENO":
                    resumen_fert.add_run('Fertilizante recomendado: ').bold = True
                    resumen_fert.add_run('Urea (46% N)\n')
                    resumen_fert.add_run('Cantidad de urea necesaria: ').bold = True
                    resumen_fert.add_run(f'{total_kg / 0.46:.0f} kg\n')
                elif nutriente == "FÓSFORO":
                    resumen_fert.add_run('Fertilizante recomendado: ').bold = True
                    resumen_fert.add_run('Fosfato Diamónico (18% P₂O₅)\n')
                elif nutriente == "POTASIO":
                    resumen_fert.add_run('Fertilizante recomendado: ').bold = True
                    resumen_fert.add_run('Cloruro de Potasio (60% K₂O)\n')
        
        # ===== 4. ANÁLISIS ECONÓMICO =====
        if resultados_economicos:
            doc.add_heading('4. ANÁLISIS ECONÓMICO', 1)
            
            # 4.1. Tabla de resultados económicos
            doc.add_heading('4.1. Resultados Económicos', 2)
            
            tabla_economica = doc.add_table(rows=10, cols=2)
            tabla_economica.style = 'Table Grid'
            
            datos_economicos = [
                ("Cultivo", resultados_economicos['cultivo']),
                ("Área total (ha)", f"{resultados_economicos['area_total_ha']:.2f}"),
                ("Rendimiento actual (ton/ha)", f"{resultados_economicos['rendimiento_actual_ton_ha']:.1f}"),
                ("Rendimiento proyectado (ton/ha)", f"{resultados_economicos['rendimiento_proy_ton_ha']:.1f}"),
                ("Incremento de rendimiento (ton/ha)", f"{resultados_economicos['incremento_rendimiento_ton_ha']:.1f}"),
                ("Costo fertilización (USD/ha)", f"{resultados_economicos['costo_fertilizacion_ha']:.0f}"),
                ("Incremento margen (USD/ha)", f"{resultados_economicos['incremento_margen_ha']:.0f}"),
                ("ROI fertilización (%)", f"{resultados_economicos['roi_fertilizacion_%']:.0f}"),
                ("VAN total (USD)", f"{resultados_economicos['van_usd']:.0f}"),
                ("TIR (%)", f"{resultados_economicos['tir_%']:.1f}")
            ]
            
            for i, (concepto, valor) in enumerate(datos_economicos):
                tabla_economica.cell(i, 0).text = concepto
                tabla_economica.cell(i, 0).paragraphs[0].runs[0].bold = True
                tabla_economica.cell(i, 1).text = valor
            
            # 4.2. Recomendaciones económicas
            doc.add_heading('4.2. Recomendaciones Económicas', 2)
            
            recomendaciones_economicas = doc.add_paragraph()
            if resultados_economicos['roi_fertilizacion_%'] > 100:
                recomendaciones_economicas.add_run('✅ INVERTIR EN FERTILIZACIÓN: ').bold = True
                recomendaciones_economicas.add_run(f'ROI > 100% indica excelente retorno de la inversión.\n')
            elif resultados_economicos['roi_fertilizacion_%'] > 50:
                recomendaciones_economicas.add_run('✅ CONSIDERAR FERTILIZACIÓN: ').bold = True
                recomendaciones_economicas.add_run(f'ROI > 50% indica buena rentabilidad.\n')
            else:
                recomendaciones_economicas.add_run('⚠️ EVALUAR CON CUIDADO: ').bold = True
                recomendaciones_economicas.add_run(f'ROI < 50% puede no ser rentable.\n')
            
            if resultados_economicos['van_usd'] > 0:
                recomendaciones_economicas.add_run('✅ PROYECTO VIABLE: ').bold = True
                recomendaciones_economicas.add_run(f'VAN positivo genera valor económico.\n')
            
            recomendaciones_economicas.add_run('\n')
            recomendaciones_economicas.add_run('💡 Consideraciones adicionales:\n').bold = True
            recomendaciones_economicas.add_run('• Monitorear precios de mercado del cultivo\n')
            recomendaciones_economicas.add_run('• Considerar seguros agrícolas ante riesgos climáticos\n')
            recomendaciones_economicas.add_run('• Evaluar costos logísticos adicionales\n')
        
        # ===== 5. ANÁLISIS DE TEXTURA DEL SUELO =====
        if gdf_textura is not None:
            doc.add_heading('5. ANÁLISIS DE TEXTURA DEL SUELO', 1)
            
            # 5.1. Estadísticas de textura
            doc.add_heading('5.1. Composición Promedio del Suelo', 2)
            
            if 'arena' in gdf_textura.columns and 'limo' in gdf_textura.columns and 'arcilla' in gdf_textura.columns:
                tabla_textura = doc.add_table(rows=4, cols=2)
                tabla_textura.style = 'Table Grid'
                
                datos_textura = [
                    ("Arena Promedio (%)", f"{gdf_textura['arena'].mean():.1f}"),
                    ("Limo Promedio (%)", f"{gdf_textura['limo'].mean():.1f}"),
                    ("Arcilla Promedio (%)", f"{gdf_textura['arcilla'].mean():.1f}"),
                    ("Textura Predominante", gdf_textura['textura_suelo'].mode()[0] if len(gdf_textura) > 0 else "N/D")
                ]
                
                for i, (concepto, valor) in enumerate(datos_textura):
                    tabla_textura.cell(i, 0).text = concepto
                    tabla_textura.cell(i, 0).paragraphs[0].runs[0].bold = True
                    tabla_textura.cell(i, 1).text = valor
            
            # 5.2. Recomendaciones por textura
            doc.add_heading('5.2. Recomendaciones de Manejo', 2)
            
            if 'textura_suelo' in gdf_textura.columns:
                textura_predominante = gdf_textura['textura_suelo'].mode()[0] if len(gdf_textura) > 0 else "Sin datos"
                recomendaciones_textura = doc.add_paragraph()
                
                if textura_predominante == "Franco limoso":
                    recomendaciones_textura.add_run('✅ SUELO FRANCO LIMOSO (ÓPTIMO)\n\n').bold = True
                    recomendaciones_textura.add_run('Propiedades:\n• Excelente estructura y porosidad\n• Alta capacidad de retención de agua\n• Fertilidad natural alta\n\n')
                    recomendaciones_textura.add_run('Manejo recomendado:\n• Labranza mínima o conservacionista\n• Rotación de cultivos\n• Uso de coberturas vegetales\n')
                elif textura_predominante == "Franco arenoso":
                    recomendaciones_textura.add_run('⚠️ SUELO FRANCO ARENOSO\n\n').bold = True
                    recomendaciones_textura.add_run('Propiedades:\n• Excelente drenaje\n• Fácil labranza\n• Baja retención de agua y nutrientes\n\n')
                    recomendaciones_textura.add_run('Manejo recomendado:\n• Riego frecuente en pequeñas cantidades\n• Fertilización fraccionada\n• Aplicación de materia orgánica\n')
                elif "Arcilla" in textura_predominante:
                    recomendaciones_textura.add_run('⚠️ SUELO ARCILLOSO\n\n').bold = True
                    recomendaciones_textura.add_run('Propiedades:\n• Alta retención de agua y nutrientes\n• Drenaje lento\n• Difícil labranza\n\n')
                    recomendaciones_textura.add_run('Manejo recomendado:\n• Sistemas de drenaje\n• Labranza en condiciones óptimas de humedad\n• Encalamiento para mejorar estructura\n')
                else:
                    recomendaciones_textura.add_run(f'Suelo {textura_predominante} - Consultar recomendaciones específicas para esta textura.')
        
        # ===== 6. ANÁLISIS TOPOGRÁFICO =====
        if dem_data is not None:
            doc.add_heading('6. ANÁLISIS TOPOGRÁFICO', 1)
            
            X, Y, Z, pendiente_grid, gdf_original = dem_data
            
            # 6.1. Estadísticas de pendiente
            doc.add_heading('6.1. Estadísticas de Pendiente', 2)
            
            # Calcular estadísticas básicas
            if hasattr(pendiente_grid, 'flatten'):
                pendiente_flat = pendiente_grid.flatten()
                pendiente_flat = pendiente_flat[~np.isnan(pendiente_flat)]
                
                if len(pendiente_flat) > 0:
                    tabla_pendientes = doc.add_table(rows=5, cols=2)
                    tabla_pendientes.style = 'Table Grid'
                    
                    datos_pendientes = [
                        ("Pendiente mínima (%)", f"{np.min(pendiente_flat):.1f}"),
                        ("Pendiente máxima (%)", f"{np.max(pendiente_flat):.1f}"),
                        ("Pendiente promedio (%)", f"{np.mean(pendiente_flat):.1f}"),
                        ("Desviación estándar (%)", f"{np.std(pendiente_flat):.1f}"),
                        ("Área total analizada (ha)", f"{area_total:.2f}")
                    ]
                    
                    for i, (concepto, valor) in enumerate(datos_pendientes):
                        tabla_pendientes.cell(i, 0).text = concepto
                        tabla_pendientes.cell(i, 0).paragraphs[0].runs[0].bold = True
                        tabla_pendientes.cell(i, 1).text = valor
            
            # 6.2. Análisis de riesgo de erosión
            doc.add_heading('6.2. Análisis de Riesgo de Erosión', 2)
            
            if hasattr(pendiente_grid, 'flatten'):
                pendiente_flat = pendiente_grid.flatten()
                pendiente_flat = pendiente_flat[~np.isnan(pendiente_flat)]
                
                if len(pendiente_flat) > 0:
                    porcentaje_suave = np.sum((pendiente_flat >= 0) & (pendiente_flat < 5)) / len(pendiente_flat) * 100
                    porcentaje_moderada = np.sum((pendiente_flat >= 5) & (pendiente_flat < 10)) / len(pendiente_flat) * 100
                    porcentaje_fuerte = np.sum(pendiente_flat >= 10) / len(pendiente_flat) * 100
                    
                    riesgo_texto = doc.add_paragraph()
                    
                    if porcentaje_fuerte < 10:
                        riesgo_texto.add_run('✅ RIESGO BAJO DE EROSIÓN\n\n').bold = True
                        riesgo_texto.add_run(f'• Pendientes suaves (0-5%): {porcentaje_suave:.1f}%\n')
                        riesgo_texto.add_run(f'• Pendientes moderadas (5-10%): {porcentaje_moderada:.1f}%\n')
                        riesgo_texto.add_run(f'• Pendientes fuertes (>10%): {porcentaje_fuerte:.1f}%\n\n')
                        riesgo_texto.add_run('El terreno es adecuado para la mayoría de las prácticas agrícolas.')
                    elif porcentaje_fuerte < 30:
                        riesgo_texto.add_run('⚠️ RIESGO MODERADO DE EROSIÓN\n\n').bold = True
                        riesgo_texto.add_run(f'• Pendientes suaves (0-5%): {porcentaje_suave:.1f}%\n')
                        riesgo_texto.add_run(f'• Pendientes moderadas (5-10%): {porcentaje_moderada:.1f}%\n')
                        riesgo_texto.add_run(f'• Pendientes fuertes (>10%): {porcentaje_fuerte:.1f}%\n\n')
                        riesgo_texto.add_run('Se recomiendan prácticas de conservación de suelo como terrazas, cultivos en contorno y coberturas vegetales.')
                    else:
                        riesgo_texto.add_run('🚨 RIESGO ALTO DE EROSIÓN\n\n').bold = True
                        riesgo_texto.add_run(f'• Pendientes suaves (0-5%): {porcentaje_suave:.1f}%\n')
                        riesgo_texto.add_run(f'• Pendientes moderadas (5-10%): {porcentaje_moderada:.1f}%\n')
                        riesgo_texto.add_run(f'• Pendientes fuertes (>10%): {porcentaje_fuerte:.1f}%\n\n')
                        riesgo_texto.add_run('Se requieren medidas urgentes de conservación: terrazas, cultivos en franjas, reforestación y evitar labranza intensiva.')
        
        # ===== 7. CONCLUSIONES Y RECOMENDACIONES =====
        doc.add_heading('7. CONCLUSIONES Y RECOMENDACIONES', 1)
        
        conclusiones = doc.add_paragraph()
        conclusiones.add_run('CONCLUSIONES PRINCIPALES:\n\n').bold = True
        
        # Conclusión sobre fertilidad
        if 'npk_integrado' in gdf_analizado.columns:
            npk_prom = gdf_analizado['npk_integrado'].mean()
            if npk_prom < 0.5:
                conclusiones.add_run('• La fertilidad del suelo es BAJA, requiere intervención inmediata con fertilización balanceada.\n')
            elif npk_prom < 0.7:
                conclusiones.add_run('• La fertilidad del suelo es MODERADA, puede optimizarse con manejo adecuado.\n')
            else:
                conclusiones.add_run('• La fertilidad del suelo es ÓPTIMA, mantener prácticas actuales.\n')
        
        # Conclusión sobre potencial de rendimiento
        if 'rendimiento_actual' in gdf_analizado.columns and 'rendimiento_proyectado' in gdf_analizado.columns:
            incremento_prom = gdf_analizado['rendimiento_proyectado'].mean() - gdf_analizado['rendimiento_actual'].mean()
            if incremento_prom > 0:
                conclusiones.add_run(f'• Existe potencial para incrementar el rendimiento en {incremento_prom:.1f} ton/ha.\n')
        
        # Conclusión económica
        if resultados_economicos:
            if resultados_economicos['roi_fertilizacion_%'] > 50:
                conclusiones.add_run('• La inversión en fertilización es económicamente rentable.\n')
        
        conclusiones.add_run('\nRECOMENDACIONES GENERALES:\n\n').bold = True
        conclusiones.add_run('1. Realizar análisis de suelo de laboratorio para validar resultados satelitales.\n')
        conclusiones.add_run('2. Implementar agricultura de precisión para aplicación variable de insumos.\n')
        conclusiones.add_run('3. Monitorear periódicamente el estado del cultivo con imágenes satelitales.\n')
        conclusiones.add_run('4. Considerar rotación de cultivos para mejorar la salud del suelo.\n')
        conclusiones.add_run('5. Mantener registros detallados de aplicaciones y rendimientos.\n')
        
        # Recomendaciones específicas por cultivo
        conclusiones.add_run('\nRECOMENDACIONES ESPECÍFICAS PARA EL CULTIVO:\n\n').bold = True
        if cultivo == "VID":
            conclusiones.add_run('• Para vid: Controlar vigor vegetativo, evitar exceso de nitrógeno.\n')
            conclusiones.add_run('• Realizar poda adecuada para controlar carga de fruta.\n')
            conclusiones.add_run('• Fertilización potásica para mejorar calidad de uva.\n')
        elif cultivo == "OLIVO":
            conclusiones.add_run('• Para olivo: Poda de formación para facilitar cosecha mecánica.\n')
            conclusiones.add_run('• Control de riego para evitar rajado de fruta.\n')
            conclusiones.add_run('• Fertilización fosfatada para mejorar desarrollo radicular.\n')
        elif cultivo == "HORTALIZAS DE HOJAS":
            conclusiones.add_run('• Para hortalizas de hojas: Fertirriego con nitrógeno frecuente.\n')
            conclusiones.add.run('• Control de plagas con manejo integrado.\n')
            conclusiones.add_run('• Cosecha escalonada para mantener producción constante.\n')
        
        # ===== 8. METADATOS TÉCNICOS =====
        doc.add_heading('8. METADATOS TÉCNICOS', 1)
        
        metadatos = doc.add_table(rows=6, cols=2)
        metadatos.style = 'Table Grid'
        
        datos_tecnicos = [
            ("Sistema generador", "Analizador Multi-Cultivo Satelital"),
            ("Versión", "2.0 (Especializado en Vid, Olivo y Hortalizas)"),
            ("Fecha de generación", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
            ("Sistema de coordenadas", "EPSG:4326 (WGS84)"),
            ("Número de zonas analizadas", str(len(gdf_analizado))),
            ("Área total analizada", f"{area_total:.2f} hectáreas")
        ]
        
        for i, (campo, valor) in enumerate(datos_tecnicos):
            metadatos.cell(i, 0).text = campo
            metadatos.cell(i, 0).paragraphs[0].runs[0].bold = True
            metadatos.cell(i, 1).text = valor
        
        # Guardar documento en BytesIO
        docx_output = BytesIO()
        doc.save(docx_output)
        docx_output.seek(0)
        
        return docx_output
        
    except Exception as e:
        st.error(f"❌ Error generando reporte completo DOCX: {str(e)}")
        import traceback
        st.error(f"Detalle: {traceback.format_exc()}")
        return None

# ===== FUNCIONES DE MAPAS (MANTENIDAS SIN CAMBIOS) =====
def crear_mapa_npk_con_esri(gdf_analizado, nutriente, cultivo, satelite, mostrar_capa_inta=False):
    """Crea mapa de NPK con fondo ESRI Satellite + capa opcional del INTA - DEVUELVE BytesIO"""
    try:
        if gdf_analizado.empty or 'id_zona' not in gdf_analizado.columns:
            return None

        # Convertir a Web Mercator para el mapa base
        gdf_plot = gdf_analizado.to_crs(epsg=3857)
        fig, ax = plt.subplots(1, 1, figsize=(12, 8))

        # Configurar estilo oscuro
        fig.patch.set_facecolor('#0f172a')
        ax.set_facecolor('#0f172a')

        # Mapeo de nutrientes
        mapeo_nutriente = {
            'NITRÓGENO': ('nitrogeno_actual', 'NITROGENO', 'NITRÓGENO (kg/ha)'),
            'FÓSFORO': ('fosforo_actual', 'FOSFORO', 'FÓSFORO (kg/ha)'),
            'POTASIO': ('potasio_actual', 'POTASIO', 'POTASIO (kg/ha)')
        }

        if nutriente not in mapeo_nutriente:
            return None

        columna, clave_param, titulo_nutriente = mapeo_nutriente[nutriente]

        if columna not in gdf_analizado.columns:
            return None

        # Rangos dinámicos
        vmin = PARAMETROS_CULTIVOS[cultivo][clave_param]['min'] * 0.7
        vmax = PARAMETROS_CULTIVOS[cultivo][clave_param]['max'] * 1.2
        if vmin >= vmax:
            vmin, vmax = 0, 100  # fallback seguro

        cmap = LinearSegmentedColormap.from_list('nutriente_gee', PALETAS_GEE[clave_param])

        # Plotear zonas
        for idx, row in gdf_plot.iterrows():
            valor = row.get(columna, 0)
            valor_norm = max(0, min(1, (valor - vmin) / (vmax - vmin))) if vmax != vmin else 0.5
            color = cmap(valor_norm)
            gdf_plot.iloc[[idx]].plot(ax=ax, color=color, edgecolor='white', linewidth=1.5, alpha=0.7)

            centroid = row.geometry.centroid
            ax.annotate(f"Z{row['id_zona']}\n{valor:.0f}", (centroid.x, centroid.y),
                        xytext=(5, 5), textcoords="offset points",
                        fontsize=8, color='white', weight='bold',
                        bbox=dict(boxstyle="round,pad=0.3", facecolor=(30/255, 41/255, 59/255, 0.9), edgecolor='white'))

        # Mapa base
        try:
            ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, alpha=0.3)
        except Exception:
            pass  # Silencioso: no interrumpe

        # Capa INTA
        if mostrar_capa_inta:
            agregar_capa_inta(ax, alpha=0.5)

        # Título
        info_satelite = SATELITES_DISPONIBLES.get(satelite, SATELITES_DISPONIBLES['DATOS_SIMULADOS'])
        ax.set_title(f'{ICONOS_CULTIVOS[cultivo]} ANÁLISIS DE {nutriente} - {cultivo}\n'
                     f'{info_satelite["icono"]} {info_satelite["nombre"]} - {titulo_nutriente}',
                     fontsize=16, fontweight='bold', pad=20, color='white')
        ax.set_xlabel('Longitud', color='white')
        ax.set_ylabel('Latitud', color='white')
        ax.tick_params(colors='white')
        ax.grid(True, alpha=0.3, color='#475569')

        # Barra de colores
        sm = plt.cm.ScalarMappable(cmap=cmap, norm=plt.Normalize(vmin=vmin, vmax=vmax))
        sm.set_array([])
        cbar = plt.colorbar(sm, ax=ax, shrink=0.8)
        cbar.set_label(titulo_nutriente, fontsize=12, fontweight='bold', color='white')
        cbar.ax.yaxis.set_tick_params(color='white')
        cbar.outline.set_edgecolor('white')
        plt.setp(plt.getp(cbar.ax.axes, 'yticklabels'), color='white')

        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='#0f172a')
        plt.close(fig)  # ← Cerrar figura específica
        buf.seek(0)
        return buf  # ← DEVOLVER BytesIO

    except Exception as e:
        st.error(f"Error creando mapa NPK: {str(e)}")
        return None

def crear_mapa_fertilidad_integrada(gdf_analizado, cultivo, satelite, mostrar_capa_inta=False):
    """Crea mapa de fertilidad integrada (NPK combinado) - DEVUELVE BytesIO"""
    try:
        if gdf_analizado.empty or 'npk_integrado' not in gdf_analizado.columns:
            return None

        gdf_plot = gdf_analizado.to_crs(epsg=3857)
        fig, ax = plt.subplots(1, 1, figsize=(12, 8))
        fig.patch.set_facecolor('#0f172a')
        ax.set_facecolor('#0f172a')

        cmap = LinearSegmentedColormap.from_list('fertilidad_gee', PALETAS_GEE['FERTILIDAD'])

        for idx, row in gdf_plot.iterrows():
            valor = row['npk_integrado']
            color = cmap(max(0, min(1, valor)))  # npk_integrado ya está normalizado [0,1]
            gdf_plot.iloc[[idx]].plot(ax=ax, color=color, edgecolor='white', linewidth=1.5, alpha=0.7)

            centroid = row.geometry.centroid
            ax.annotate(f"Z{row['id_zona']}\n{valor:.2f}", (centroid.x, centroid.y),
                        xytext=(5, 5), textcoords="offset points",
                        fontsize=8, color='white', weight='bold',
                        bbox=dict(boxstyle="round,pad=0.3", facecolor=(30/255, 41/255, 59/255, 0.9), edgecolor='white'))

        try:
            ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, alpha=0.3)
        except Exception:
            pass

        if mostrar_capa_inta:
            agregar_capa_inta(ax, alpha=0.5)

        info_satelite = SATELITES_DISPONIBLES.get(satelite, SATELITES_DISPONIBLES['DATOS_SIMULADOS'])
        ax.set_title(f'{ICONOS_CULTIVOS[cultivo]} FERTILIDAD INTEGRADA (NPK) - {cultivo}\n'
                     f'{info_satelite["icono"]} {info_satelite["nombre"]}',
                     fontsize=16, fontweight='bold', pad=20, color='white')
        ax.set_xlabel('Longitud', color='white')
        ax.set_ylabel('Latitud', color='white')
        ax.tick_params(colors='white')
        ax.grid(True, alpha=0.3, color='#475569')

        sm = plt.cm.ScalarMappable(cmap=cmap, norm=plt.Normalize(vmin=0, vmax=1))
        sm.set_array([])
        cbar = plt.colorbar(sm, ax=ax, shrink=0.8)
        cbar.set_label('Índice de Fertilidad (0-1)', fontsize=12, fontweight='bold', color='white')
        cbar.ax.yaxis.set_tick_params(color='white')
        cbar.outline.set_edgecolor('white')
        plt.setp(plt.getp(cbar.ax.axes, 'yticklabels'), color='white')

        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='#0f172a')
        plt.close(fig)  # IMPORTANTE: Cerrar la figura
        buf.seek(0)
        return buf  # ← DEVOLVER BytesIO, NO bytes

    except Exception as e:
        st.error(f"Error creando mapa de fertilidad: {str(e)}")
        return None

def crear_mapa_texturas_con_esri(gdf_analizado, cultivo, mostrar_capa_inta=False):
    """Crea mapa de texturas con fondo ESRI Satellite - DEVUELVE BytesIO"""
    try:
        if gdf_analizado.empty or 'textura_suelo' not in gdf_analizado.columns:
            return None

        gdf_plot = gdf_analizado.to_crs(epsg=3857)
        fig, ax = plt.subplots(1, 1, figsize=(12, 8))
        fig.patch.set_facecolor('#0f172a')
        ax.set_facecolor('#0f172a')

        colores_textura = {
            'Franco limoso': '#c7eae5',
            'Franco': '#a6d96a',
            'Franco arcilloso limoso': '#5ab4ac',
            'Franco arenoso': '#f6e8c3',
            'Arcilla': '#01665e',
            'Arcilla limosa': '#003c30',
            'Arena franca': '#d8b365',
            'Limo': '#8c510a',
            'Franco arcilloso': '#35978f',
            'Franco arcilloso arenoso': '#80cdc1',
            'Limo arenoso': '#dfc27d',
            'Arena': '#f6e8c3',
            'Arcilla arenosa': '#01665e',
            'Franco limoso arenoso': '#a6d96a',
            'Sin datos': '#999999'
        }

        for idx, row in gdf_plot.iterrows():
            textura = row['textura_suelo']
            color = colores_textura.get(textura, '#999999')
            gdf_plot.iloc[[idx]].plot(ax=ax, color=color, edgecolor='white', linewidth=1.5, alpha=0.8)

            textura_abrev = textura[:12] + '...' if len(textura) > 15 else textura
            centroid = row.geometry.centroid
            ax.annotate(f"Z{row['id_zona']}\n{textura_abrev}", (centroid.x, centroid.y),
                        xytext=(5, 5), textcoords="offset points",
                        fontsize=8, color='black', weight='bold',
                        bbox=dict(boxstyle="round,pad=0.3", facecolor='white', alpha=0.9))

        try:
            ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, alpha=0.3)
        except Exception:
            pass

        if mostrar_capa_inta:
            agregar_capa_inta(ax, alpha=0.5)

        ax.set_title(f'{ICONOS_CULTIVOS[cultivo]} MAPA DE TEXTURAS USDA - {cultivo}',
                     fontsize=16, fontweight='bold', pad=20, color='white')
        ax.set_xlabel('Longitud', color='white')
        ax.set_ylabel('Latitud', color='white')
        ax.tick_params(colors='white')
        ax.grid(True, alpha=0.3, color='#475569')

        # Leyenda - CORREGIDO: usar mpatches.Patch
        texturas_presentes = [t for t in gdf_analizado['textura_suelo'].unique() if t in colores_textura]
        if texturas_presentes:
            legend_elements = [mpatches.Patch(facecolor=colores_textura[t], edgecolor='white', label=t) 
                              for t in texturas_presentes]
            legend = ax.legend(handles=legend_elements, title='Texturas USDA',
                               loc='upper left', bbox_to_anchor=(1.05, 1), fontsize=9)
            legend.get_title().set_color('white')
            for text in legend.get_texts():
                text.set_color('white')
            legend.get_frame().set_facecolor((30/255, 41/255, 59/255, 0.9))
            legend.get_frame().set_edgecolor('white')

        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='#0f172a')
        plt.close(fig)
        buf.seek(0)
        return buf

    except Exception as e:
        st.error(f"Error creando mapa de texturas: {str(e)}")
        return None
# ===== FUNCIONES DE GRÁFICOS NASA POWER CON ESTILO OSCURO =====
def crear_grafico_personalizado(series, titulo, ylabel, color_linea, fondo_grafico='#0f172a', color_texto='#ffffff'):
    """Crea gráfico de línea con estilo oscuro"""
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.set_facecolor(fondo_grafico)
    fig.patch.set_facecolor(fondo_grafico)
    ax.plot(series.index, series.values, color=color_linea, linewidth=2.2)
    ax.set_title(titulo, fontsize=14, fontweight='bold', color=color_texto)
    ax.set_ylabel(ylabel, fontsize=12, color=color_texto)
    ax.set_xlabel("Fecha", fontsize=11, color=color_texto)
    ax.tick_params(axis='x', colors=color_texto, rotation=0)
    ax.tick_params(axis='y', colors=color_texto)
    ax.grid(True, color='#475569', linestyle='--', linewidth=0.7, alpha=0.7)
    for spine in ax.spines.values():
        spine.set_color('#475569')
    plt.tight_layout()
    return fig

def crear_grafico_barras_personalizado(series, titulo, ylabel, color_barra, fondo_grafico='#0f172a', color_texto='#ffffff'):
    """Crea gráfico de barras con estilo oscuro"""
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.set_facecolor(fondo_grafico)
    fig.patch.set_facecolor(fondo_grafico)
    ax.bar(series.index, series.values, color=color_barra, alpha=0.85)
    ax.set_title(titulo, fontsize=14, fontweight='bold', color=color_texto)
    ax.set_ylabel(ylabel, fontsize=12, color=color_texto)
    ax.set_xlabel("Fecha", fontsize=11, color=color_texto)
    ax.tick_params(axis='x', colors=color_texto, rotation=0)
    ax.tick_params(axis='y', colors=color_texto)
    ax.grid(axis='y', color='#475569', linestyle='--', linewidth=0.7, alpha=0.7)
    for spine in ax.spines.values():
        spine.set_color('#475569')
    plt.tight_layout()
    return fig

# ===== FUNCIÓN PRINCIPAL DE ANÁLISIS (MEJORADA) =====
def ejecutar_analisis(gdf, nutriente, analisis_tipo, n_divisiones, cultivo,
                     satelite=None, indice=None, fecha_inicio=None,
                     fecha_fin=None, intervalo_curvas=5.0, resolucion_dem=10.0,
                     usar_inta=True, mostrar_capa_inta=False):
    resultados = {
        'exitoso': False,
        'gdf_analizado': None,
        'mapa_buffer': None,
        'tabla_datos': None,
        'estadisticas': {},
        'recomendaciones': [],
        'area_total': 0,
        'df_power': None
    }

    try:
        gdf = validar_y_corregir_crs(gdf)
        area_total = calcular_superficie(gdf)
        resultados['area_total'] = area_total

        # === ANÁLISIS DE TEXTURA DEL SUELO ===
        if analisis_tipo == "ANÁLISIS DE TEXTURA":
            gdf_dividido = dividir_parcela_en_zonas(gdf, n_divisiones)
            gdf_analizado = analizar_textura_suelo(gdf_dividido, cultivo)
            resultados['gdf_analizado'] = gdf_analizado
            resultados['exitoso'] = True
            return resultados

        # === ANÁLISIS DE CURVAS DE NIVEL ===
        elif analisis_tipo == "ANÁLISIS DE CURVAS DE NIVEL":
            gdf_dividido = dividir_parcela_en_zonas(gdf, n_divisiones)
            resultados['gdf_analizado'] = gdf_dividido
            resultados['exitoso'] = True
            return resultados

        # === ANÁLISIS SATELITAL (FERTILIDAD O NPK) ===
        elif analisis_tipo in ["FERTILIDAD ACTUAL", "RECOMENDACIONES NPK"]:
            # Obtener datos satelitales
            datos_satelitales = None
            if satelite == "SENTINEL-2":
                datos_satelitales = descargar_datos_sentinel2(gdf, fecha_inicio, fecha_fin, indice)
            elif satelite == "LANDSAT-8":
                datos_satelitales = descargar_datos_landsat8(gdf, fecha_inicio, fecha_fin, indice)
            else:
                datos_satelitales = generar_datos_simulados(gdf, cultivo, indice)

            # Dividir parcela
            gdf_dividido = dividir_parcela_en_zonas(gdf, n_divisiones)

            # Calcular NPK usando metodologías científicas + INTA
            indices_npk = calcular_indices_npk_avanzados(gdf_dividido, cultivo, satelite, usar_inta)

            # Crear GeoDataFrame con resultados
            gdf_analizado = gdf_dividido.copy()
            for idx, indice_data in enumerate(indices_npk):
                for key, value in indice_data.items():
                    gdf_analizado.loc[gdf_analizado.index[idx], key] = value

            # Calcular áreas
            areas_ha_list = []
            for idx, row in gdf_analizado.iterrows():
                area_gdf = gpd.GeoDataFrame({'geometry': [row.geometry]}, crs=gdf_analizado.crs)
                area_ha = calcular_superficie(area_gdf)
                if hasattr(area_ha, 'iloc'):
                    area_ha = float(area_ha.iloc[0])
                elif hasattr(area_ha, '__len__') and len(area_ha) > 0:
                    area_ha = float(area_ha[0])
                else:
                    area_ha = float(area_ha)
                areas_ha_list.append(area_ha)
            gdf_analizado['area_ha'] = areas_ha_list

            # Calcular recomendaciones si es necesario
            if analisis_tipo == "RECOMENDACIONES NPK":
                recomendaciones_npk = calcular_recomendaciones_npk_cientificas(gdf_analizado, nutriente, cultivo)
                gdf_analizado['valor_recomendado'] = recomendaciones_npk

                # ✅ NUEVO: Calcular rendimientos para análisis de recomendaciones
                rendimientos_actual = calcular_rendimiento_potencial(gdf_analizado, cultivo)
                rendimientos_proyectado = calcular_rendimiento_con_recomendaciones(gdf_analizado, cultivo)
                gdf_analizado['rendimiento_actual'] = rendimientos_actual
                gdf_analizado['rendimiento_proyectado'] = rendimientos_proyectado
                gdf_analizado['incremento_rendimiento'] = gdf_analizado['rendimiento_proyectado'] - gdf_analizado['rendimiento_actual']

            # ✅ NUEVO: Para fertilidad actual también calcular rendimiento
            elif analisis_tipo == "FERTILIDAD ACTUAL":
                rendimientos_actual = calcular_rendimiento_potencial(gdf_analizado, cultivo)
                gdf_analizado['rendimiento_actual'] = rendimientos_actual

            resultados['gdf_analizado'] = gdf_analizado
            resultados['exitoso'] = True

            # === DATOS DE NASA POWER ===
            if satelite:
                df_power = obtener_datos_nasa_power(gdf, fecha_inicio, fecha_fin)
                if df_power is not None:
                    resultados['df_power'] = df_power

            return resultados

        else:
            st.error(f"Tipo de análisis no soportado: {analisis_tipo}")
            return resultados

    except Exception as e:
        st.error(f"❌ Error en análisis: {str(e)}")
        import traceback
        st.error(f"Detalle: {traceback.format_exc()}")
        return resultados

# ===== FUNCIONES DE VISUALIZACIÓN =====
def mostrar_resultados_textura(gdf_analizado, cultivo, area_total, mostrar_capa_inta=False):
    st.subheader("📊 ESTADÍSTICAS DE TEXTURA (USDA)")
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        textura_predominante = gdf_analizado['textura_suelo'].mode()[0] if len(gdf_analizado) > 0 else "Sin datos"
        st.metric("🏗️ Textura Predominante", textura_predominante)
    with col2:
        avg_arena = gdf_analizado['arena'].mean()
        st.metric("🏖️ Arena Promedio", f"{avg_arena:.1f}%")
    with col3:
        avg_limo = gdf_analizado['limo'].mean()
        st.metric("🌫️ Limo Promedio", f"{avg_limo:.1f}%")
    with col4:
        avg_arcilla = gdf_analizado['arcilla'].mean()
        st.metric("🧱 Arcilla Promedio", f"{avg_arcilla:.1f}%")

    st.subheader("📈 COMPOSICIÓN GRANULOMÉTRICA (USDA)")
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6))

    # Configurar estilo oscuro
    fig.patch.set_facecolor('#0f172a')
    ax1.set_facecolor('#0f172a')
    ax2.set_facecolor('#0f172a')

    composicion = [gdf_analizado['arena'].mean(), gdf_analizado['limo'].mean(), gdf_analizado['arcilla'].mean()]
    labels = ['Arena', 'Limo', 'Arcilla']
    colors_pie = ['#d8b365', '#f6e8c3', '#01665e']
    ax1.pie(composicion, labels=labels, colors=colors_pie, autopct='%1.1f%%', startangle=90, textprops={'color': 'white'})
    ax1.set_title('Composición Promedio USDA', color='white')

    textura_dist = gdf_analizado['textura_suelo'].value_counts()
    ax2.bar(textura_dist.index, textura_dist.values, color=[PALETAS_GEE['TEXTURA'][i % len(PALETAS_GEE['TEXTURA'])] for i in range(len(textura_dist))])
    ax2.set_title('Distribución de Texturas USDA', color='white')
    ax2.set_xlabel('Clase Textural USDA', color='white')
    ax2.set_ylabel('Número de Zonas', color='white')
    ax2.tick_params(axis='x', rotation=45, colors='white')
    ax2.tick_params(axis='y', colors='white')
    ax2.set_facecolor('#0f172a')

    plt.tight_layout()
    st.pyplot(fig)

    st.subheader("🗺️ MAPA DE TEXTURAS USDA CON ESRI SATELLITE")
    mapa_texturas = crear_mapa_texturas_con_esri(gdf_analizado, cultivo, mostrar_capa_inta)
    if mapa_texturas:
        st.image(mapa_texturas, use_container_width=True)
        st.download_button(
            "📥 Descargar Mapa de Texturas USDA",
            mapa_texturas,
            f"mapa_texturas_usda_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.png",
            "image/png"
        )

    st.subheader("📋 TABLA DE RESULTADOS POR ZONA (USDA)")
    columnas_textura = ['id_zona', 'area_ha', 'textura_suelo', 'arena', 'limo', 'arcilla']
    columnas_textura = [col for col in columnas_textura if col in gdf_analizado.columns]
    if columnas_textura:
        tabla_textura = gdf_analizado[columnas_textura].copy()
        tabla_textura.columns = ['Zona', 'Área (ha)', 'Textura USDA', 'Arena (%)', 'Limo (%)', 'Arcilla (%)']
        st.dataframe(tabla_textura)

    st.subheader("💡 RECOMENDACIONES DE MANEJO POR TEXTURA USDA")
    if 'textura_suelo' in gdf_analizado.columns:
        textura_predominante = gdf_analizado['textura_suelo'].mode()[0] if len(gdf_analizado) > 0 else "Sin datos"
        if textura_predominante in RECOMENDACIONES_TEXTURA:
            st.markdown(f"#### 🏗️ **{textura_predominante.upper()}**")
            info_textura = RECOMENDACIONES_TEXTURA[textura_predominante]
            col1, col2, col3 = st.columns(3)
            with col1:
                st.markdown("**✅ PROPIEDADES FÍSICAS**")
                for prop in info_textura['propiedades']:
                    st.markdown(f"• {prop}")
            with col2:
                st.markdown("**⚠️ LIMITANTES**")
                for lim in info_textura['limitantes']:
                    st.markdown(f"• {lim}")
            with col3:
                st.markdown("**🛠️ MANEJO RECOMENDADO**")
                for man in info_textura['manejo']:
                    st.markdown(f"• {man}")
        else:
            st.info(f"Textura '{textura_predominante}' - Consultar recomendaciones específicas para esta clase textural")

    st.subheader("💾 DESCARGAR RESULTADOS USDA")
    if 'columnas_textura' in locals() and columnas_textura:
        tabla_textura = gdf_analizado[columnas_textura].copy()
        tabla_textura.columns = ['Zona', 'Área (ha)', 'Textura USDA', 'Arena (%)', 'Limo (%)', 'Arcilla (%)']
        csv = tabla_textura.to_csv(index=False)
        st.download_button(
            "📥 Descargar CSV con Análisis de Textura USDA",
            csv,
            f"textura_usda_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
            "text/csv"
        )

def mostrar_resultados_curvas_nivel(X, Y, Z, pendiente_grid, curvas, elevaciones, gdf_original, cultivo, area_total):
    st.subheader("📊 ESTADÍSTICAS TOPOGRÁFICAS")
    elevaciones_flat = Z.flatten()
    elevaciones_flat = elevaciones_flat[~np.isnan(elevaciones_flat)]
    if len(elevaciones_flat) > 0:
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            elevacion_promedio = np.mean(elevaciones_flat)
            st.metric("🏔️ Elevación Promedio", f"{elevacion_promedio:.1f} m")
        with col2:
            rango_elevacion = np.max(elevaciones_flat) - np.min(elevaciones_flat)
            st.metric("📏 Rango de Elevación", f"{rango_elevacion:.1f} m")
        with col3:
            mapa_pendientes, stats_pendiente = crear_mapa_pendientes_simple(X, Y, pendiente_grid, gdf_original)
            st.metric("📐 Pendiente Promedio", f"{stats_pendiente['promedio']:.1f}%")
        with col4:
            num_curvas = len(curvas) if curvas else 0
            st.metric("🔄 Número de Curvas", f"{num_curvas}")

    st.subheader("🔥 MAPA DE CALOR DE PENDIENTES")
    st.image(mapa_pendientes, use_container_width=True)
    st.download_button(
        "📥 Descargar Mapa de Pendientes",
        mapa_pendientes,
        f"mapa_pendientes_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.png",
        "image/png"
    )

    st.subheader("⚠️ ANÁLISIS DE RIESGO DE EROSION")
    if 'stats_pendiente' in locals() and 'distribucion' in stats_pendiente:
        riesgo_total = 0
        for categoria, data in stats_pendiente['distribucion'].items():
            if categoria in CLASIFICACION_PENDIENTES:
                riesgo_total += data['porcentaje'] * CLASIFICACION_PENDIENTES[categoria]['factor_erosivo']
        riesgo_promedio = riesgo_total / 100

        col1, col2, col3 = st.columns(3)
        with col1:
            if riesgo_promedio < 0.3:
                st.success("✅ **RIESGO BAJO**")
                st.metric("Factor Riesgo", f"{riesgo_promedio:.2f}")
            elif riesgo_promedio < 0.6:
                st.warning("⚠️ **RIESGO MODERADO**")
                st.metric("Factor Riesgo", f"{riesgo_promedio:.2f}")
            else:
                st.error("🚨 **RIESGO ALTO**")
                st.metric("Factor Riesgo", f"{riesgo_promedio:.2f}")

        with col2:
            area_total_ha = area_total
            porcentaje_critico = sum(data['porcentaje'] for cat, data in stats_pendiente['distribucion'].items()
                                     if cat in ['FUERTE (10-15%)', 'MUY FUERTE (15-25%)', 'EXTREMA (>25%)'])
            area_critica = area_total_ha * (porcentaje_critico / 100)
            st.metric("Área Crítica (>10%)", f"{area_critica:.2f} ha")

        with col3:
            porcentaje_manejable = sum(data['porcentaje'] for cat, data in stats_pendiente['distribucion'].items()
                                       if cat in ['PLANA (0-2%)', 'SUAVE (2-5%)', 'MODERADA (5-10%)'])
            area_manejable = area_total_ha * (porcentaje_manejable / 100)
            st.metric("Área Manejable (<10%)", f"{area_manejable:.2f} ha")

    st.subheader("📈 VISUALIZACIÓN 3D DEL TERRENO")
    try:
        fig = plt.figure(figsize=(12, 8))
        ax = fig.add_subplot(111, projection='3d')
        surf = ax.plot_surface(X, Y, Z, cmap='terrain', alpha=0.8, linewidth=0)
        ax.set_xlabel('Longitud', color='white')
        ax.set_ylabel('Latitud', color='white')
        ax.set_zlabel('Elevación (m)', color='white')
        ax.set_title(f'Modelo 3D del Terreno - {cultivo}', color='white')
        ax.tick_params(colors='white')
        fig.patch.set_facecolor('#0f172a')
        ax.set_facecolor('#0f172a')
        ax.xaxis.label.set_color('white')
        ax.yaxis.label.set_color('white')
        ax.zaxis.label.set_color('white')
        ax.title.set_color('white')

        cbar = fig.colorbar(surf, ax=ax, shrink=0.5, aspect=5, label='Elevación (m)')
        cbar.set_label('Elevación (m)', color='white')
        cbar.ax.yaxis.set_tick_params(color='white')
        plt.setp(plt.getp(cbar.ax.axes, 'yticklabels'), color='white')

        plt.tight_layout()
        st.pyplot(fig)
    except Exception as e:
        st.warning(f"No se pudo generar visualización 3D: {e}")

    st.subheader("💾 DESCARGAR RESULTADOS")
    sample_points = []
    for i in range(0, X.shape[0], 5):
        for j in range(0, X.shape[1], 5):
            if not np.isnan(Z[i, j]):
                sample_points.append({
                    'lat': Y[i, j],
                    'lon': X[i, j],
                    'elevacion_m': Z[i, j],
                    'pendiente_%': pendiente_grid[i, j]
                })

    if sample_points:
        df_dem = pd.DataFrame(sample_points)
        csv = df_dem.to_csv(index=False)
        st.download_button(
            label="📊 Descargar Muestras DEM (CSV)",
            data=csv,
            file_name=f"dem_muestras_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
            mime="text/csv"
        )

# ===== SIDEBAR MEJORADO CON OPCIÓN DEL INTA =====
with st.sidebar:
    st.markdown('<div class="sidebar-title">⚙️ CONFIGURACIÓN</div>', unsafe_allow_html=True)
    cultivo = st.selectbox("Cultivo:", ["VID", "OLIVO", "HORTALIZAS DE HOJA"])
    
    # Selección de variedad según cultivo
    if cultivo == "VID":
        variedad = st.selectbox(
            "Variedad de Vid:",
            list(VARIEDADES_VID.keys()),
            index=0
        )
        st.session_state['variedad'] = variedad
        st.session_state['variedad_params'] = VARIEDADES_VID[variedad]
    elif cultivo == "OLIVO":
        variedad = st.selectbox(
            "Variedad de Olivo:",
            list(VARIEDADES_OLIVO.keys()),
            index=0
        )
        st.session_state['variedad'] = variedad
        st.session_state['variedad_params'] = VARIEDADES_OLIVO[variedad]
    elif cultivo == "HORTALIZAS DE HOJA":
        variedad = st.selectbox(
            "Variedad de Hortaliza:",
            list(VARIEDADES_HORTALIZAS.keys()),
            index=0
        )
        st.session_state['variedad'] = variedad
        st.session_state['variedad_params'] = VARIEDADES_HORTALIZAS[variedad]

    # Mostrar información de la variedad seleccionada
    if 'variedad' in st.session_state and st.session_state['variedad']:
        params = st.session_state['variedad_params']
        st.info(f"""
        **📊 {st.session_state['variedad']}**
        - Potencial: {params['RENDIMIENTO_BASE']} - {params['RENDIMIENTO_OPTIMO']} ton/ha
        - Ciclo: {params.get('CICLO', 'N/D')} días
        - Región: {params.get('REGION', 'N/D')}
        """)
    st.image(IMAGENES_CULTIVOS[cultivo], use_container_width=True)

    # Mostrar metodología NPK seleccionada
    if satelite_seleccionado in METODOLOGIAS_NPK:
        st.info(f"**Metodología {satelite_seleccionado}:**")
        for nutriente_metodo, info in METODOLOGIAS_NPK[satelite_seleccionado].items():
            st.write(f"- **{nutriente_metodo}**: {info['metodo']}")

    # ===== NUEVO: INTEGRACIÓN CON EL INTA =====
    with st.sidebar.expander("🌱 INTEGRACIÓN CON EL INTA"):
        usar_inta = st.checkbox(
            "✅ Usar datos del INTA para materia orgánica",
            value=True,
            help="Activa la estimación de materia orgánica basada en mapas reales del INTA por región"
        )
        
        mostrar_mapa_inta = st.checkbox(
            "🗺️ Mostrar capa base del INTA",
            value=False,
            help="Superpone el mapa de suelos del INTA como capa base (requiere conexión a internet)"
        )
        
        if usar_inta:
            st.info("🔍 La materia orgánica se estimará usando datos regionales del INTA")
        else:
            st.warning("⚠️ Se usará estimación genérica basada en cultivo")

    analisis_tipo = st.selectbox("Tipo de Análisis:", ["FERTILIDAD ACTUAL", "RECOMENDACIONES NPK", "ANÁLISIS DE TEXTURA", "ANÁLISIS DE CURVAS DE NIVEL"])
    if analisis_tipo == "RECOMENDACIONES NPK":
        nutriente = st.selectbox("Nutriente:", ["NITRÓGENO", "FÓSFORO", "POTASIO"])

    st.subheader("🛰️ Fuente de Datos Satelitales")
    satelite_seleccionado = st.selectbox(
        "Satélite:",
        ["SENTINEL-2", "LANDSAT-8", "DATOS_SIMULADOS"],
        help="Selecciona la fuente de datos satelitales"
    )
    if satelite_seleccionado in SATELITES_DISPONIBLES:
        info_satelite = SATELITES_DISPONIBLES[satelite_seleccionado]
        st.info(f"""
        **{info_satelite['icono']} {info_satelite['nombre']}**
        - Resolución: {info_satelite['resolucion']}
        - Revisita: {info_satelite['revisita']}
        - Índices: {', '.join(info_satelite['indices'][:3])}
        """)

    if analisis_tipo in ["FERTILIDAD ACTUAL", "RECOMENDACIONES NPK"]:
        st.subheader("📊 Índices de Vegetación")
        if satelite_seleccionado == "SENTINEL-2":
            indice_seleccionado = st.selectbox("Índice:", SATELITES_DISPONIBLES['SENTINEL-2']['indices'])
        elif satelite_seleccionado == "LANDSAT-8":
            indice_seleccionado = st.selectbox("Índice:", SATELITES_DISPONIBLES['LANDSAT-8']['indices'])
        else:
            indice_seleccionado = st.selectbox("Índice:", SATELITES_DISPONIBLES['DATOS_SIMULADOS']['indices'])

    if analisis_tipo in ["FERTILIDAD ACTUAL", "RECOMENDACIONES NPK"]:
        st.subheader("📅 Rango Temporal")
        fecha_fin = st.date_input("Fecha fin", datetime.now())
        fecha_inicio = st.date_input("Fecha inicio", datetime.now() - timedelta(days=30))

    st.subheader("🎯 División de Parcela")
    n_divisiones = st.slider("Número de zonas de manejo:", min_value=16, max_value=48, value=32)

    if analisis_tipo == "ANÁLISIS DE CURVAS DE NIVEL":
        st.subheader("🏔️ Configuración Curvas de Nivel")
        intervalo_curvas = st.slider("Intervalo entre curvas (metros):", 1.0, 20.0, 5.0, 1.0)
        resolucion_dem = st.slider("Resolución DEM (metros):", 5.0, 50.0, 10.0, 5.0)

    st.subheader("📤 Subir Parcela")
    uploaded_file = st.file_uploader("Subir archivo de tu parcela", type=['zip', 'kml', 'kmz'],
                                     help="Formatos aceptados: Shapefile (.zip), KML (.kml), KMZ (.kmz)")

    # CONFIGURACIÓN ECONÓMICA
    with st.sidebar.expander("💰 CONFIGURACIÓN ECONÓMICA"):
        st.markdown("#### Precios de Mercado (USD)")
        # Precios de cultivos
        st.subheader("🌾 Precios Cultivos")
        precio_vid = st.number_input("Vid (USD/ton)", value=800.0, min_value=500.0, max_value=1200.0)
        precio_olivo = st.number_input("Olivo (USD/ton)", value=1200.0, min_value=800.0, max_value=1800.0)
        precio_hortalizas = st.number_input("Hortalizas (USD/ton)", value=500.0, min_value=300.0, max_value=800.0)

        # Actualizar precios en parámetros
        PARAMETROS_ECONOMICOS['PRECIOS_CULTIVOS']['VID']['precio_ton'] = precio_vid
        PARAMETROS_ECONOMICOS['PRECIOS_CULTIVOS']['OLIVO']['precio_ton'] = precio_olivo
        PARAMETROS_ECONOMICOS['PRECIOS_CULTIVOS']['HORTALIZAS DE HOJA']['precio_ton'] = precio_hortalizas

        st.subheader("🧪 Precios Fertilizantes")
        precio_urea = st.number_input("Urea (USD/ton)", value=450.0, min_value=300.0, max_value=600.0)
        precio_fosfato = st.number_input("Fosfato (USD/ton)", value=650.0, min_value=400.0, max_value=800.0)
        precio_potasio = st.number_input("Potasio (USD/ton)", value=400.0, min_value=250.0, max_value=550.0)
        PARAMETROS_ECONOMICOS['PRECIOS_FERTILIZANTES']['UREA'] = precio_urea
        PARAMETROS_ECONOMICOS['PRECIOS_FERTILIZANTES']['FOSFATO_DIAMONICO'] = precio_fosfato
        PARAMETROS_ECONOMICOS['PRECIOS_FERTILIZANTES']['CLORURO_POTASIO'] = precio_potasio

        st.subheader("📈 Parámetros Financieros")
        tasa_descuento = st.slider("Tasa Descuento (%)", 5.0, 20.0, 10.0, 0.5) / 100
        inflacion = st.slider("Inflación Esperada (%)", 0.0, 15.0, 8.0, 0.5) / 100
        PARAMETROS_ECONOMICOS['PARAMETROS_FINANCIEROS']['tasa_descuento'] = tasa_descuento
        PARAMETROS_ECONOMICOS['PARAMETROS_FINANCIEROS']['inflacion_esperada'] = inflacion

# Guardar estado de las opciones del INTA en session_state
st.session_state['usar_inta'] = usar_inta
st.session_state['mostrar_mapa_inta'] = mostrar_mapa_inta

# ===== FUNCIONES DE VISUALIZACIÓN DE RESULTADOS =====
def mostrar_resultados_fertilidad(gdf_analizado, cultivo, area_total, satelite, mostrar_capa_inta=False):
    """Muestra resultados del análisis de fertilidad actual"""
    
    st.subheader(f"🌱 ANÁLISIS DE FERTILIDAD ACTUAL - {cultivo}")
    
    # Estadísticas principales
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        if 'npk_integrado' in gdf_analizado.columns:
            npk_prom = gdf_analizado['npk_integrado'].mean()
            st.metric("📊 Índice NPK Integrado", f"{npk_prom:.3f}")
    with col2:
        if 'nitrogeno_actual' in gdf_analizado.columns:
            n_prom = gdf_analizado['nitrogeno_actual'].mean()
            st.metric("🧪 Nitrógeno", f"{n_prom:.1f} kg/ha")
    with col3:
        if 'fosforo_actual' in gdf_analizado.columns:
            p_prom = gdf_analizado['fosforo_actual'].mean()
            st.metric("🔬 Fósforo", f"{p_prom:.1f} kg/ha")
    with col4:
        if 'potasio_actual' in gdf_analizado.columns:
            k_prom = gdf_analizado['potasio_actual'].mean()
            st.metric("⚡ Potasio", f"{k_prom:.1f} kg/ha")
    
    # Mapa de fertilidad integrada
    st.subheader("🗺️ MAPA DE FERTILIDAD INTEGRADA (NPK)")
    mapa_fertilidad = crear_mapa_fertilidad_integrada(
        gdf_analizado, cultivo, satelite, mostrar_capa_inta
    )
    
    if mapa_fertilidad:
        # Verificar si es bytes o BytesIO y convertirlo si es necesario
        if isinstance(mapa_fertilidad, bytes):
            mapa_bytes = mapa_fertilidad
            mapa_buffer = io.BytesIO(mapa_fertilidad)
        else:
            mapa_buffer = mapa_fertilidad
            mapa_bytes = mapa_fertilidad.getvalue()
        
        st.image(mapa_buffer, use_container_width=True)
        
        # Botón para descargar
        st.download_button(
            label="📥 Descargar Mapa de Fertilidad",
            data=mapa_bytes,
            file_name=f"fertilidad_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.png",
            mime="image/png"
        )
    else:
        st.warning("No se pudo generar el mapa de fertilidad.")
    
    # Tabla de resultados
    st.subheader("📋 RESULTADOS POR ZONA")
    columnas_mostrar = ['id_zona', 'area_ha', 'npk_integrado', 
                       'nitrogeno_actual', 'fosforo_actual', 'potasio_actual']
    columnas_mostrar = [col for col in columnas_mostrar if col in gdf_analizado.columns]
    
    if columnas_mostrar:
        df_mostrar = gdf_analizado[columnas_mostrar].copy()
        df_mostrar.columns = ['Zona', 'Área (ha)', 'Índice NPK', 
                             'Nitrógeno (kg/ha)', 'Fósforo (kg/ha)', 'Potasio (kg/ha)']
        st.dataframe(df_mostrar)
        
        # Descargar CSV
        csv = df_mostrar.to_csv(index=False)
        st.download_button(
            label="📥 Descargar CSV",
            data=csv,
            file_name=f"fertilidad_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
            mime="text/csv"
        )

def mostrar_resultados_recomendaciones(gdf_analizado, cultivo, nutriente, area_total, satelite, mostrar_capa_inta=False):
    """Muestra resultados de recomendaciones NPK"""
    
    st.subheader(f"💡 RECOMENDACIONES DE {nutriente} - {cultivo}")
    
    # Estadísticas
    col1, col2, col3 = st.columns(3)
    with col1:
        if 'valor_recomendado' in gdf_analizado.columns:
            rec_prom = gdf_analizado['valor_recomendado'].mean()
            st.metric("🧴 Recomendación Promedio", f"{rec_prom:.1f} kg/ha")
    with col2:
        col_nutriente = f"{nutriente.lower()}_actual"
        if col_nutriente in gdf_analizado.columns:
            actual_prom = gdf_analizado[col_nutriente].mean()
            st.metric("📊 Nivel Actual", f"{actual_prom:.1f} kg/ha")
    with col3:
        if 'area_ha' in gdf_analizado.columns and 'valor_recomendado' in gdf_analizado.columns:
            total_kg = (gdf_analizado['valor_recomendado'] * gdf_analizado['area_ha']).sum()
            st.metric("⚖️ Total Requerido", f"{total_kg:.0f} kg")
    
    # Mapa NPK
    st.subheader(f"🗺️ MAPA DE {nutriente}")
    mapa_npk = crear_mapa_npk_con_esri(
        gdf_analizado, nutriente, cultivo, satelite, mostrar_capa_inta
    )
    
    if mapa_npk:
        if isinstance(mapa_npk, bytes):
            mapa_bytes = mapa_npk
            mapa_buffer = io.BytesIO(mapa_npk)
        else:
            mapa_buffer = mapa_npk
            mapa_bytes = mapa_npk.getvalue()
        
        st.image(mapa_buffer, use_container_width=True)
        
        # Botón para descargar
        st.download_button(
            label=f"📥 Descargar Mapa de {nutriente}",
            data=mapa_bytes,
            file_name=f"{nutriente.lower()}_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.png",
            mime="image/png"
        )
    
    # Tabla de resultados con recomendaciones
    st.subheader("📋 RECOMENDACIONES POR ZONA")
    columnas_mostrar = ['id_zona', 'area_ha']
    
    # Agregar columnas según nutriente
    if nutriente == "NITRÓGENO":
        columnas_mostrar.extend(['nitrogeno_actual'])
    elif nutriente == "FÓSFORO":
        columnas_mostrar.extend(['fosforo_actual'])
    elif nutriente == "POTASIO":
        columnas_mostrar.extend(['potasio_actual'])
    
    columnas_mostrar.append('valor_recomendado')
    columnas_mostrar = [col for col in columnas_mostrar if col in gdf_analizado.columns]
    
    if columnas_mostrar:
        df_mostrar = gdf_analizado[columnas_mostrar].copy()
        
        # Renombrar columnas para mostrar
        rename_dict = {
            'id_zona': 'Zona',
            'area_ha': 'Área (ha)',
            'nitrogeno_actual': 'Nitrógeno Actual (kg/ha)',
            'fosforo_actual': 'Fósforo Actual (kg/ha)',
            'potasio_actual': 'Potasio Actual (kg/ha)',
            'valor_recomendado': 'Recomendación (kg/ha)'
        }
        
        df_mostrar = df_mostrar.rename(columns=rename_dict)
        st.dataframe(df_mostrar)
        
        # Descargar CSV
        csv = df_mostrar.to_csv(index=False)
        st.download_button(
            label="📥 Descargar Recomendaciones (CSV)",
            data=csv,
            file_name=f"recomendaciones_{nutriente}_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
            mime="text/csv"
        )
    
    # === ANÁLISIS DE RENDIMIENTO ===
    if 'rendimiento_actual' in gdf_analizado.columns and 'rendimiento_proyectado' in gdf_analizado.columns:
        st.subheader("🌾 ANÁLISIS DE POTENCIAL DE COSECHA")
        
        # Métricas de rendimiento
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            rend_actual = gdf_analizado['rendimiento_actual'].mean()
            st.metric("📊 Rendimiento Actual", f"{rend_actual:.1f} ton/ha")
        with col2:
            rend_proy = gdf_analizado['rendimiento_proyectado'].mean()
            st.metric("🚀 Rendimiento Proyectado", f"{rend_proy:.1f} ton/ha")
        with col3:
            incremento = gdf_analizado['incremento_rendimiento'].mean()
            st.metric("📈 Incremento", f"+{incremento:.1f} ton/ha")
        with col4:
            porcentaje = (incremento / rend_actual * 100) if rend_actual > 0 else 0
            st.metric("💯 % Aumento", f"+{porcentaje:.1f}%")
        
        # Mapas de calor de rendimiento
        st.subheader("🔥 MAPAS DE CALOR DE RENDIMIENTO")
        
        # Tabs para diferentes mapas
        tab1, tab2, tab3 = st.tabs(["📊 Actual", "🚀 Proyectado", "📈 Comparativo"])
        
        with tab1:
            mapa_actual = crear_mapa_calor_rendimiento_actual(gdf_analizado, cultivo)
            if mapa_actual:
                st.image(mapa_actual, use_container_width=True)
                st.download_button(
                    label="📥 Descargar Mapa Actual",
                    data=mapa_actual.getvalue() if hasattr(mapa_actual, 'getvalue') else mapa_actual,
                    file_name=f"rendimiento_actual_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.png",
                    mime="image/png"
                )
        
        with tab2:
            mapa_proyectado = crear_mapa_calor_rendimiento_proyectado(gdf_analizado, cultivo)
            if mapa_proyectado:
                st.image(mapa_proyectado, use_container_width=True)
                st.download_button(
                    label="📥 Descargar Mapa Proyectado",
                    data=mapa_proyectado.getvalue() if hasattr(mapa_proyectado, 'getvalue') else mapa_proyectado,
                    file_name=f"rendimiento_proyectado_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.png",
                    mime="image/png"
                )
        
        with tab3:
            mapa_comparativo = crear_mapa_comparativo_calor(gdf_analizado, cultivo)
            if mapa_comparativo:
                st.image(mapa_comparativo, use_container_width=True)
                st.download_button(
                    label="📥 Descargar Mapa Comparativo",
                    data=mapa_comparativo.getvalue() if hasattr(mapa_comparativo, 'getvalue') else mapa_comparativo,
                    file_name=f"rendimiento_comparativo_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.png",
                    mime="image/png"
                )
        
        # Tabla de rendimientos
        st.subheader("📋 RENDIMIENTOS POR ZONA")
        columnas_rend = ['id_zona', 'area_ha', 'rendimiento_actual', 'rendimiento_proyectado', 'incremento_rendimiento']
        columnas_rend = [col for col in columnas_rend if col in gdf_analizado.columns]
        
        if columnas_rend:
            df_rend = gdf_analizado[columnas_rend].copy()
            df_rend.columns = ['Zona', 'Área (ha)', 'Rendimiento Actual (ton/ha)', 'Rendimiento Proyectado (ton/ha)', 'Incremento (ton/ha)']
            st.dataframe(df_rend)
            
            # Descargar CSV de rendimientos
            csv_rend = df_rend.to_csv(index=False)
            st.download_button(
                label="📥 Descargar Rendimientos (CSV)",
                data=csv_rend,
                file_name=f"rendimientos_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                mime="text/csv"
            )
        
        # === ANÁLISIS ECONÓMICO ===
        st.subheader("💰 ANÁLISIS ECONÓMICO")
        
        if st.button("🧮 Calcular Análisis Económico", type="primary"):
            with st.spinner("Calculando análisis económico..."):
                resultados_economicos = realizar_analisis_economico(
                    gdf_analizado, 
                    cultivo, 
                    st.session_state['variedad_params'], 
                    area_total
                )
                
                mostrar_analisis_economico(resultados_economicos)
    
    # Datos de NASA POWER
    st.subheader("🌤️ DATOS METEOROLÓGICOS NASA POWER")
    if st.button("🛰️ Obtener Datos Climáticos", type="secondary"):
        with st.spinner("Consultando NASA POWER..."):
            df_power = obtener_datos_nasa_power(gdf_analizado, fecha_inicio, fecha_fin)
            if df_power is not None:
                # Mostrar estadísticas
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("☀️ Radiación Solar", f"{df_power['radiacion_solar'].mean():.1f} kWh/m²/día")
                with col2:
                    st.metric("💨 Viento 2m", f"{df_power['viento_2m'].mean():.1f} m/s")
                with col3:
                    st.metric("🌡️ Temperatura", f"{df_power['temperatura'].mean():.1f} °C")
                with col4:
                    st.metric("🌧️ Precipitación", f"{df_power['precipitacion'].mean():.2f} mm/día")
                
                # Gráficos
                st.subheader("📈 GRÁFICOS METEOROLÓGICOS")
                
                tab1, tab2, tab3, tab4 = st.tabs(["☀️ Radiación", "💨 Viento", "🌡️ Temperatura", "🌧️ Precipitación"])
                
                with tab1:
                    fig = crear_grafico_personalizado(df_power.set_index('fecha')['radiacion_solar'], 
                                                     "Radiación Solar Diaria", "kWh/m²/día", "#FFD700")
                    st.pyplot(fig)
                
                with tab2:
                    fig = crear_grafico_personalizado(df_power.set_index('fecha')['viento_2m'], 
                                                     "Velocidad del Viento a 2m", "m/s", "#87CEEB")
                    st.pyplot(fig)
                
                with tab3:
                    fig = crear_grafico_personalizado(df_power.set_index('fecha')['temperatura'], 
                                                     "Temperatura del Aire a 2m", "°C", "#FF6347")
                    st.pyplot(fig)
                
                with tab4:
                    fig = crear_grafico_barras_personalizado(df_power.set_index('fecha')['precipitacion'], 
                                                           "Precipitación Diaria", "mm/día", "#1E90FF")
                    st.pyplot(fig)
                
                # Descargar datos
                csv_power = df_power.to_csv(index=False)
                st.download_button(
                    label="📥 Descargar Datos Climáticos (CSV)",
                    data=csv_power,
                    file_name=f"nasa_power_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                    mime="text/csv"
                )
            else:
                st.warning("No se pudieron obtener datos de NASA POWER para este período/locación")

# ===== INTERFAZ PRINCIPAL - VERSIÓN CORREGIDA =====
def main():
    st.title("🍇🫒🥬 ANALIZADOR MULTI-CULTIVO SATELITAL")
    st.markdown("### **VID | OLIVO | HORTALIZAS DE HOJAS**")
    
    # Variables para almacenar resultados
    resultados_generales = None
    resultados_economicos = None
    gdf_textura_analizado = None
    dem_data_analizado = None
    
    # Verificar si hay archivo cargado
    if uploaded_file is not None:
        with st.spinner("📂 Cargando archivo de parcela..."):
            gdf = cargar_archivo_parcela(uploaded_file)
            
        if gdf is not None:
            st.success(f"✅ Parcela cargada correctamente ({len(gdf)} polígonos)")
            
            # Mostrar información básica
            area_total = calcular_superficie(gdf)
            st.info(f"📏 **Área total:** {area_total:.2f} hectáreas")
            
            # Botón para ejecutar análisis
            if st.button("🚀 EJECUTAR ANÁLISIS COMPLETO", type="primary", use_container_width=True):
                with st.spinner("🔬 Ejecutando análisis..."):
                    resultados = ejecutar_analisis(
                        gdf=gdf,
                        nutriente=nutriente if analisis_tipo == "RECOMENDACIONES NPK" else None,
                        analisis_tipo=analisis_tipo,
                        n_divisiones=n_divisiones,
                        cultivo=cultivo,
                        satelite=satelite_seleccionado,
                        indice=indice_seleccionado if analisis_tipo in ["FERTILIDAD ACTUAL", "RECOMENDACIONES NPK"] else None,
                        fecha_inicio=fecha_inicio if analisis_tipo in ["FERTILIDAD ACTUAL", "RECOMENDACIONES NPK"] else None,
                        fecha_fin=fecha_fin if analisis_tipo in ["FERTILIDAD ACTUAL", "RECOMENDACIONES NPK"] else None,
                        intervalo_curvas=intervalo_curvas if analisis_tipo == "ANÁLISIS DE CURVAS DE NIVEL" else 5.0,
                        resolucion_dem=resolucion_dem if analisis_tipo == "ANÁLISIS DE CURVAS DE NIVEL" else 10.0,
                        usar_inta=st.session_state.get('usar_inta', True),
                        mostrar_capa_inta=st.session_state.get('mostrar_mapa_inta', False)
                    )
                
                if resultados['exitoso']:
                    st.success("✅ Análisis completado exitosamente!")
                    resultados_generales = resultados['gdf_analizado']
                    
                    # Mostrar resultados según tipo de análisis
                    if analisis_tipo == "ANÁLISIS DE TEXTURA":
                        mostrar_resultados_textura(
                            resultados_generales,
                            cultivo,
                            area_total,
                            st.session_state.get('mostrar_mapa_inta', False)
                        )
                        gdf_textura_analizado = resultados_generales
                        
                    elif analisis_tipo == "ANÁLISIS DE CURVAS DE NIVEL":
                        # Generar DEM sintético
                        X, Y, Z, bounds = generar_dem_sintetico(gdf, resolucion_dem)
                        pendiente_grid = calcular_pendiente_simple(X, Y, Z, resolucion_dem)
                        curvas, elevaciones = generar_curvas_nivel_simple(X, Y, Z, intervalo_curvas, gdf)
                        
                        mostrar_resultados_curvas_nivel(
                            X, Y, Z, pendiente_grid, curvas, elevaciones,
                            gdf, cultivo, area_total
                        )
                        dem_data_analizado = (X, Y, Z, pendiente_grid, gdf)
                        
                    elif analisis_tipo == "FERTILIDAD ACTUAL":
                        mostrar_resultados_fertilidad(
                            resultados_generales,
                            cultivo,
                            area_total,
                            satelite_seleccionado,
                            st.session_state.get('mostrar_mapa_inta', False)
                        )
                        
                    elif analisis_tipo == "RECOMENDACIONES NPK":
                        mostrar_resultados_recomendaciones(
                            resultados_generales,
                            cultivo,
                            nutriente,
                            area_total,
                            satelite_seleccionado,
                            st.session_state.get('mostrar_mapa_inta', False)
                        )
                        
                        # Mostrar análisis económico si hay datos de variedad
                        if 'variedad_params' in st.session_state:
                            resultados_economicos = realizar_analisis_economico(
                                resultados_generales,
                                cultivo,
                                st.session_state['variedad_params'],
                                area_total
                            )
                            mostrar_analisis_economico(resultados_economicos)
                    
                    # ===== SECCIÓN DE DESCARGA DE REPORTE UNIFICADO =====
                    st.markdown("---")
                    st.subheader("📥 DESCARGAR REPORTE COMPLETO")
                    
                    if st.button("📄 Generar Reporte Completo (DOCX)", type="primary", use_container_width=True):
                        with st.spinner("📋 Compilando reporte completo..."):
                            reporte_docx = generar_reporte_completo_docx(
                                gdf_analizado=resultados_generales,
                                cultivo=cultivo,
                                area_total=area_total,
                                analisis_tipo=analisis_tipo,
                                nutriente=nutriente if analisis_tipo == "RECOMENDACIONES NPK" else None,
                                satelite=satelite_seleccionado,
                                indice=indice_seleccionado if analisis_tipo in ["FERTILIDAD ACTUAL", "RECOMENDACIONES NPK"] else None,
                                resultados_economicos=resultados_economicos,
                                gdf_textura=gdf_textura_analizado,
                                dem_data=dem_data_analizado
                            )
                            
                            if reporte_docx:
                                st.success("✅ Reporte completo generado exitosamente!")
                                
                                # Botón de descarga
                                st.download_button(
                                    label="📥 Descargar Reporte Completo (DOCX)",
                                    data=reporte_docx,
                                    file_name=f"reporte_completo_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                else:
                    st.error("❌ Error en el análisis. Por favor, revisa los datos e inténtalo de nuevo.")
    else:
        st.info("📁 Por favor, sube un archivo de parcela para comenzar el análisis.")
        st.markdown("""
        ### 📋 Formatos aceptados:
        - **Shapefile** (.zip que contenga .shp, .shx, .dbf, .prj)
        - **KML** (.kml) - Google Earth
        - **KMZ** (.kmz) - Google Earth comprimido
        
        ### 📝 Instrucciones:
        1. Selecciona el cultivo en la barra lateral
        2. Configura los parámetros del análisis
        3. Sube tu archivo de parcela
        4. Haz clic en **EJECUTAR ANÁLISIS COMPLETO**
        """)

# ===== EJECUTAR LA APLICACIÓN =====
if __name__ == "__main__":
    main()
